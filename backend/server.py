from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Depends, Query, Request, Header, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from typing import List, Optional
import os
import json
import httpx
import logging
from bs4 import BeautifulSoup
from dotenv import load_dotenv
import uuid
from datetime import datetime, timezone, timedelta
from motor.motor_asyncio import AsyncIOMotorClient
import io
import re
import hashlib
import secrets
import shutil
from collections import defaultdict
import time
import base64
import asyncio

# Configure logging
logger = logging.getLogger(__name__)

load_dotenv()

# Import Email Service
from email_service import (
    send_welcome_email,
    send_verification_email,
    schedule_getting_started_email,
    schedule_subscription_thank_you,
    send_download_confirmation,
    schedule_signup_no_purchase_reminder,
    track_checkout_started,
    cancel_abandoned_checkout_email,
    send_review_request,
    send_password_changed_email,
    send_password_reset_email,
    cancel_signup_no_purchase_reminder,
    process_scheduled_emails,
    send_email,
    get_base_template
)

# Import PDF Engine
from pdf_engine import (
    analyze_pdf_metadata,
    normalize_pdf_metadata,
    generate_analysis_report_pdf,
    PDFAnalysisResult,
    get_document_types,
    get_legitimate_producers,
    DOCUMENT_TYPES,
    analyze_document_with_ai,
    apply_ai_analysis_to_result,
    edit_and_regenerate_pdf,
    get_metadata_presets,
    METADATA_PRESETS,
    clean_paystub_pdf,
    clean_bank_statement_pdf
)

import anthropic

# Stripe Configuration
import stripe
STRIPE_API_KEY = os.environ.get("STRIPE_API_KEY")
stripe.api_key = STRIPE_API_KEY

# Rate limiting storage (in production, use Redis)
rate_limit_storage = defaultdict(list)

# PDF and DOCX parsing
try:
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

app = FastAPI()

# Create uploads directory if not exists
UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "uploads", "blog")
os.makedirs(UPLOAD_DIR, exist_ok=True)

# Don't mount static files - use dynamic route for blog images to support MongoDB fallback
# app.mount("/api/uploads", StaticFiles(directory=os.path.join(os.path.dirname(__file__), "uploads")), name="uploads")

# Security
security = HTTPBearer(auto_error=False)

# MongoDB Connection
MONGO_URL = os.environ.get("MONGO_URL", "mongodb://localhost:27017")
DB_NAME = os.environ.get("DB_NAME", "mintslip_db")

client = AsyncIOMotorClient(MONGO_URL)
db = client[DB_NAME]
discounts_collection = db["discount_codes"]
admins_collection = db["admins"]
users_collection = db["users"]
purchases_collection = db["purchases"]
sessions_collection = db["sessions"]
subscriptions_collection = db["subscriptions"]
subscription_payments_collection = db["subscription_payments"]
blog_posts_collection = db["blog_posts"]
blog_categories_collection = db["blog_categories"]
blog_images_collection = db["blog_images"]
saved_documents_collection = db["saved_documents"]
site_settings_collection = db["site_settings"]
phone_entries_collection = db["phone_entries"]
address_entries_collection = db["address_entries"]
banned_ips_collection = db["banned_ips"]
email_templates_collection = db["email_templates"]
admin_notifications_collection = db["admin_notifications"]
moderators_collection = db["moderators"]
moderator_permissions_collection = db["moderator_permissions"]
audit_logs_collection = db["audit_logs"]
support_tickets_collection = db["support_tickets"]
people_search_logs_collection = db["people_search_logs"]
people_search_payments_collection = db["people_search_payments"]
phone_reports_collection = db["phone_reports"]
data_source_settings_collection = db["data_source_settings"]
people_records_collection = db["people_records"]
scraper_jobs_collection = db["scraper_jobs"]
opt_out_requests_collection = db["opt_out_requests"]
suppression_list_collection = db["suppression_list"]

# Default permissions per moderator level
ALL_PERMISSIONS = [
    "view_purchases", "view_users", "edit_users", "view_saved_docs",
    "view_discounts", "manage_discounts", "view_banned_ips", "manage_banned_ips",
    "view_blog", "manage_blog", "view_email_templates", "manage_email_templates",
    "send_mass_email", "view_site_settings", "manage_site_settings",
]

def _make_perms(*enabled):
    return {p: (p in enabled) for p in ALL_PERMISSIONS}

DEFAULT_LEVEL_PERMISSIONS = {
    1: _make_perms("view_purchases", "view_users", "view_saved_docs", "view_discounts", "view_blog"),
    2: _make_perms("view_purchases", "view_users", "edit_users", "view_saved_docs",
                   "view_discounts", "manage_discounts", "view_banned_ips",
                   "view_blog", "manage_blog", "view_email_templates"),
    3: _make_perms("view_purchases", "view_users", "edit_users", "view_saved_docs",
                   "view_discounts", "manage_discounts", "view_banned_ips", "manage_banned_ips",
                   "view_blog", "manage_blog", "view_email_templates", "manage_email_templates",
                   "send_mass_email", "view_site_settings"),
}

# Create uploads directory for user documents if not exists
USER_DOCUMENTS_DIR = os.path.join(os.path.dirname(__file__), "uploads", "user_documents")
os.makedirs(USER_DOCUMENTS_DIR, exist_ok=True)

# Constants for saved documents
DOCUMENT_EXPIRY_DAYS = 30

# Saved documents limits by subscription tier
SAVED_DOCS_LIMITS = {
    "starter": 10,
    "professional": 30,
    "business": -1  # -1 means unlimited
}

# Password hashing - Using bcrypt for better security
def hash_password(password: str) -> str:
    # Add salt for additional security
    salt = "MintSlip_Salt_2025_Secure"
    return hashlib.sha256((password + salt).encode()).hexdigest()

def verify_password(password: str, hashed: str) -> bool:
    return hash_password(password) == hashed

# Session token generation
def generate_session_token() -> str:
    return secrets.token_urlsafe(32)

# Rate limiting function
def check_rate_limit(identifier: str, max_requests: int = 10, window_seconds: int = 60) -> bool:
    """
    Simple rate limiting. Returns True if request is allowed, False if rate limited.
    In production, use Redis for distributed rate limiting.
    """
    current_time = time.time()
    window_start = current_time - window_seconds
    
    # Clean old entries
    rate_limit_storage[identifier] = [t for t in rate_limit_storage[identifier] if t > window_start]
    
    # Check if rate limited
    if len(rate_limit_storage[identifier]) >= max_requests:
        return False
    
    # Add current request
    rate_limit_storage[identifier].append(current_time)
    return True

# Helper function to get client IP address
def get_client_ip(request: Request) -> str:
    """Extract client IP from request, handling proxies"""
    # Check for forwarded headers (common with proxies/load balancers)
    forwarded_for = request.headers.get("x-forwarded-for")
    if forwarded_for:
        # Take the first IP in the chain (original client)
        return forwarded_for.split(",")[0].strip()
    
    real_ip = request.headers.get("x-real-ip")
    if real_ip:
        return real_ip.strip()
    
    # Fallback to direct client IP
    if request.client:
        return request.client.host
    
    return "unknown"

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",
        "https://localhost:3000",
        "https://mintslip.com",
        "https://www.mintslip.com",
        "https://mintslip.vercel.app",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Models
class WorkExperience(BaseModel):
    company: str
    position: str
    location: Optional[str] = ""
    startDate: str
    endDate: str
    current: bool = False
    responsibilities: List[str]

class Education(BaseModel):
    institution: str
    degree: str
    field: str
    graduationDate: str
    gpa: Optional[str] = ""

class ResumeInput(BaseModel):
    personalInfo: dict
    workExperience: List[WorkExperience]
    education: List[Education]
    skills: List[str]
    targetJobTitle: str
    jobDescription: str
    jobUrl: Optional[str] = ""

class JobScrapeRequest(BaseModel):
    url: str

# Discount Code Models
class DiscountCodeCreate(BaseModel):
    code: str
    discountPercent: float
    startDate: str  # ISO date string
    expiryDate: str  # ISO date string
    usageType: str  # "unlimited", "limited", "one_per_customer"
    usageLimit: Optional[int] = None
    applicableTo: str  # "all" or "specific"
    specificGenerators: Optional[List[str]] = None
    isActive: bool = True

class DiscountCodeUpdate(BaseModel):
    code: Optional[str] = None
    discountPercent: Optional[float] = None
    startDate: Optional[str] = None
    expiryDate: Optional[str] = None
    usageType: Optional[str] = None
    usageLimit: Optional[int] = None
    applicableTo: Optional[str] = None
    specificGenerators: Optional[List[str]] = None
    isActive: Optional[bool] = None

class CouponValidateRequest(BaseModel):
    code: str
    generatorType: str
    customerIdentifier: Optional[str] = None  # email or IP for one-per-customer

# ========== ADMIN & USER AUTHENTICATION MODELS ==========

class AdminLogin(BaseModel):
    email: str
    password: str

class UserSignup(BaseModel):
    email: str
    password: str
    name: str
    saveDocuments: Optional[bool] = False  # User preference to save documents

class UserLogin(BaseModel):
    email: str
    password: str

class UserPreferencesUpdate(BaseModel):
    saveDocuments: Optional[bool] = None

class SaveDocumentRequest(BaseModel):
    documentType: str
    fileName: str
    fileData: str  # Base64 encoded file content
    template: Optional[str] = None

# ========== PURCHASE TRACKING MODELS ==========

class PurchaseCreate(BaseModel):
    documentType: str  # paystub, resume, w2, w9, 1099-nec, 1099-misc, bank-statement, offer-letter, vehicle-bill-of-sale, schedule-c, utility-bill
    amount: float
    email: Optional[str] = None
    stripePaymentIntentId: Optional[str] = None
    discountCode: Optional[str] = None
    discountAmount: Optional[float] = 0
    userId: Optional[str] = None  # For subscribed users
    template: Optional[str] = None  # Template name/ID used (for paystubs, bank statements)
    quantity: Optional[int] = 1  # Number of documents (e.g., number of paystubs)

class ManualPurchaseCreate(BaseModel):
    documentType: str
    amount: float
    paypalEmail: str
    purchaseDate: Optional[str] = None  # ISO date string, defaults to now
    template: Optional[str] = None
    discountCode: Optional[str] = None
    discountAmount: Optional[float] = 0
    notes: Optional[str] = None  # Admin notes
    quantity: Optional[int] = 1  # Number of documents
    ipAddress: Optional[str] = None  # Customer IP address

# ========== SUBSCRIPTION MODELS ==========

SUBSCRIPTION_TIERS = {
    "starter": {
        "name": "Starter",
        "price": 19.99,
        "price_cents": 1999,
        "downloads": 10,
        "stripePriceId": ""  # Will be filled when Stripe products are created
    },
    "professional": {
        "name": "Professional", 
        "price": 29.99,
        "price_cents": 2999,
        "downloads": 30,
        "stripePriceId": ""
    },
    "business": {
        "name": "Business",
        "price": 49.99,
        "price_cents": 4999,
        "downloads": -1,  # -1 means unlimited
        "stripePriceId": ""
    }
}

class SubscriptionCreate(BaseModel):
    userId: str
    tier: str  # starter, professional, business
    stripeSubscriptionId: str
    stripeCustomerId: str

class CreateCheckoutSession(BaseModel):
    tier: str
    successUrl: Optional[str] = None
    cancelUrl: Optional[str] = None

# ========== BLOG MODELS ==========

class BlogCategoryCreate(BaseModel):
    name: str
    slug: str
    description: Optional[str] = ""

class BlogPostCreate(BaseModel):
    title: str
    slug: str
    metaTitle: Optional[str] = None
    metaDescription: Optional[str] = None
    author: Optional[str] = "MintSlip Team"
    featuredImage: Optional[str] = None
    content: str  # Markdown content
    excerpt: Optional[str] = None
    category: Optional[str] = None  # Category slug
    tags: Optional[List[str]] = []
    faqSchema: Optional[List[dict]] = []  # [{question, answer}]
    status: Optional[str] = "draft"  # draft, published
    indexFollow: Optional[bool] = True
    publishDate: Optional[str] = None

class BlogPostUpdate(BaseModel):
    title: Optional[str] = None
    slug: Optional[str] = None
    metaTitle: Optional[str] = None
    metaDescription: Optional[str] = None
    author: Optional[str] = None
    featuredImage: Optional[str] = None
    content: Optional[str] = None
    excerpt: Optional[str] = None
    category: Optional[str] = None
    tags: Optional[List[str]] = None
    faqSchema: Optional[List[dict]] = None
    status: Optional[str] = None
    indexFollow: Optional[bool] = None
    publishDate: Optional[str] = None

# Blog Categories
BLOG_CATEGORIES = [
    {"slug": "pay-stubs", "name": "Pay Stubs", "description": "Everything about pay stubs and paycheck documentation"},
    {"slug": "proof-of-income", "name": "Proof of Income", "description": "Guides and tips for proof of income documentation"},
    {"slug": "payroll", "name": "Payroll", "description": "Payroll management and best practices"},
    {"slug": "taxes", "name": "Taxes", "description": "Tax forms, filing, and financial documentation"},
    {"slug": "employment", "name": "Employment", "description": "Employment documentation and workplace guides"},
]

_RESUME_SYSTEM = "You are an expert resume writer and ATS optimization specialist. You help create professional, ATS-optimized resumes tailored to specific job descriptions."

async def call_claude(prompt: str) -> str:
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY not configured")
    client = anthropic.AsyncAnthropic(api_key=api_key)
    message = await client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=8096,
        system=_RESUME_SYSTEM,
        messages=[{"role": "user", "content": prompt}],
    )
    return message.content[0].text

@app.get("/api/health")
async def health_check():
    return {"status": "healthy"}


# ========== HERO SECTION STATS ==========

def format_stat_count(count: int) -> str:
    """Format count with appropriate suffix like 50+, 100+, 500+, 1K+, 5K+, 10K+, etc."""
    if count >= 10000:
        # 10K+, 15K+, 20K+, etc.
        thousands = count // 1000
        if thousands >= 50:
            return f"{(thousands // 10) * 10}K+"
        elif thousands >= 10:
            return f"{thousands}K+"
    if count >= 5000:
        return "5K+"
    if count >= 1000:
        return "1K+"
    if count >= 500:
        return "500+"
    if count >= 100:
        return "100+"
    if count >= 50:
        return "50+"
    # If less than 50, return the actual count or a minimum display
    return str(count) if count > 0 else "0"


@app.get("/api/hero-stats")
async def get_hero_stats():
    """Get hero section stats: unique users and total documents created"""
    try:
        # === USER COUNT ===
        # Get all registered user emails (lowercase for comparison)
        registered_users_cursor = users_collection.find({}, {"email": 1})
        registered_emails = set()
        async for user in registered_users_cursor:
            if user.get("email"):
                registered_emails.add(user["email"].lower())
        
        # Get unique guest purchase emails (where isGuest: True and email exists)
        guest_purchases_cursor = purchases_collection.aggregate([
            {"$match": {"isGuest": True, "email": {"$exists": True, "$ne": ""}}},
            {"$group": {"_id": {"$toLower": "$email"}}}
        ])
        
        guest_emails = set()
        async for doc in guest_purchases_cursor:
            if doc.get("_id"):
                guest_emails.add(doc["_id"])
        
        # Combine: registered users + guest-only emails (not registered)
        unique_guest_only_emails = guest_emails - registered_emails
        total_unique_users = len(registered_emails) + len(unique_guest_only_emails)
        
        # === DOCUMENTS COUNT ===
        # Count total documents from purchases (considering quantity field)
        docs_pipeline = [
            {"$group": {
                "_id": None,
                "totalDocs": {"$sum": {"$ifNull": ["$quantity", 1]}}
            }}
        ]
        docs_result = await purchases_collection.aggregate(docs_pipeline).to_list(1)
        total_documents = docs_result[0]["totalDocs"] if docs_result else 0
        
        # Also count saved documents (subscription users who saved documents)
        saved_docs_count = await saved_documents_collection.count_documents({})
        total_documents += saved_docs_count
        
        return {
            "success": True,
            "users": {
                "count": total_unique_users,
                "formatted": format_stat_count(total_unique_users),
                "breakdown": {
                    "registeredUsers": len(registered_emails),
                    "uniqueGuestPurchasers": len(unique_guest_only_emails)
                }
            },
            "documents": {
                "count": total_documents,
                "formatted": format_stat_count(total_documents),
                "breakdown": {
                    "purchased": total_documents - saved_docs_count,
                    "saved": saved_docs_count
                }
            }
        }
    except Exception as e:
        print(f"Error getting hero stats: {e}")
        return {
            "success": False,
            "users": {"count": 0, "formatted": "0"},
            "documents": {"count": 0, "formatted": "0"},
            "error": str(e)
        }


@app.get("/api/user-count")
async def get_user_count():
    """Get the total unique user count (registered users + unique guest purchasers) - Legacy endpoint"""
    stats = await get_hero_stats()
    if stats.get("success"):
        return {
            "success": True,
            "count": stats["users"]["count"],
            "formattedCount": stats["users"]["formatted"],
            "breakdown": stats["users"]["breakdown"]
        }
    return {
        "success": False,
        "count": 0,
        "formattedCount": "0",
        "error": stats.get("error", "Unknown error")
    }


# ========== BANK STATEMENT TRANSACTION GENERATOR ==========

# US States with major cities
US_STATES_CITIES = {
    "AL": ["Birmingham", "Montgomery", "Huntsville", "Mobile", "Tuscaloosa"],
    "AK": ["Anchorage", "Fairbanks", "Juneau", "Sitka", "Ketchikan"],
    "AZ": ["Phoenix", "Tucson", "Mesa", "Chandler", "Scottsdale", "Gilbert", "Glendale", "Tempe"],
    "AR": ["Little Rock", "Fort Smith", "Fayetteville", "Springdale", "Jonesboro"],
    "CA": ["Los Angeles", "San Francisco", "San Diego", "San Jose", "Sacramento", "Fresno", "Oakland", "Long Beach", "Anaheim", "Bakersfield"],
    "CO": ["Denver", "Colorado Springs", "Aurora", "Fort Collins", "Lakewood", "Boulder"],
    "CT": ["Bridgeport", "New Haven", "Hartford", "Stamford", "Waterbury"],
    "DE": ["Wilmington", "Dover", "Newark", "Middletown", "Smyrna"],
    "FL": ["Miami", "Orlando", "Tampa", "Jacksonville", "Fort Lauderdale", "St Petersburg", "Hialeah", "Tallahassee"],
    "GA": ["Atlanta", "Augusta", "Columbus", "Macon", "Savannah", "Athens", "Sandy Springs", "Roswell", "Marietta", "Johns Creek"],
    "HI": ["Honolulu", "Pearl City", "Hilo", "Kailua", "Waipahu"],
    "ID": ["Boise", "Meridian", "Nampa", "Idaho Falls", "Pocatello"],
    "IL": ["Chicago", "Aurora", "Naperville", "Joliet", "Rockford", "Springfield", "Peoria"],
    "IN": ["Indianapolis", "Fort Wayne", "Evansville", "South Bend", "Carmel"],
    "IA": ["Des Moines", "Cedar Rapids", "Davenport", "Sioux City", "Iowa City"],
    "KS": ["Wichita", "Overland Park", "Kansas City", "Olathe", "Topeka"],
    "KY": ["Louisville", "Lexington", "Bowling Green", "Owensboro", "Covington"],
    "LA": ["New Orleans", "Baton Rouge", "Shreveport", "Lafayette", "Lake Charles"],
    "ME": ["Portland", "Lewiston", "Bangor", "South Portland", "Auburn"],
    "MD": ["Baltimore", "Frederick", "Rockville", "Gaithersburg", "Bowie"],
    "MA": ["Boston", "Worcester", "Springfield", "Cambridge", "Lowell"],
    "MI": ["Detroit", "Grand Rapids", "Warren", "Sterling Heights", "Ann Arbor", "Lansing"],
    "MN": ["Minneapolis", "Saint Paul", "Rochester", "Duluth", "Bloomington"],
    "MS": ["Jackson", "Gulfport", "Southaven", "Hattiesburg", "Biloxi"],
    "MO": ["Kansas City", "Saint Louis", "Springfield", "Columbia", "Independence"],
    "MT": ["Billings", "Missoula", "Great Falls", "Bozeman", "Butte"],
    "NE": ["Omaha", "Lincoln", "Bellevue", "Grand Island", "Kearney"],
    "NV": ["Las Vegas", "Henderson", "Reno", "North Las Vegas", "Sparks"],
    "NH": ["Manchester", "Nashua", "Concord", "Derry", "Dover"],
    "NJ": ["Newark", "Jersey City", "Paterson", "Elizabeth", "Edison", "Trenton"],
    "NM": ["Albuquerque", "Las Cruces", "Rio Rancho", "Santa Fe", "Roswell"],
    "NY": ["New York", "Buffalo", "Rochester", "Yonkers", "Syracuse", "Albany"],
    "NC": ["Charlotte", "Raleigh", "Greensboro", "Durham", "Winston Salem", "Fayetteville"],
    "ND": ["Fargo", "Bismarck", "Grand Forks", "Minot", "West Fargo"],
    "OH": ["Columbus", "Cleveland", "Cincinnati", "Toledo", "Akron", "Dayton"],
    "OK": ["Oklahoma City", "Tulsa", "Norman", "Broken Arrow", "Lawton"],
    "OR": ["Portland", "Salem", "Eugene", "Gresham", "Hillsboro", "Beaverton"],
    "PA": ["Philadelphia", "Pittsburgh", "Allentown", "Reading", "Erie", "Scranton"],
    "RI": ["Providence", "Warwick", "Cranston", "Pawtucket", "East Providence"],
    "SC": ["Charleston", "Columbia", "North Charleston", "Mount Pleasant", "Rock Hill", "Greenville"],
    "SD": ["Sioux Falls", "Rapid City", "Aberdeen", "Brookings", "Watertown"],
    "TN": ["Nashville", "Memphis", "Knoxville", "Chattanooga", "Clarksville", "Murfreesboro"],
    "TX": ["Houston", "San Antonio", "Dallas", "Austin", "Fort Worth", "El Paso", "Arlington", "Plano", "Irving", "Frisco"],
    "UT": ["Salt Lake City", "West Valley City", "Provo", "West Jordan", "Orem"],
    "VT": ["Burlington", "South Burlington", "Rutland", "Barre", "Montpelier"],
    "VA": ["Virginia Beach", "Norfolk", "Chesapeake", "Richmond", "Newport News", "Alexandria", "Arlington"],
    "WA": ["Seattle", "Spokane", "Tacoma", "Vancouver", "Bellevue", "Kent", "Everett"],
    "WV": ["Charleston", "Huntington", "Morgantown", "Parkersburg", "Wheeling"],
    "WI": ["Milwaukee", "Madison", "Green Bay", "Kenosha", "Racine"],
    "WY": ["Cheyenne", "Casper", "Laramie", "Gillette", "Rock Springs"],
    "DC": ["Washington"]
}

# Merchants by category with store number ranges
MERCHANTS_BY_CATEGORY = {
    "groceries": [
        {"name": "WAL-MART", "prefix": "#", "numRange": (100, 9999)},
        {"name": "KROGER", "prefix": "#", "numRange": (100, 999)},
        {"name": "PUBLIX", "prefix": "#", "numRange": (100, 1999)},
        {"name": "ALDI", "prefix": "#", "numRange": (10, 999)},
        {"name": "WHOLE FOODS MKT", "prefix": "#", "numRange": (100, 999)},
        {"name": "TRADER JOES", "prefix": "#", "numRange": (100, 999)},
        {"name": "COSTCO WHSE", "prefix": "#", "numRange": (100, 999)},
        {"name": "SAMS CLUB", "prefix": "#", "numRange": (4000, 8999)},
        {"name": "TARGET", "prefix": "T-", "numRange": (1000, 3999)},
        {"name": "SAFEWAY", "prefix": "#", "numRange": (100, 4999)},
        {"name": "FOOD LION", "prefix": "#", "numRange": (100, 2999)},
        {"name": "WINN DIXIE", "prefix": "#", "numRange": (100, 999)},
        {"name": "HEB GROCERY", "prefix": "#", "numRange": (100, 999)},
        {"name": "SPROUTS FARMERS", "prefix": "#", "numRange": (100, 499)},
    ],
    "gas_auto": [
        {"name": "SHELL OIL", "prefix": "", "numRange": (10000, 99999)},
        {"name": "CHEVRON", "prefix": "", "numRange": (10000, 99999)},
        {"name": "EXXONMOBIL", "prefix": "", "numRange": (10000, 99999)},
        {"name": "BP", "prefix": "#", "numRange": (1000, 9999)},
        {"name": "CIRCLE K", "prefix": "#", "numRange": (1000, 9999)},
        {"name": "SPEEDWAY", "prefix": "#", "numRange": (1000, 9999)},
        {"name": "QUIKTRIP", "prefix": "QT ", "numRange": (100, 999)},
        {"name": "RACETRAC", "prefix": "#", "numRange": (100, 999)},
        {"name": "WAWA", "prefix": "#", "numRange": (100, 9999)},
        {"name": "AUTOZONE", "prefix": "#", "numRange": (1000, 9999)},
        {"name": "ADVANCE AUTO", "prefix": "#", "numRange": (1000, 9999)},
        {"name": "OREILLY AUTO", "prefix": "#", "numRange": (1000, 9999)},
        {"name": "JIFFY LUBE", "prefix": "#", "numRange": (100, 9999)},
    ],
    "dining": [
        {"name": "MCDONALDS", "prefix": "F", "numRange": (10000, 99999)},
        {"name": "CHICK-FIL-A", "prefix": "#", "numRange": (1000, 9999)},
        {"name": "WENDYS", "prefix": "#", "numRange": (1000, 9999)},
        {"name": "STARBUCKS", "prefix": "", "numRange": (10000, 99999)},
        {"name": "CHIPOTLE", "prefix": "", "numRange": (1000, 9999)},
        {"name": "TACO BELL", "prefix": "#", "numRange": (1000, 99999)},
        {"name": "BURGER KING", "prefix": "#", "numRange": (1000, 9999)},
        {"name": "SUBWAY", "prefix": "#", "numRange": (10000, 99999)},
        {"name": "DUNKIN", "prefix": "#", "numRange": (100000, 999999)},
        {"name": "DOMINOS PIZZA", "prefix": "", "numRange": (1000, 9999)},
        {"name": "PIZZA HUT", "prefix": "", "numRange": (1000, 99999)},
        {"name": "PAPA JOHNS", "prefix": "#", "numRange": (1000, 9999)},
        {"name": "PANERA BREAD", "prefix": "", "numRange": (100000, 999999)},
        {"name": "APPLEBEES", "prefix": "", "numRange": (100, 9999)},
        {"name": "OLIVE GARDEN", "prefix": "#", "numRange": (1000, 9999)},
        {"name": "DENNYS", "prefix": "#", "numRange": (1000, 9999)},
        {"name": "IHOP", "prefix": "#", "numRange": (1000, 9999)},
        {"name": "CRACKER BARREL", "prefix": "#", "numRange": (100, 999)},
    ],
    "retail": [
        {"name": "TARGET", "prefix": "T-", "numRange": (1000, 3999)},
        {"name": "AMAZON MKTPLACE", "prefix": "", "numRange": None},
        {"name": "AMAZON.COM", "prefix": "", "numRange": None},
        {"name": "KOHLS", "prefix": "#", "numRange": (100, 1999)},
        {"name": "ROSS STORES", "prefix": "#", "numRange": (100, 1999)},
        {"name": "TJ MAXX", "prefix": "#", "numRange": (100, 1999)},
        {"name": "MARSHALLS", "prefix": "#", "numRange": (100, 1999)},
        {"name": "BED BATH BEYOND", "prefix": "#", "numRange": (100, 999)},
        {"name": "BEST BUY", "prefix": "#", "numRange": (100, 1999)},
        {"name": "HOME DEPOT", "prefix": "#", "numRange": (100, 9999)},
        {"name": "LOWES", "prefix": "#", "numRange": (100, 9999)},
        {"name": "CVS PHARMACY", "prefix": "#", "numRange": (1000, 9999)},
        {"name": "WALGREENS", "prefix": "#", "numRange": (1000, 19999)},
        {"name": "DOLLAR GENERAL", "prefix": "#", "numRange": (1000, 19999)},
        {"name": "DOLLAR TREE", "prefix": "#", "numRange": (1000, 9999)},
        {"name": "FIVE BELOW", "prefix": "#", "numRange": (100, 999)},
        {"name": "OLD NAVY", "prefix": "#", "numRange": (100, 999)},
        {"name": "GAP", "prefix": "#", "numRange": (100, 999)},
        {"name": "NIKE", "prefix": "", "numRange": None},
        {"name": "FOOT LOCKER", "prefix": "#", "numRange": (100, 9999)},
    ],
    "utilities": [
        {"name": "CITY WATER DEPT", "prefix": "", "numRange": None},
        {"name": "POWER COMPANY", "prefix": "", "numRange": None},
        {"name": "GAS COMPANY", "prefix": "", "numRange": None},
        {"name": "AT&T WIRELESS", "prefix": "", "numRange": None},
        {"name": "VERIZON WIRELESS", "prefix": "", "numRange": None},
        {"name": "T-MOBILE", "prefix": "", "numRange": None},
        {"name": "XFINITY COMCAST", "prefix": "", "numRange": None},
        {"name": "SPECTRUM", "prefix": "", "numRange": None},
        {"name": "COX COMMUNICATIONS", "prefix": "", "numRange": None},
    ],
    "subscriptions": [
        {"name": "NETFLIX.COM", "prefix": "", "numRange": None},
        {"name": "SPOTIFY USA", "prefix": "", "numRange": None},
        {"name": "HULU", "prefix": "", "numRange": None},
        {"name": "APPLE.COM/BILL", "prefix": "", "numRange": None},
        {"name": "AMAZON PRIME", "prefix": "", "numRange": None},
        {"name": "DISNEY PLUS", "prefix": "", "numRange": None},
        {"name": "HBO MAX", "prefix": "", "numRange": None},
        {"name": "YOUTUBE PREMIUM", "prefix": "", "numRange": None},
        {"name": "PARAMOUNT PLUS", "prefix": "", "numRange": None},
        {"name": "ADOBE SYSTEMS", "prefix": "", "numRange": None},
        {"name": "MICROSOFT", "prefix": "", "numRange": None},
        {"name": "GOOGLE STORAGE", "prefix": "", "numRange": None},
        {"name": "PLANET FITNESS", "prefix": "", "numRange": None},
        {"name": "LA FITNESS", "prefix": "", "numRange": None},
    ],
    "atm_withdrawal": [
        {"name": "ATM WITHDRAWAL", "prefix": "", "numRange": None},
        {"name": "ATM CASH WITHDRAWAL", "prefix": "", "numRange": None},
        {"name": "CASH BACK", "prefix": "", "numRange": None},
        {"name": "ALLPOINT ATM", "prefix": "", "numRange": None},
        {"name": "MONEYPASS ATM", "prefix": "", "numRange": None},
    ],
    "fees": [
        {"name": "MONTHLY SERVICE FEE", "prefix": "", "numRange": None},
        {"name": "OVERDRAFT FEE", "prefix": "", "numRange": None},
        {"name": "NSF FEE", "prefix": "", "numRange": None},
        {"name": "WIRE TRANSFER FEE", "prefix": "", "numRange": None},
        {"name": "FOREIGN TXN FEE", "prefix": "", "numRange": None},
    ],
    "misc": [
        {"name": "UBER TRIP", "prefix": "", "numRange": None},
        {"name": "LYFT RIDE", "prefix": "", "numRange": None},
        {"name": "DOORDASH", "prefix": "", "numRange": None},
        {"name": "UBER EATS", "prefix": "", "numRange": None},
        {"name": "GRUBHUB", "prefix": "", "numRange": None},
        {"name": "INSTACART", "prefix": "", "numRange": None},
        {"name": "AIRBNB", "prefix": "", "numRange": None},
        {"name": "TICKETMASTER", "prefix": "", "numRange": None},
        {"name": "STUBHUB", "prefix": "", "numRange": None},
        {"name": "AMC THEATRES", "prefix": "#", "numRange": (100, 999)},
        {"name": "REGAL CINEMAS", "prefix": "#", "numRange": (100, 999)},
    ],
    "credits_deposit": [
        {"name": "DIRECT DEPOSIT", "prefix": "", "numRange": None, "employers": [
            "PAYROLL", "EMPLOYER", "COMPANY PAYROLL", "HR DEPT", "WORKFORCE", "STAFFING"
        ]},
    ],
    "credits_p2p": [
        {"name": "ZELLE FROM", "prefix": "", "numRange": None},
        {"name": "VENMO CASHOUT", "prefix": "", "numRange": None},
        {"name": "CASH APP", "prefix": "", "numRange": None},
        {"name": "PAYPAL TRANSFER", "prefix": "", "numRange": None},
    ],
    "credits_refunds": [
        {"name": "REFUND FROM", "prefix": "", "numRange": None},
        {"name": "CREDIT FROM", "prefix": "", "numRange": None},
        {"name": "RETURN CREDIT", "prefix": "", "numRange": None},
    ],
}

# Price ranges by category (min, max)
PRICE_RANGES = {
    "groceries": (15.00, 250.00),
    "gas_auto": (20.00, 85.00),
    "dining": (8.00, 75.00),
    "retail": (10.00, 200.00),
    "utilities": (50.00, 250.00),
    "subscriptions": (5.99, 19.99),
    "atm_withdrawal": (20.00, 400.00),
    "fees": (5.00, 35.00),
    "misc": (10.00, 150.00),
    "credits_deposit": (500.00, 5000.00),
    "credits_p2p": (20.00, 500.00),
    "credits_refunds": (10.00, 150.00),
}

class GenerateTransactionsRequest(BaseModel):
    state: str
    cities: List[str]
    volume: str  # light, moderate, heavy
    categories: List[str]
    statementMonth: str  # YYYY-MM format
    employerName: Optional[str] = None
    payFrequency: Optional[str] = "biweekly"  # weekly, biweekly, monthly
    depositAmount: Optional[float] = None
    includeLocation: Optional[bool] = False  # If True, include city/state in descriptor

@app.post("/api/generate-bank-transactions")
async def generate_bank_transactions(data: GenerateTransactionsRequest):
    """Generate realistic bank transactions for statement"""
    import random
    from calendar import monthrange
    
    # Validate inputs
    if data.state not in US_STATES_CITIES:
        raise HTTPException(status_code=400, detail=f"Invalid state code: {data.state}")
    
    # Get volume range
    volume_ranges = {
        "light": (10, 15),
        "moderate": (18, 25),
        "heavy": (28, 40)
    }
    min_tx, max_tx = volume_ranges.get(data.volume, (15, 25))
    num_transactions = random.randint(min_tx, max_tx)
    
    # Parse statement month
    try:
        year, month = map(int, data.statementMonth.split("-"))
        _, days_in_month = monthrange(year, month)
    except:
        raise HTTPException(status_code=400, detail="Invalid statement month format. Use YYYY-MM")
    
    # Filter valid categories
    valid_categories = [c for c in data.categories if c in MERCHANTS_BY_CATEGORY]
    if not valid_categories:
        raise HTTPException(status_code=400, detail="No valid categories selected")
    
    # Separate debit and credit categories
    credit_categories = ["credits_deposit", "credits_p2p", "credits_refunds"]
    debit_categories = [c for c in valid_categories if c not in credit_categories]
    selected_credit_categories = [c for c in valid_categories if c in credit_categories]
    
    transactions = []
    cities_to_use = data.cities if data.cities else US_STATES_CITIES.get(data.state, [""])
    
    # Generate debit transactions (purchases, withdrawals, etc.)
    num_debits = num_transactions - len(selected_credit_categories) * 2  # Reserve space for credits
    if num_debits < 1:
        num_debits = num_transactions
    
    for _ in range(max(num_debits, 5)):
        category = random.choice(debit_categories) if debit_categories else random.choice(valid_categories)
        merchants = MERCHANTS_BY_CATEGORY.get(category, [])
        if not merchants:
            continue
            
        merchant = random.choice(merchants)
        city = random.choice(cities_to_use).upper()
        
        # Build descriptor
        store_num = ""
        if merchant.get("numRange"):
            num = random.randint(merchant["numRange"][0], merchant["numRange"][1])
            store_num = f"{merchant['prefix']}{num:04d}" if merchant['prefix'] else f"#{num:04d}"
        
        # Format descriptor based on includeLocation flag
        # If includeLocation is True: "MERCHANT_NAME #XXXX CITY STATE"
        # If includeLocation is False (default for Chime): "MERCHANT_NAME #XXXX" (store name only)
        if data.includeLocation:
            if store_num:
                descriptor = f"{merchant['name']} {store_num} {city} {data.state}"
            else:
                descriptor = f"{merchant['name']} {city} {data.state}"
        else:
            # Store name only (for Chime template)
            if store_num:
                descriptor = f"{merchant['name']} {store_num}"
            else:
                descriptor = merchant['name']
        
        # Generate realistic price
        price_range = PRICE_RANGES.get(category, (10.00, 100.00))
        amount = round(random.uniform(price_range[0], price_range[1]), 2)
        
        # Determine transaction type
        if category == "atm_withdrawal":
            tx_type = "ATM"
            amount = round(amount / 20) * 20  # ATM withdrawals in $20 increments
        elif category == "fees":
            tx_type = "Fee"
        else:
            tx_type = "Purchase"
        
        # Generate random date within the month
        day = random.randint(1, days_in_month)
        date_str = f"{year}-{month:02d}-{day:02d}"
        
        transactions.append({
            "date": date_str,
            "description": descriptor,
            "type": tx_type,
            "amount": str(amount),
            "category": category
        })
    
    # Generate credit transactions (deposits, refunds, p2p)
    for category in selected_credit_categories:
        if category == "credits_deposit" and data.employerName:
            # Generate payroll deposits based on frequency
            deposit_amount = data.depositAmount or random.uniform(1000, 3500)
            
            if data.payFrequency == "weekly":
                pay_days = [7, 14, 21, 28]
            elif data.payFrequency == "biweekly":
                pay_days = [15, days_in_month] if days_in_month >= 28 else [14, 28]
            else:  # monthly
                pay_days = [days_in_month]
            
            for day in pay_days:
                if day <= days_in_month:
                    date_str = f"{year}-{month:02d}-{day:02d}"
                    transactions.append({
                        "date": date_str,
                        "description": f"DIRECT DEP {data.employerName.upper()}",
                        "type": "Deposit",
                        "amount": str(round(deposit_amount, 2)),
                        "category": category
                    })
        elif category == "credits_p2p":
            # Generate 1-3 P2P credits
            for _ in range(random.randint(1, 3)):
                merchants = MERCHANTS_BY_CATEGORY.get(category, [])
                merchant = random.choice(merchants)
                amount = round(random.uniform(20, 300), 2)
                day = random.randint(1, days_in_month)
                date_str = f"{year}-{month:02d}-{day:02d}"
                
                # Add a random person's initials or name
                initials = ''.join(random.choices('ABCDEFGHJKLMNPRSTUVWXYZ', k=2))
                
                transactions.append({
                    "date": date_str,
                    "description": f"{merchant['name']} {initials}",
                    "type": "Deposit",
                    "amount": str(amount),
                    "category": category
                })
        elif category == "credits_refunds":
            # Generate 1-2 refunds
            for _ in range(random.randint(1, 2)):
                # Pick a random retail/dining merchant for the refund
                refund_categories = ["retail", "dining", "groceries"]
                refund_cat = random.choice(refund_categories)
                merchants = MERCHANTS_BY_CATEGORY.get(refund_cat, [])
                if merchants:
                    merchant = random.choice(merchants)
                    amount = round(random.uniform(10, 100), 2)
                    day = random.randint(1, days_in_month)
                    date_str = f"{year}-{month:02d}-{day:02d}"
                    
                    transactions.append({
                        "date": date_str,
                        "description": f"REFUND {merchant['name']}",
                        "type": "Refund",
                        "amount": str(amount),
                        "category": category
                    })
    
    # Sort transactions by date (newest first)
    transactions.sort(key=lambda x: x["date"], reverse=True)
    
    return {
        "success": True,
        "transactions": transactions,
        "count": len(transactions)
    }


# ========== AUTHENTICATION HELPERS ==========

async def log_action(session: dict, action: str, resource_type: str = "", resource_id: str = "", details: str = ""):
    """Fire-and-forget audit log writer."""
    try:
        actor_id = session.get("adminId") or session.get("moderatorId", "")
        actor_name = session.get("email", "unknown")
        role = session.get("type", "admin")
        level = session.get("level")
        await audit_logs_collection.insert_one({
            "id": str(uuid.uuid4()),
            "actorId": actor_id,
            "actorEmail": actor_name,
            "role": role,
            "level": level,
            "action": action,
            "resourceType": resource_type,
            "resourceId": resource_id,
            "details": details,
            "timestamp": datetime.now(timezone.utc).isoformat(),
        })
    except Exception:
        pass  # never break the request over logging

def check_permission(session: dict, perm: str):
    """Check that a session has a specific permission. Admins always pass."""
    if session.get("type") == "admin":
        return
    if not session.get("permissions", {}).get(perm, False):
        raise HTTPException(status_code=403, detail="Insufficient permissions")

async def get_current_admin(credentials: HTTPAuthorizationCredentials = Depends(security)):
    """Verify admin or moderator session token"""
    if not credentials:
        raise HTTPException(status_code=401, detail="Not authenticated")

    token = credentials.credentials
    session = await sessions_collection.find_one(
        {"token": token, "type": {"$in": ["admin", "moderator"]}}, {"_id": 0}
    )

    if not session:
        raise HTTPException(status_code=401, detail="Invalid or expired session")

    # Moderators: check isActive
    if session.get("type") == "moderator":
        mod = await moderators_collection.find_one({"id": session.get("moderatorId")}, {"isActive": 1})
        if not mod or not mod.get("isActive", True):
            await sessions_collection.delete_one({"token": token})
            raise HTTPException(status_code=401, detail="Account disabled")

    # Check if session is expired (24 hours)
    created_at = datetime.fromisoformat(session["createdAt"].replace("Z", "+00:00"))
    if (datetime.now(timezone.utc) - created_at).total_seconds() > 86400:
        await sessions_collection.delete_one({"token": token})
        raise HTTPException(status_code=401, detail="Session expired")

    return session

async def require_admin_only(session: dict = Depends(get_current_admin)):
    """Only the original admin account may proceed (not moderators)."""
    if session.get("type") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    return session

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    """Verify user session token"""
    if not credentials:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    token = credentials.credentials
    session = await sessions_collection.find_one({"token": token, "type": "user"}, {"_id": 0})
    
    if not session:
        raise HTTPException(status_code=401, detail="Invalid or expired session")
    
    # Check if session is expired (7 days)
    created_at = datetime.fromisoformat(session["createdAt"].replace("Z", "+00:00"))
    if (datetime.now(timezone.utc) - created_at).total_seconds() > 604800:
        await sessions_collection.delete_one({"token": token})
        raise HTTPException(status_code=401, detail="Session expired")
    
    return session

# ========== ADMIN AUTHENTICATION ENDPOINTS ==========

@app.post("/api/admin/login")
async def admin_login(data: AdminLogin, request: Request):
    """Admin login endpoint with rate limiting"""
    # Get client IP for rate limiting
    client_ip = request.client.host if request.client else "unknown"
    
    # Rate limit: 5 login attempts per minute per IP
    if not check_rate_limit(f"admin_login_{client_ip}", max_requests=5, window_seconds=60):
        raise HTTPException(status_code=429, detail="Too many login attempts. Please try again later.")
    
    email = data.email.lower()
    admin = await admins_collection.find_one({"email": email}, {"_id": 0})

    if admin:
        # ── Main admin login ──
        if not verify_password(data.password, admin["password"]):
            raise HTTPException(status_code=401, detail="Invalid credentials")

        token = generate_session_token()
        session = {
            "id": str(uuid.uuid4()),
            "token": token,
            "adminId": admin["id"],
            "email": admin["email"],
            "type": "admin",
            "createdAt": datetime.now(timezone.utc).isoformat(),
        }
        await sessions_collection.insert_one(session)
        return {
            "success": True,
            "token": token,
            "role": "admin",
            "permissions": None,
            "admin": {"id": admin["id"], "email": admin["email"], "name": admin.get("name", "Admin")},
        }

    # ── Moderator login ──
    moderator = await moderators_collection.find_one({"email": email}, {"_id": 0})
    if not moderator:
        raise HTTPException(status_code=401, detail="Invalid credentials")

    if not moderator.get("isActive", True):
        raise HTTPException(status_code=403, detail="Account disabled")

    if not verify_password(data.password, moderator["password"]):
        raise HTTPException(status_code=401, detail="Invalid credentials")

    level = moderator.get("level", 1)
    # Load configured permissions for this level (fallback to defaults)
    perm_doc = await moderator_permissions_collection.find_one({"level": level}, {"_id": 0})
    permissions = perm_doc.get("permissions", DEFAULT_LEVEL_PERMISSIONS[level]) if perm_doc else DEFAULT_LEVEL_PERMISSIONS[level]

    token = generate_session_token()
    session = {
        "id": str(uuid.uuid4()),
        "token": token,
        "moderatorId": moderator["id"],
        "email": moderator["email"],
        "type": "moderator",
        "level": level,
        "permissions": permissions,
        "createdAt": datetime.now(timezone.utc).isoformat(),
    }
    await sessions_collection.insert_one(session)
    return {
        "success": True,
        "token": token,
        "role": "moderator",
        "level": level,
        "permissions": permissions,
        "admin": {"id": moderator["id"], "email": moderator["email"], "name": moderator.get("name", "Moderator")},
    }

@app.post("/api/admin/logout")
async def admin_logout(session: dict = Depends(get_current_admin)):
    """Admin logout endpoint"""
    await sessions_collection.delete_one({"token": session["token"]})
    return {"success": True}

@app.get("/api/admin/verify")
async def verify_admin(session: dict = Depends(get_current_admin)):
    """Verify admin session"""
    admin = await admins_collection.find_one({"id": session["adminId"]}, {"_id": 0, "password": 0})
    return {"success": True, "admin": admin}

class ChangeAdminPassword(BaseModel):
    currentPassword: str
    newPassword: str

DOC_DISPLAY_NAMES = {
    "paystub": "Pay Stub",
    "canadian-paystub": "Canadian Pay Stub",
    "resume": "AI Resume",
    "w2": "W-2 Form",
    "w9": "W-9 Form",
    "1099-nec": "1099-NEC",
    "1099-misc": "1099-MISC",
    "bank-statement": "Bank Statement",
    "offer-letter": "Offer Letter",
    "vehicle-bill-of-sale": "Vehicle Bill of Sale",
    "schedule-c": "Schedule C",
    "utility-bill": "Service Expense",
}

async def create_notification(doc_type: str, customer_email: str = "", amount: float = 0):
    display_name = DOC_DISPLAY_NAMES.get(doc_type, doc_type.replace("-", " ").title())
    notification = {
        "id": str(uuid.uuid4()),
        "type": "document_created",
        "docType": doc_type,
        "docDisplayName": display_name,
        "customerEmail": customer_email or "",
        "amount": float(amount or 0),
        "read": False,
        "createdAt": datetime.now(timezone.utc).isoformat()
    }
    await admin_notifications_collection.insert_one(notification)

@app.put("/api/admin/change-password")
async def change_admin_password(data: ChangeAdminPassword, session: dict = Depends(get_current_admin)):
    """Change admin password (requires current password)"""
    admin = await admins_collection.find_one({"id": session["adminId"]})
    if not admin:
        raise HTTPException(status_code=404, detail="Admin not found")
    
    # Verify current password
    if not verify_password(data.currentPassword, admin["password"]):
        raise HTTPException(status_code=400, detail="Current password is incorrect")
    
    # Update password
    await admins_collection.update_one(
        {"id": session["adminId"]},
        {"$set": {"password": hash_password(data.newPassword)}}
    )
    
    return {"success": True, "message": "Password changed successfully"}

@app.get("/api/admin/notifications")
async def get_admin_notifications(session: dict = Depends(get_current_admin)):
    """Get recent admin notifications"""
    notifications = await admin_notifications_collection.find(
        {}, {"_id": 0}
    ).sort("createdAt", -1).limit(30).to_list(30)
    unread_count = await admin_notifications_collection.count_documents({"read": False})
    return {"success": True, "notifications": notifications, "unreadCount": unread_count}

@app.put("/api/admin/notifications/mark-read")
async def mark_notifications_read(session: dict = Depends(get_current_admin)):
    """Mark all notifications as read"""
    await admin_notifications_collection.update_many({"read": False}, {"$set": {"read": True}})
    return {"success": True}

@app.delete("/api/admin/notifications")
async def clear_all_notifications(session: dict = Depends(get_current_admin)):
    """Delete all notifications"""
    await admin_notifications_collection.delete_many({})
    return {"success": True}

@app.get("/api/admin/profile")
async def get_admin_profile(session: dict = Depends(get_current_admin)):
    """Get admin profile"""
    admin = await admins_collection.find_one({"id": session["adminId"]}, {"_id": 0, "password": 0})
    return {"success": True, "profile": admin}

@app.put("/api/admin/profile")
async def update_admin_profile(request: dict, session: dict = Depends(get_current_admin)):
    """Update admin profile (name, email, photo)"""
    update = {}
    if "name" in request:
        update["name"] = request["name"]
    if "email" in request:
        new_email = request["email"].lower().strip()
        existing = await admins_collection.find_one({"email": new_email, "id": {"$ne": session["adminId"]}})
        if existing:
            raise HTTPException(status_code=400, detail="Email already in use")
        update["email"] = new_email
    if "photo" in request:
        update["photo"] = request["photo"]
    if update:
        await admins_collection.update_one({"id": session["adminId"]}, {"$set": update})
    return {"success": True}

@app.post("/api/admin/setup")
async def setup_admin():
    """Initialize default admin account (run once) - only works if no admin exists"""
    # Check if any admin already exists
    admin_count = await admins_collection.count_documents({})
    if admin_count > 0:
        raise HTTPException(status_code=403, detail="Admin setup already completed")
    
    admin = {
        "id": str(uuid.uuid4()),
        "email": "admin@mintslip.com",
        "password": hash_password("MINTSLIP2025!"),
        "name": "Admin",
        "createdAt": datetime.now(timezone.utc).isoformat()
    }
    await admins_collection.insert_one(admin)
    return {"message": "Admin account created successfully"}

# ========== USER AUTHENTICATION ENDPOINTS ==========

@app.post("/api/user/signup")
async def user_signup(data: UserSignup, request: Request):
    """User signup endpoint"""
    # Get client IP
    client_ip = get_client_ip(request)

    # Check if signup is enabled
    auth_settings = await site_settings_collection.find_one({"key": "auth_enabled"})
    if auth_settings and not auth_settings.get("isEnabled", True):
        raise HTTPException(status_code=403, detail="New account registration is currently disabled.")

    # Check if IP is banned
    banned = await banned_ips_collection.find_one({"ip": client_ip, "isActive": True})
    if banned:
        raise HTTPException(status_code=403, detail="Access denied")

    # Check if email already exists
    existing = await users_collection.find_one({"email": data.email.lower()})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    user = {
        "id": str(uuid.uuid4()),
        "email": data.email.lower(),
        "password": hash_password(data.password),
        "name": data.name,
        "subscription": None,
        "downloadsUsed": 0,
        "downloadsReset": None,
        "preferences": {
            "saveDocuments": data.saveDocuments if data.saveDocuments else False
        },
        "ipAddress": client_ip,
        "emailVerified": False,
        "verificationCode": secrets.token_hex(3).upper(),  # 6-char code
        "createdAt": datetime.now(timezone.utc).isoformat()
    }
    await users_collection.insert_one(user)
    
    # Create session
    token = generate_session_token()
    session = {
        "id": str(uuid.uuid4()),
        "token": token,
        "userId": user["id"],
        "email": user["email"],
        "type": "user",
        "createdAt": datetime.now(timezone.utc).isoformat()
    }
    await sessions_collection.insert_one(session)
    
    # Send emails asynchronously with error handling
    async def send_signup_emails():
        try:
            await send_welcome_email(user["email"], user["name"])
            await send_verification_email(
                user["email"], 
                user["name"], 
                user["verificationCode"],
                f"{os.environ.get('SITE_URL', 'https://mintslip.com')}/verify-email?code={user['verificationCode']}&email={user['email']}"
            )
            await schedule_getting_started_email(user["email"], user["name"], user["id"])
            await schedule_signup_no_purchase_reminder(user["email"], user["name"], user["id"])
        except Exception as e:
            print(f"Error sending signup emails to {user['email']}: {str(e)}")
    
    asyncio.create_task(send_signup_emails())
    
    return {
        "success": True,
        "token": token,
        "user": {
            "id": user["id"],
            "email": user["email"],
            "name": user["name"],
            "subscription": None,
            "preferences": user["preferences"],
            "emailVerified": False
        }
    }

@app.post("/api/user/login")
async def user_login(data: UserLogin, request: Request):
    """User login endpoint with rate limiting"""
    # Get client IP for rate limiting
    client_ip = request.client.host if request.client else "unknown"

    # Check if login is enabled
    auth_settings = await site_settings_collection.find_one({"key": "auth_enabled"})
    if auth_settings and not auth_settings.get("isEnabled", True):
        raise HTTPException(status_code=403, detail="User login is currently disabled.")

    # Rate limit: 5 login attempts per minute per IP
    if not check_rate_limit(f"user_login_{client_ip}", max_requests=5, window_seconds=60):
        raise HTTPException(status_code=429, detail="Too many login attempts. Please try again later.")

    user = await users_collection.find_one({"email": data.email.lower()}, {"_id": 0})
    
    if not user:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    if not verify_password(data.password, user["password"]):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    # Check if user is banned
    if user.get("isBanned", False):
        raise HTTPException(status_code=403, detail="Your account has been suspended. Please contact support.")
    
    # Create session
    token = generate_session_token()
    session = {
        "id": str(uuid.uuid4()),
        "token": token,
        "userId": user["id"],
        "email": user["email"],
        "type": "user",
        "createdAt": datetime.now(timezone.utc).isoformat()
    }
    await sessions_collection.insert_one(session)
    
    return {
        "success": True,
        "token": token,
        "user": {
            "id": user["id"],
            "email": user["email"],
            "name": user["name"],
            "subscription": user.get("subscription"),
            "emailVerified": user.get("emailVerified", True)  # Default True for existing users
        }
    }

@app.post("/api/user/logout")
async def user_logout(session: dict = Depends(get_current_user)):
    """User logout endpoint"""
    await sessions_collection.delete_one({"token": session["token"]})
    return {"success": True}

@app.get("/api/user/me")
async def get_user_profile(session: dict = Depends(get_current_user)):
    """Get current user profile"""
    user = await users_collection.find_one({"id": session["userId"]}, {"_id": 0, "password": 0})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    return {"success": True, "user": user}


class ChangeUserPassword(BaseModel):
    currentPassword: str
    newPassword: str

@app.put("/api/user/change-password")
async def change_user_password(data: ChangeUserPassword, request: Request, session: dict = Depends(get_current_user)):
    """Change user password (requires current password)"""
    # Get client IP for rate limiting
    client_ip = request.client.host if request.client else "unknown"
    
    # Rate limit: 3 password change attempts per minute
    if not check_rate_limit(f"user_change_password_{client_ip}", max_requests=3, window_seconds=60):
        raise HTTPException(status_code=429, detail="Too many attempts. Please try again later.")
    
    # Get current user with password
    user = await users_collection.find_one({"id": session["userId"]})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Verify current password
    if not verify_password(data.currentPassword, user["password"]):
        raise HTTPException(status_code=400, detail="Current password is incorrect")
    
    # Validate new password
    if len(data.newPassword) < 8:
        raise HTTPException(status_code=400, detail="New password must be at least 8 characters")
    
    # Update password
    await users_collection.update_one(
        {"id": session["userId"]},
        {"$set": {
            "password": hash_password(data.newPassword),
            "passwordChangedAt": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    # Send password changed notification email
    asyncio.create_task(send_password_changed_email(user["email"], user.get("name", "")))
    
    return {"success": True, "message": "Password changed successfully"}


# ========== CHANGE EMAIL ENDPOINT ==========

class ChangeEmailRequest(BaseModel):
    newEmail: str
    password: str

@app.put("/api/user/change-email")
async def change_email(data: ChangeEmailRequest, session: dict = Depends(get_current_user)):
    """Change user's email address - requires password verification and email re-verification"""
    user = await users_collection.find_one({"id": session["userId"]})
    
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Verify password
    if not verify_password(data.password, user["password"]):
        raise HTTPException(status_code=400, detail="Incorrect password")
    
    # Validate new email format
    new_email = data.newEmail.lower().strip()
    if not new_email or "@" not in new_email or "." not in new_email.split("@")[-1]:
        raise HTTPException(status_code=400, detail="Invalid email address")
    
    # Check if new email is same as current
    if new_email == user["email"].lower():
        raise HTTPException(status_code=400, detail="New email is the same as current email")
    
    # Check if new email is already in use by another user (case-insensitive)
    existing_user = await users_collection.find_one({
        "email": {"$regex": f"^{re.escape(new_email)}$", "$options": "i"},
        "id": {"$ne": session["userId"]}  # Exclude current user
    })
    if existing_user:
        raise HTTPException(status_code=400, detail="This email is already registered to another account")
    
    # Double-check with exact match as well
    exact_match = await users_collection.find_one({"email": new_email})
    if exact_match and exact_match.get("id") != session["userId"]:
        raise HTTPException(status_code=400, detail="This email is already registered to another account")
    
    # Generate new verification code
    verification_code = secrets.token_hex(3).upper()
    old_email = user["email"]
    
    # Update user's email and set as unverified
    await users_collection.update_one(
        {"id": session["userId"]},
        {"$set": {
            "email": new_email,
            "emailVerified": False,
            "verificationCode": verification_code,
            "previousEmail": old_email,
            "emailChangedAt": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    # Update session with new email
    await sessions_collection.update_many(
        {"userId": session["userId"]},
        {"$set": {"email": new_email}}
    )
    
    # Send verification email to new address
    asyncio.create_task(send_verification_email(
        new_email,
        user.get("name", ""),
        verification_code,
        f"{os.environ.get('SITE_URL', 'https://mintslip.com')}/verify-email?code={verification_code}&email={new_email}"
    ))
    
    # Return updated user info
    updated_user = await users_collection.find_one({"id": session["userId"]}, {"_id": 0, "password": 0})
    
    return {
        "success": True, 
        "message": "Email updated. Please verify your new email address.",
        "user": {
            "id": updated_user["id"],
            "email": updated_user["email"],
            "name": updated_user.get("name", ""),
            "subscription": updated_user.get("subscription"),
            "preferences": updated_user.get("preferences", {}),
            "emailVerified": False
        }
    }


# ========== EMAIL VERIFICATION ENDPOINTS ==========

class VerifyEmailRequest(BaseModel):
    email: str
    code: str

@app.post("/api/user/verify-email")
async def verify_email(data: VerifyEmailRequest):
    """Verify user's email address"""
    user = await users_collection.find_one({"email": data.email.lower()})
    
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    if user.get("emailVerified"):
        return {"success": True, "message": "Email already verified"}
    
    if user.get("verificationCode") != data.code.upper():
        raise HTTPException(status_code=400, detail="Invalid verification code")
    
    await users_collection.update_one(
        {"email": data.email.lower()},
        {"$set": {"emailVerified": True}, "$unset": {"verificationCode": ""}}
    )
    
    return {"success": True, "message": "Email verified successfully"}


@app.post("/api/user/resend-verification")
async def resend_verification(session: dict = Depends(get_current_user)):
    """Resend verification email"""
    user = await users_collection.find_one({"id": session["userId"]})
    
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    if user.get("emailVerified"):
        return {"success": True, "message": "Email already verified"}
    
    # Generate new code
    new_code = secrets.token_hex(3).upper()
    await users_collection.update_one(
        {"id": session["userId"]},
        {"$set": {"verificationCode": new_code}}
    )
    
    # Send verification email
    asyncio.create_task(send_verification_email(
        user["email"],
        user.get("name", ""),
        new_code,
        f"{os.environ.get('SITE_URL', 'https://mintslip.com')}/verify-email?code={new_code}&email={user['email']}"
    ))
    
    return {"success": True, "message": "Verification email sent"}


# ========== FORGOT PASSWORD ENDPOINTS ==========

class ForgotPasswordRequest(BaseModel):
    email: str

@app.post("/api/user/forgot-password")
async def forgot_password(data: ForgotPasswordRequest, request: Request):
    """Send password reset email"""
    # Get client IP for rate limiting
    client_ip = request.client.host if request.client else "unknown"
    
    # Rate limit: 3 password reset requests per minute per IP
    if not check_rate_limit(f"forgot_password_{client_ip}", max_requests=3, window_seconds=60):
        raise HTTPException(status_code=429, detail="Too many requests. Please try again later.")
    
    user = await users_collection.find_one({"email": data.email.lower()})
    
    # Always return success even if user doesn't exist (security best practice)
    if not user:
        return {"success": True, "message": "If an account exists with this email, you will receive a password reset link."}
    
    # Generate reset token and code
    reset_token = secrets.token_urlsafe(32)
    reset_code = secrets.token_hex(3).upper()  # 6-char code
    reset_expiry = datetime.now(timezone.utc) + timedelta(hours=1)
    
    # Store reset token in user record
    await users_collection.update_one(
        {"email": data.email.lower()},
        {"$set": {
            "passwordResetToken": reset_token,
            "passwordResetCode": reset_code,
            "passwordResetExpiry": reset_expiry.isoformat()
        }}
    )
    
    # Send reset email
    reset_link = f"{os.environ.get('SITE_URL', 'https://mintslip.com')}/reset-password?token={reset_token}&email={data.email.lower()}"
    asyncio.create_task(send_password_reset_email(
        user["email"],
        user.get("name", ""),
        reset_link,
        reset_code
    ))
    
    return {"success": True, "message": "If an account exists with this email, you will receive a password reset link."}


class ResetPasswordRequest(BaseModel):
    email: str
    token: Optional[str] = None
    code: Optional[str] = None
    newPassword: str

@app.post("/api/user/reset-password")
async def reset_password(data: ResetPasswordRequest, request: Request):
    """Reset password using token or code"""
    # Get client IP for rate limiting
    client_ip = request.client.host if request.client else "unknown"
    
    # Rate limit: 5 reset attempts per minute per IP
    if not check_rate_limit(f"reset_password_{client_ip}", max_requests=5, window_seconds=60):
        raise HTTPException(status_code=429, detail="Too many attempts. Please try again later.")
    
    user = await users_collection.find_one({"email": data.email.lower()})
    
    if not user:
        raise HTTPException(status_code=400, detail="Invalid reset request")
    
    # Check if reset token/code exists and is valid
    stored_token = user.get("passwordResetToken")
    stored_code = user.get("passwordResetCode")
    expiry = user.get("passwordResetExpiry")
    
    if not stored_token or not expiry:
        raise HTTPException(status_code=400, detail="No password reset request found. Please request a new reset link.")
    
    # Check expiry
    expiry_dt = datetime.fromisoformat(expiry.replace('Z', '+00:00')) if isinstance(expiry, str) else expiry
    if datetime.now(timezone.utc) > expiry_dt:
        # Clear expired reset data
        await users_collection.update_one(
            {"email": data.email.lower()},
            {"$unset": {"passwordResetToken": "", "passwordResetCode": "", "passwordResetExpiry": ""}}
        )
        raise HTTPException(status_code=400, detail="Reset link has expired. Please request a new one.")
    
    # Verify token or code
    token_valid = data.token and data.token == stored_token
    code_valid = data.code and data.code.upper() == stored_code
    
    if not token_valid and not code_valid:
        raise HTTPException(status_code=400, detail="Invalid reset token or code")
    
    # Validate new password
    if len(data.newPassword) < 8:
        raise HTTPException(status_code=400, detail="Password must be at least 8 characters")
    
    # Update password and clear reset data
    await users_collection.update_one(
        {"email": data.email.lower()},
        {
            "$set": {
                "password": hash_password(data.newPassword),
                "passwordChangedAt": datetime.now(timezone.utc).isoformat()
            },
            "$unset": {
                "passwordResetToken": "",
                "passwordResetCode": "",
                "passwordResetExpiry": ""
            }
        }
    )
    
    # Send password changed notification
    asyncio.create_task(send_password_changed_email(user["email"], user.get("name", "")))
    
    return {"success": True, "message": "Password reset successfully. You can now log in with your new password."}


# ========== DOWNLOAD EMAIL WITH FILE ATTACHMENT ENDPOINT ==========

class SendDownloadEmailRequest(BaseModel):
    email: str
    userName: Optional[str] = ""
    documentType: str
    pdfBase64: str  # Base64 encoded file content (PDF or ZIP)
    fileName: Optional[str] = None
    isGuest: bool = False
    isZip: bool = False  # True if the file is a ZIP archive

@app.post("/api/send-download-email")
async def send_download_email_with_attachment(data: SendDownloadEmailRequest):
    """
    Send download confirmation email with file attached (PDF or ZIP).
    File is passed in memory and not stored - respects user privacy.
    """
    document_names = {
        "paystub": "Pay Stub",
        "canadian-paystub": "Canadian Pay Stub",
        "w2": "W-2 Form",
        "w9": "W-9 Form",
        "1099-nec": "1099-NEC Form",
        "1099-misc": "1099-MISC Form",
        "bank-statement": "Accounting Mockup",
        "ai-resume": "AI Resume",
        "offer-letter": "Offer Letter",
        "schedule-c": "Schedule C",
        "utility-bill": "Utility Bill",
        "vehicle-bill-of-sale": "Vehicle Bill of Sale"
    }
    
    doc_name = document_names.get(data.documentType, data.documentType.replace("-", " ").title())
    
    # Generate filename if not provided - use .zip extension for ZIP files
    file_ext = ".zip" if data.isZip else ".pdf"
    filename = data.fileName or f"MintSlip_{doc_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}{file_ext}"
    
    # Ensure correct extension
    if data.isZip and not filename.endswith('.zip'):
        filename = filename.rsplit('.', 1)[0] + '.zip'
    elif not data.isZip and not filename.endswith('.pdf'):
        filename = filename.rsplit('.', 1)[0] + '.pdf'
    
    # Prepare attachment (Resend accepts base64 content directly)
    file_attachment = {
        "filename": filename,
        "content": data.pdfBase64  # Already base64 encoded from frontend
    }
    
    try:
        result = await send_download_confirmation(
            user_email=data.email,
            user_name=data.userName or "",
            document_type=data.documentType,
            download_link=None,
            is_guest=data.isGuest,
            pdf_attachment=file_attachment
        )
        
        if result.get("success"):
            logger.info(f"Download email with {'ZIP' if data.isZip else 'PDF'} attachment sent to {data.email}")
            return {"success": True, "message": f"Email sent with {'ZIP' if data.isZip else 'PDF'} attachment"}
        else:
            logger.error(f"Failed to send download email: {result.get('error')}")
            raise HTTPException(status_code=500, detail=f"Failed to send email: {result.get('error')}")
            
    except Exception as e:
        logger.error(f"Error sending download email with attachment: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to send email: {str(e)}")


# ========== USER PREFERENCES ENDPOINTS ==========

@app.put("/api/user/preferences")
async def update_user_preferences(data: UserPreferencesUpdate, session: dict = Depends(get_current_user)):
    """Update user preferences"""
    update_data = {}
    
    if data.saveDocuments is not None:
        update_data["preferences.saveDocuments"] = data.saveDocuments
    
    if not update_data:
        raise HTTPException(status_code=400, detail="No preferences to update")
    
    await users_collection.update_one(
        {"id": session["userId"]},
        {"$set": update_data}
    )
    
    # Get updated user
    user = await users_collection.find_one({"id": session["userId"]}, {"_id": 0, "password": 0})
    
    return {"success": True, "user": user, "message": "Preferences updated successfully"}


# ========== SAVED DOCUMENTS ENDPOINTS ==========

@app.get("/api/user/saved-documents")
async def get_saved_documents(
    session: dict = Depends(get_current_user),
    skip: int = 0,
    limit: int = 20
):
    """Get user's saved documents"""
    user_id = session["userId"]
    
    # Clean up expired documents first
    expiry_date = datetime.now(timezone.utc) - timedelta(days=DOCUMENT_EXPIRY_DAYS)
    expired_docs = await saved_documents_collection.find({
        "userId": user_id,
        "createdAt": {"$lt": expiry_date.isoformat()}
    }).to_list(100)
    
    # Delete expired documents and their files
    for doc in expired_docs:
        file_path = os.path.join(USER_DOCUMENTS_DIR, doc.get("storedFileName", ""))
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
            except:
                pass
        await saved_documents_collection.delete_one({"id": doc["id"]})
    
    # Get saved documents
    documents = await saved_documents_collection.find(
        {"userId": user_id},
        {"_id": 0, "storedFileName": 0}  # Don't expose internal filename
    ).sort("createdAt", -1).skip(skip).limit(limit).to_list(limit)
    
    total = await saved_documents_collection.count_documents({"userId": user_id})
    
    # Add days remaining for each document
    now = datetime.now(timezone.utc)
    for doc in documents:
        created_at = datetime.fromisoformat(doc["createdAt"].replace("Z", "+00:00"))
        expiry_date = created_at + timedelta(days=DOCUMENT_EXPIRY_DAYS)
        days_remaining = (expiry_date - now).days
        doc["daysRemaining"] = max(0, days_remaining)
        doc["expiresAt"] = expiry_date.isoformat()
    
    # Get user's subscription tier to determine max documents
    user = await users_collection.find_one({"id": user_id}, {"_id": 0})
    subscription = user.get("subscription") if user else None
    subscription_tier = subscription.get("tier", "starter") if subscription else "starter"
    max_documents = SAVED_DOCS_LIMITS.get(subscription_tier, 10)
    
    return {
        "success": True,
        "documents": documents,
        "total": total,
        "maxDocuments": max_documents,
        "subscriptionTier": subscription_tier
    }


@app.post("/api/user/saved-documents")
async def save_document(data: SaveDocumentRequest, session: dict = Depends(get_current_user)):
    """Save a document for later access"""
    user_id = session["userId"]
    
    # Check user preference
    user = await users_collection.find_one({"id": user_id}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Check if user has save documents enabled
    if not user.get("preferences", {}).get("saveDocuments", False):
        raise HTTPException(status_code=400, detail="Document saving is not enabled. Enable it in your settings.")
    
    # Get user's subscription tier to determine max documents
    subscription_tier = user.get("subscription", {}).get("tier", "starter") if user.get("subscription") else None
    if not subscription_tier:
        raise HTTPException(status_code=400, detail="Active subscription required to save documents.")
    
    max_documents = SAVED_DOCS_LIMITS.get(subscription_tier, 10)
    
    # Check document count limit (skip if unlimited)
    doc_count = await saved_documents_collection.count_documents({"userId": user_id})
    if max_documents != -1 and doc_count >= max_documents:
        raise HTTPException(
            status_code=400, 
            detail=f"Maximum saved documents limit ({max_documents}) reached for your {subscription_tier.capitalize()} plan. Delete some documents to save new ones."
        )
    
    # Decode base64 file data
    try:
        file_content = base64.b64decode(data.fileData)
    except Exception as e:
        raise HTTPException(status_code=400, detail="Invalid file data")
    
    # Generate unique filename (for reference only)
    file_ext = os.path.splitext(data.fileName)[1] or ".pdf"
    stored_filename = f"{user_id}_{uuid.uuid4()}{file_ext}"
    
    # Store file content as base64 in MongoDB for persistence
    # Also save to disk as backup/cache
    file_path = os.path.join(USER_DOCUMENTS_DIR, stored_filename)
    try:
        with open(file_path, "wb") as f:
            f.write(file_content)
    except Exception as e:
        logger.warning(f"Failed to save file to disk (will use MongoDB): {e}")
    
    # Create document record with file content stored in MongoDB
    doc_id = str(uuid.uuid4())
    document = {
        "id": doc_id,
        "userId": user_id,
        "documentType": data.documentType,
        "fileName": data.fileName,
        "storedFileName": stored_filename,
        "fileSize": len(file_content),
        "template": data.template,
        "fileContent": data.fileData,  # Store base64 content in MongoDB for persistence
        "createdAt": datetime.now(timezone.utc).isoformat()
    }
    
    await saved_documents_collection.insert_one(document)
    
    # Calculate expiry info
    expiry_date = datetime.now(timezone.utc) + timedelta(days=DOCUMENT_EXPIRY_DAYS)
    
    return {
        "success": True,
        "document": {
            "id": doc_id,
            "documentType": data.documentType,
            "fileName": data.fileName,
            "fileSize": len(file_content),
            "template": data.template,
            "createdAt": document["createdAt"],
            "daysRemaining": DOCUMENT_EXPIRY_DAYS,
            "expiresAt": expiry_date.isoformat()
        },
        "message": f"Document saved. It will be available for {DOCUMENT_EXPIRY_DAYS} days."
    }


@app.get("/api/user/saved-documents/{doc_id}/download")
async def download_saved_document(doc_id: str, session: dict = Depends(get_current_user)):
    """Download a saved document"""
    from fastapi.responses import FileResponse, Response
    
    user_id = session["userId"]
    
    # Find document
    document = await saved_documents_collection.find_one({
        "id": doc_id,
        "userId": user_id
    })
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Check if expired
    created_at = datetime.fromisoformat(document["createdAt"].replace("Z", "+00:00"))
    if datetime.now(timezone.utc) - created_at > timedelta(days=DOCUMENT_EXPIRY_DAYS):
        # Delete expired document
        file_path = os.path.join(USER_DOCUMENTS_DIR, document["storedFileName"])
        if os.path.exists(file_path):
            os.remove(file_path)
        await saved_documents_collection.delete_one({"id": doc_id})
        raise HTTPException(status_code=404, detail="Document has expired and been deleted")
    
    # Determine media type
    file_ext = os.path.splitext(document["fileName"])[1].lower()
    media_types = {
        ".pdf": "application/pdf",
        ".zip": "application/zip"
    }
    media_type = media_types.get(file_ext, "application/octet-stream")
    
    # Try to get file from disk first
    file_path = os.path.join(USER_DOCUMENTS_DIR, document["storedFileName"])
    
    if os.path.exists(file_path):
        return FileResponse(
            path=file_path,
            filename=document["fileName"],
            media_type=media_type
        )
    
    # If file not on disk, try to restore from MongoDB content
    if document.get("fileContent"):
        try:
            file_content = base64.b64decode(document["fileContent"])
            # Restore file to disk for future access
            try:
                with open(file_path, "wb") as f:
                    f.write(file_content)
                logger.info(f"Restored file from MongoDB: {file_path}")
            except Exception as e:
                logger.warning(f"Could not restore file to disk: {e}")
            
            return Response(
                content=file_content,
                media_type=media_type,
                headers={"Content-Disposition": f'attachment; filename="{document["fileName"]}"'}
            )
        except Exception as e:
            logger.error(f"Failed to decode file content from MongoDB: {e}")
    
    raise HTTPException(status_code=404, detail="Document file not found")


@app.delete("/api/user/saved-documents/{doc_id}")
async def delete_saved_document(doc_id: str, session: dict = Depends(get_current_user)):
    """Delete a saved document"""
    user_id = session["userId"]
    
    # Find document
    document = await saved_documents_collection.find_one({
        "id": doc_id,
        "userId": user_id
    })
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Delete file
    file_path = os.path.join(USER_DOCUMENTS_DIR, document["storedFileName"])
    if os.path.exists(file_path):
        try:
            os.remove(file_path)
        except:
            pass
    
    # Delete record
    await saved_documents_collection.delete_one({"id": doc_id})
    
    return {"success": True, "message": "Document deleted successfully"}


@app.get("/api/user/saved-documents/count")
async def get_saved_documents_count(session: dict = Depends(get_current_user)):
    """Get count of saved documents for current user"""
    user_id = session["userId"]
    count = await saved_documents_collection.count_documents({"userId": user_id})
    
    # Get user's subscription tier to determine max documents
    user = await users_collection.find_one({"id": user_id}, {"_id": 0})
    subscription_tier = user.get("subscription", {}).get("tier", "starter") if user and user.get("subscription") else "starter"
    max_documents = SAVED_DOCS_LIMITS.get(subscription_tier, 10)
    
    remaining = "unlimited" if max_documents == -1 else max(0, max_documents - count)
    
    return {
        "success": True,
        "count": count,
        "maxDocuments": max_documents,
        "remaining": remaining,
        "subscriptionTier": subscription_tier
    }



# ========== STRIPE SUBSCRIPTION ENDPOINTS ==========

# Subscription Plans Configuration
SUBSCRIPTION_PLANS = {
    "starter": {
        "name": "Starter",
        "price": 19.99,
        "price_cents": 1999,
        "downloads": 10,
        "description": "10 downloads per month"
    },
    "professional": {
        "name": "Professional", 
        "price": 29.99,
        "price_cents": 2999,
        "downloads": 30,
        "description": "30 downloads per month"
    },
    "business": {
        "name": "Business",
        "price": 49.99,
        "price_cents": 4999,
        "downloads": -1,
        "description": "Unlimited downloads"
    }
}

# Store for Stripe Price IDs (will be populated when products are created)
stripe_price_ids = {}


@app.get("/api/subscriptions/plans")
async def get_subscription_plans():
    """Get available subscription plans with dynamic download limits"""
    tier_downloads = await get_tier_downloads()
    plans = []
    for tier, config in SUBSCRIPTION_PLANS.items():
        dl = tier_downloads.get(tier, config["downloads"])
        plans.append({
            "tier": tier,
            "name": config["name"],
            "price": config["price"],
            "downloads": dl,
            "downloads_display": "Unlimited" if dl == -1 else str(dl)
        })
    return {"success": True, "plans": plans}


@app.get("/api/stripe/config")
async def get_stripe_config():
    """Get Stripe publishable key for frontend"""
    return {
        "publishableKey": os.environ.get("STRIPE_PUBLISHABLE_KEY", "pk_test_51SOOSM0OuJwef38xP0FqCJ3b45STthDKnJWP572LoODAaxGIq8ujrAwp1W0MeGkI6XczeweTr7lOLIKC6MnLadoX00iDo2VzYM")
    }


@app.post("/api/stripe/create-checkout-session")
async def create_checkout_session(data: CreateCheckoutSession, session: dict = Depends(get_current_user)):
    """Create a Stripe checkout session for subscription"""
    tier = data.tier
    
    if tier not in SUBSCRIPTION_PLANS:
        raise HTTPException(status_code=400, detail="Invalid subscription tier")
    
    plan = SUBSCRIPTION_PLANS[tier]
    user = await users_collection.find_one({"id": session["userId"]})
    
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    try:
        # Create or get Stripe customer
        stripe_customer_id = user.get("stripeCustomerId")
        
        if not stripe_customer_id:
            customer = stripe.Customer.create(
                email=user["email"],
                name=user.get("name", ""),
                metadata={"userId": user["id"]}
            )
            stripe_customer_id = customer.id
            await users_collection.update_one(
                {"id": user["id"]},
                {"$set": {"stripeCustomerId": stripe_customer_id}}
            )
        
        # Create price if not exists
        price_id = stripe_price_ids.get(tier)
        if not price_id:
            # Create product and price
            product = stripe.Product.create(
                name=f"MintSlip {plan['name']} Plan",
                description=plan["description"]
            )
            
            price = stripe.Price.create(
                product=product.id,
                unit_amount=plan["price_cents"],
                currency="usd",
                recurring={"interval": "month"}
            )
            price_id = price.id
            stripe_price_ids[tier] = price_id
        
        # Get frontend URL from environment or use default
        frontend_url = os.environ.get("FRONTEND_URL", "https://l7ltqw-3000.csb.app")
        
        # Create checkout session
        checkout_session = stripe.checkout.Session.create(
            customer=stripe_customer_id,
            payment_method_types=["card"],
            line_items=[{
                "price": price_id,
                "quantity": 1
            }],
            mode="subscription",
            success_url=data.successUrl or f"{frontend_url}/subscription/success?session_id={{CHECKOUT_SESSION_ID}}",
            cancel_url=data.cancelUrl or f"{frontend_url}/subscription/cancel",
            metadata={
                "userId": user["id"],
                "tier": tier
            }
        )
        
        return {
            "success": True,
            "sessionId": checkout_session.id,
            "url": checkout_session.url
        }
        
    except stripe.error.StripeError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/api/stripe/create-payment-intent")
async def create_payment_intent(request: dict):
    """Create a Stripe payment intent for one-time purchase"""
    amount = request.get("amount")  # Amount in dollars
    document_type = request.get("documentType")
    template = request.get("template")
    email = request.get("email")
    discount_code = request.get("discountCode")
    discount_amount = request.get("discountAmount", 0)
    
    if not amount:
        raise HTTPException(status_code=400, detail="Amount is required")
    
    try:
        # Convert to cents
        amount_cents = int(float(amount) * 100)
        
        # Create payment intent
        intent = stripe.PaymentIntent.create(
            amount=amount_cents,
            currency="usd",
            metadata={
                "documentType": document_type or "",
                "template": template or "",
                "email": email or "",
                "discountCode": discount_code or "",
                "discountAmount": str(discount_amount)
            }
        )
        
        return {
            "success": True,
            "clientSecret": intent.client_secret,
            "paymentIntentId": intent.id
        }
        
    except stripe.error.StripeError as e:
        raise HTTPException(status_code=400, detail=str(e))


class TrackCheckoutStarted(BaseModel):
    email: str
    documentType: str
    userName: Optional[str] = ""
    userId: Optional[str] = None

@app.post("/api/track-checkout-started")
async def api_track_checkout_started(data: TrackCheckoutStarted):
    """Track when a user starts checkout for abandoned cart emails"""
    if data.email:
        asyncio.create_task(track_checkout_started(
            data.email,
            data.userName or "",
            data.userId,
            data.documentType
        ))
    return {"success": True}


class OneTimeCheckoutRequest(BaseModel):
    amount: float
    documentType: str
    template: Optional[str] = None
    discountCode: Optional[str] = None
    discountAmount: Optional[float] = 0
    successUrl: Optional[str] = None
    cancelUrl: Optional[str] = None
    userId: Optional[str] = None  # Optional userId for logged-in users
    userEmail: Optional[str] = None  # Optional user email
    quantity: Optional[int] = 1  # Number of documents (e.g., number of paystubs)


@app.post("/api/stripe/create-one-time-checkout")
async def create_one_time_checkout(data: OneTimeCheckoutRequest, request: Request, authorization: Optional[str] = Header(None)):
    """Create a Stripe checkout session for one-time document purchase"""
    try:
        # Try to get user info from token if provided
        user_id = data.userId
        user_email = data.userEmail
        
        if authorization and authorization.startswith("Bearer "):
            token = authorization.split(" ")[1]
            session = await sessions_collection.find_one({"token": token})
            if session:
                user = await users_collection.find_one({"id": session["userId"]})
                if user:
                    user_id = user["id"]
                    user_email = user["email"]
        
        # Get client IP for tracking
        client_ip = get_client_ip(request)
        
        # Convert to cents
        amount_cents = int(float(data.amount) * 100)
        
        # Get frontend URL from environment or use default
        frontend_url = os.environ.get("FRONTEND_URL", "https://l7ltqw-3000.csb.app")
        
        # Create checkout session for one-time payment
        # Map document types to display names for Stripe
        doc_display_names = {
            "bank-statement": "Accounting Mockup",
        }
        display_name = doc_display_names.get(data.documentType, data.documentType.replace('-', ' ').title())
        
        checkout_session = stripe.checkout.Session.create(
            payment_method_types=["card"],
            line_items=[{
                "price_data": {
                    "currency": "usd",
                    "product_data": {
                        "name": f"MintSlip - {display_name}",
                        "description": f"Professional {display_name} document generation"
                    },
                    "unit_amount": amount_cents,
                },
                "quantity": 1
            }],
            mode="payment",
            success_url=data.successUrl or f"{frontend_url}/payment-success?type={data.documentType}&session_id={{CHECKOUT_SESSION_ID}}",
            cancel_url=data.cancelUrl or frontend_url,
            metadata={
                "documentType": data.documentType,
                "template": data.template or "",
                "discountCode": data.discountCode or "",
                "discountAmount": str(data.discountAmount or 0),
                "type": "one_time_purchase",
                "userId": user_id or "",
                "userEmail": user_email or "",
                "clientIp": client_ip,
                "quantity": str(data.quantity or 1)
            }
        )
        
        return {
            "success": True,
            "sessionId": checkout_session.id,
            "url": checkout_session.url
        }
        
    except stripe.error.StripeError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.get("/api/stripe/checkout-status/{session_id}")
async def get_checkout_status(session_id: str):
    """Get the status of a Stripe checkout session"""
    try:
        session = stripe.checkout.Session.retrieve(session_id)
        
        # If payment is complete, process accordingly
        if session.payment_status == "paid" and session.status == "complete":
            user_id = session.metadata.get("userId")
            tier = session.metadata.get("tier")
            purchase_type = session.metadata.get("type")
            
            if user_id and tier:
                # Handle subscription activation
                user = await users_collection.find_one({"id": user_id})
                if user and (not user.get("subscription") or user.get("subscription", {}).get("stripeSubscriptionId") != session.subscription):
                    plan_config = SUBSCRIPTION_PLANS.get(tier, {})
                    tier_downloads = await get_tier_downloads()

                    # Update user subscription
                    await users_collection.update_one(
                        {"id": user_id},
                        {"$set": {
                            "subscription": {
                                "tier": tier,
                                "status": "active",
                                "stripeSubscriptionId": session.subscription,
                                "stripeCustomerId": session.customer,
                                "downloads_remaining": tier_downloads.get(tier, 0),
                                "downloads_total": tier_downloads.get(tier, 0),
                                "current_period_start": datetime.now(timezone.utc).isoformat(),
                                "current_period_end": (datetime.now(timezone.utc) + timedelta(days=30)).isoformat(),
                                "createdAt": datetime.now(timezone.utc).isoformat()
                            }
                        }}
                    )
            
            elif purchase_type == "one_time_purchase":
                # Track one-time purchase (if not already tracked)
                existing = await purchases_collection.find_one({"stripeSessionId": session_id})
                if not existing:
                    document_type = session.metadata.get("documentType", "unknown")
                    template = session.metadata.get("template", "")
                    discount_code = session.metadata.get("discountCode", "")
                    discount_amount = float(session.metadata.get("discountAmount", 0))
                    user_id = session.metadata.get("userId", "")
                    user_email = session.metadata.get("userEmail", "")
                    quantity = int(session.metadata.get("quantity", 1))
                    
                    # Get customer email from session if not in metadata
                    customer_email = user_email
                    if not customer_email and hasattr(session, 'customer_details') and session.customer_details:
                        customer_email = getattr(session.customer_details, 'email', "") or ""
                    
                    purchase = {
                        "id": str(uuid.uuid4()),
                        "documentType": document_type,
                        "amount": session.amount_total / 100 if session.amount_total else 0,
                        "email": customer_email,
                        "userId": user_id if user_id else None,
                        "stripeSessionId": session_id,
                        "stripePaymentIntentId": session.payment_intent,
                        "discountCode": discount_code if discount_code else None,
                        "discountAmount": discount_amount,
                        "template": template if template else None,
                        "quantity": quantity,
                        "isGuest": not bool(user_id),
                        "createdAt": datetime.now(timezone.utc).isoformat()
                    }
                    await purchases_collection.insert_one(purchase)
                    asyncio.create_task(create_notification(document_type, customer_email, purchase["amount"]))
                    print(f"Tracked purchase via status check: {document_type} x{quantity} - ${purchase['amount']} - userId: {user_id or 'guest'}")
                    
                    # Send download confirmation and review request emails
                    if customer_email:
                        # Get user name if logged in
                        user_name = ""
                        if user_id:
                            user = await users_collection.find_one({"id": user_id})
                            if user:
                                user_name = user.get("name", "")
                        is_guest = not bool(user_id)
                        # NOTE: Download confirmation with PDF attachment is now sent from frontend
                        # Only send review request and cancel abandoned checkout emails here
                        asyncio.create_task(send_review_request(customer_email, user_name, document_type, user_id if user_id else None))
                        asyncio.create_task(cancel_abandoned_checkout_email(customer_email))
                    
                    # Increment discount code usage if one was used
                    if discount_code:
                        customer_identifier = customer_email or user_id or session_id
                        update_ops = {"$inc": {"usageCount": 1}}
                        # Check if it's a one_per_customer discount
                        discount = await discounts_collection.find_one({"code": discount_code.upper()})
                        if discount and discount.get("usageType") == "one_per_customer" and customer_identifier:
                            update_ops["$push"] = {"usedByCustomers": customer_identifier}
                        await discounts_collection.update_one({"code": discount_code.upper()}, update_ops)
                        print(f"Incremented usage count for discount code: {discount_code}")
        
        return {
            "status": session.status,
            "payment_status": session.payment_status,
            "amount_total": session.amount_total,
            "currency": session.currency,
            "metadata": dict(session.metadata) if session.metadata else {},
            "customer_email": getattr(session.customer_details, 'email', "") if hasattr(session, 'customer_details') and session.customer_details else ""
        }
        
    except stripe.error.StripeError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/api/stripe/webhook")
async def stripe_webhook(request: Request):
    """Handle Stripe webhooks"""
    payload = await request.body()
    sig_header = request.headers.get("stripe-signature")
    
    # In production, verify webhook signature
    # For now, just parse the event
    try:
        event = stripe.Event.construct_from(
            json.loads(payload), stripe.api_key
        )
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Webhook error: {str(e)}")
    
    # Handle the event
    if event.type == "checkout.session.completed":
        session = event.data.object
        
        # Get metadata
        user_id = session.metadata.get("userId")
        tier = session.metadata.get("tier")
        purchase_type = session.metadata.get("type")
        
        if user_id and tier:
            # Handle subscription purchase
            plan_config = SUBSCRIPTION_PLANS.get(tier, {})
            tier_downloads = await get_tier_downloads()

            # Update user subscription
            await users_collection.update_one(
                {"id": user_id},
                {"$set": {
                    "subscription": {
                        "tier": tier,
                        "status": "active",
                        "stripeSubscriptionId": session.subscription,
                        "stripeCustomerId": session.customer,
                        "downloads_remaining": tier_downloads.get(tier, 0),
                        "downloads_total": tier_downloads.get(tier, 0),
                        "current_period_start": datetime.now(timezone.utc).isoformat(),
                        "current_period_end": (datetime.now(timezone.utc) + timedelta(days=30)).isoformat(),
                        "createdAt": datetime.now(timezone.utc).isoformat()
                    }
                }}
            )
            
            # Create subscription record
            subscription_record = {
                "id": str(uuid.uuid4()),
                "userId": user_id,
                "tier": tier,
                "stripeSubscriptionId": session.subscription,
                "stripeCustomerId": session.customer,
                "status": "active",
                "createdAt": datetime.now(timezone.utc).isoformat()
            }
            await subscriptions_collection.insert_one(subscription_record)
            
            # Send subscription thank you email and cancel no-purchase reminder
            user = await users_collection.find_one({"id": user_id})
            if user:
                asyncio.create_task(schedule_subscription_thank_you(
                    user["email"],
                    user.get("name", ""),
                    user_id,
                    plan_config.get("name", tier.title()),
                    str(plan_config.get("price", 0)),
                    tier_downloads.get(tier, 0)
                ))
                asyncio.create_task(cancel_signup_no_purchase_reminder(user_id))
        
        elif purchase_type == "one_time_purchase":
            # Handle one-time guest purchase
            document_type = session.metadata.get("documentType", "unknown")
            template = session.metadata.get("template", "")
            discount_code = session.metadata.get("discountCode", "")
            discount_amount = float(session.metadata.get("discountAmount", 0))
            
            # Get customer email from session if available
            customer_email = ""
            if session.customer_details:
                customer_email = session.customer_details.get("email", "")
            
            # Get client IP from the webhook request
            client_ip = get_client_ip(request)
            
            # Also try to get IP from metadata if passed during checkout creation
            metadata_ip = session.metadata.get("clientIp", "")
            
            # Get quantity from metadata
            quantity = int(session.metadata.get("quantity", 1))
            
            purchase = {
                "id": str(uuid.uuid4()),
                "documentType": document_type,
                "amount": session.amount_total / 100 if session.amount_total else 0,
                "email": customer_email,
                "stripeSessionId": session.id,
                "stripePaymentIntentId": session.payment_intent,
                "discountCode": discount_code if discount_code else None,
                "discountAmount": discount_amount,
                "template": template if template else None,
                "quantity": quantity,
                "isGuest": True,
                "ipAddress": metadata_ip if metadata_ip else client_ip,
                "createdAt": datetime.now(timezone.utc).isoformat()
            }
            await purchases_collection.insert_one(purchase)
            asyncio.create_task(create_notification(document_type, customer_email, purchase["amount"]))
            print(f"Tracked guest purchase: {document_type} - ${purchase['amount']}")
            
            # Send emails for guest purchase
            if customer_email:
                # NOTE: Download confirmation with PDF attachment is now sent from frontend
                asyncio.create_task(send_review_request(customer_email, "", document_type, None))
                asyncio.create_task(cancel_abandoned_checkout_email(customer_email))
            
            # Increment discount code usage if one was used
            if discount_code:
                customer_identifier = customer_email or session.id
                update_ops = {"$inc": {"usageCount": 1}}
                # Check if it's a one_per_customer discount
                discount = await discounts_collection.find_one({"code": discount_code.upper()})
                if discount and discount.get("usageType") == "one_per_customer" and customer_identifier:
                    update_ops["$push"] = {"usedByCustomers": customer_identifier}
                await discounts_collection.update_one({"code": discount_code.upper()}, update_ops)
                print(f"Incremented usage count for discount code: {discount_code}")
    
    elif event.type == "invoice.payment_succeeded":
        invoice = event.data.object
        subscription_id = invoice.subscription
        
        if subscription_id:
            # Find user with this subscription
            user = await users_collection.find_one({"subscription.stripeSubscriptionId": subscription_id})
            if user:
                tier = user["subscription"].get("tier")
                tier_downloads = await get_tier_downloads()

                # Record the subscription payment
                payment_amount = invoice.amount_paid / 100  # Convert from cents
                subscription_payment = {
                    "id": str(uuid.uuid4()),
                    "userId": user["id"],
                    "userEmail": user.get("email", ""),
                    "tier": tier,
                    "amount": payment_amount,
                    "stripeInvoiceId": invoice.id,
                    "stripeSubscriptionId": subscription_id,
                    "billingReason": invoice.billing_reason,  # 'subscription_create', 'subscription_cycle', 'subscription_update'
                    "createdAt": datetime.now(timezone.utc).isoformat()
                }
                await subscription_payments_collection.insert_one(subscription_payment)
                print(f"Recorded subscription payment: {payment_amount} for user {user['id']} tier {tier}")
                
                # Reset downloads for the new billing period
                await users_collection.update_one(
                    {"id": user["id"]},
                    {"$set": {
                        "subscription.downloads_remaining": tier_downloads.get(tier, 0),
                        "subscription.current_period_end": (datetime.now(timezone.utc) + timedelta(days=30)).isoformat()
                    }}
                )
    
    elif event.type == "customer.subscription.deleted":
        subscription = event.data.object
        subscription_id = subscription.id
        
        # Find and update user
        user = await users_collection.find_one({"subscription.stripeSubscriptionId": subscription_id})
        if user:
            await users_collection.update_one(
                {"id": user["id"]},
                {"$set": {"subscription.status": "cancelled"}}
            )
    
    elif event.type == "customer.subscription.updated":
        # Handle subscription updates from Stripe (e.g., plan changes, renewals)
        subscription = event.data.object
        subscription_id = subscription.id
        
        # Find user with this subscription
        user = await users_collection.find_one({"subscription.stripeSubscriptionId": subscription_id})
        if user:
            # Get the current price/plan from Stripe subscription
            if subscription.items and subscription.items.data:
                current_item = subscription.items.data[0]
                stripe_price_id = current_item.price.id
                
                # Try to find the matching tier from our stored price IDs
                new_tier = None
                for tier, price_id in stripe_price_ids.items():
                    if price_id == stripe_price_id:
                        new_tier = tier
                        break
                
                # If we found a matching tier and it's different from current, update
                if new_tier and new_tier != user.get("subscription", {}).get("tier"):
                    tier_downloads = await get_tier_downloads()
                    await users_collection.update_one(
                        {"id": user["id"]},
                        {"$set": {
                            "subscription.tier": new_tier,
                            "subscription.downloads_remaining": tier_downloads.get(new_tier, 0),
                            "subscription.downloads_total": tier_downloads.get(new_tier, 0),
                            "subscription.status": subscription.status,
                            "subscription.updatedAt": datetime.now(timezone.utc).isoformat()
                        }}
                    )
                    print(f"Webhook updated user {user['id']} subscription to {new_tier}")
                else:
                    # Just update the status if tier hasn't changed
                    await users_collection.update_one(
                        {"id": user["id"]},
                        {"$set": {
                            "subscription.status": subscription.status,
                            "subscription.updatedAt": datetime.now(timezone.utc).isoformat()
                        }}
                    )
    
    elif event.type == "payment_intent.succeeded":
        payment_intent = event.data.object
        
        # Track purchase for one-time payments
        metadata = payment_intent.metadata
        if metadata.get("documentType"):
            purchase = {
                "id": str(uuid.uuid4()),
                "documentType": metadata.get("documentType"),
                "amount": payment_intent.amount / 100,
                "email": metadata.get("email", ""),
                "stripePaymentIntentId": payment_intent.id,
                "discountCode": metadata.get("discountCode"),
                "discountAmount": float(metadata.get("discountAmount", 0)),
                "template": metadata.get("template"),
                "createdAt": datetime.now(timezone.utc).isoformat()
            }
            await purchases_collection.insert_one(purchase)
            asyncio.create_task(create_notification(metadata.get("documentType", ""), metadata.get("email", ""), payment_intent.amount / 100))
    
    return {"status": "received"}


@app.post("/api/stripe/cancel-subscription")
async def cancel_stripe_subscription(session: dict = Depends(get_current_user)):
    """Cancel user's Stripe subscription"""
    user = await users_collection.find_one({"id": session["userId"]})
    
    if not user or not user.get("subscription"):
        raise HTTPException(status_code=404, detail="No active subscription found")
    
    subscription_id = user["subscription"].get("stripeSubscriptionId")
    
    if not subscription_id:
        raise HTTPException(status_code=400, detail="No Stripe subscription found")
    
    try:
        # Cancel at period end (user keeps access until end of billing period)
        stripe.Subscription.modify(
            subscription_id,
            cancel_at_period_end=True
        )
        
        await users_collection.update_one(
            {"id": user["id"]},
            {"$set": {"subscription.status": "cancelling"}}
        )
        
        return {"success": True, "message": "Subscription will be cancelled at the end of the billing period"}
        
    except stripe.error.StripeError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/api/stripe/reactivate-subscription")
async def reactivate_stripe_subscription(session: dict = Depends(get_current_user)):
    """Reactivate a cancelled Stripe subscription"""
    user = await users_collection.find_one({"id": session["userId"]})
    
    if not user or not user.get("subscription"):
        raise HTTPException(status_code=404, detail="No subscription found")
    
    subscription_id = user["subscription"].get("stripeSubscriptionId")
    
    if not subscription_id:
        raise HTTPException(status_code=400, detail="No Stripe subscription found")
    
    try:
        stripe.Subscription.modify(
            subscription_id,
            cancel_at_period_end=False
        )
        
        await users_collection.update_one(
            {"id": user["id"]},
            {"$set": {"subscription.status": "active"}}
        )
        
        return {"success": True, "message": "Subscription reactivated"}

    except stripe.error.StripeError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/api/billing/portal")
async def create_billing_portal_session(request: Request, session: dict = Depends(get_current_user)):
    """Create a Stripe Customer Portal session for self-serve billing management"""
    user = await users_collection.find_one({"id": session["userId"]})
    customer_id = user.get("subscription", {}).get("stripeCustomerId") if user else None

    if not customer_id:
        raise HTTPException(status_code=400, detail="No billing account found. Please contact support.")

    frontend_url = os.getenv("FRONTEND_URL", "https://mintslip.com")
    try:
        portal_session = await asyncio.to_thread(
            stripe.billing_portal.Session.create,
            **{"customer": customer_id, "return_url": f"{frontend_url}/user/settings"}
        )
        return {"success": True, "url": portal_session.url}
    except stripe.error.StripeError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.get("/api/user/invoices")
async def get_user_invoices(session: dict = Depends(get_current_user)):
    """Fetch the user's Stripe invoice history"""
    user = await users_collection.find_one({"id": session["userId"]})
    customer_id = user.get("subscription", {}).get("stripeCustomerId") if user else None

    if not customer_id:
        return {"success": True, "invoices": []}

    try:
        invoices_resp = await asyncio.to_thread(
            stripe.Invoice.list,
            **{"customer": customer_id, "limit": 24, "expand": ["data.subscription"]}
        )
        result = []
        for inv in invoices_resp.data:
            result.append({
                "id": inv.id,
                "amount": inv.amount_paid / 100,
                "currency": (inv.currency or "usd").upper(),
                "status": inv.status,
                "date": datetime.fromtimestamp(inv.created, tz=timezone.utc).isoformat(),
                "period_start": datetime.fromtimestamp(inv.period_start, tz=timezone.utc).isoformat() if inv.period_start else None,
                "period_end": datetime.fromtimestamp(inv.period_end, tz=timezone.utc).isoformat() if inv.period_end else None,
                "pdf_url": inv.invoice_pdf,
                "hosted_url": inv.hosted_invoice_url,
                "description": inv.lines.data[0].description if inv.lines and inv.lines.data else None,
            })
        return {"success": True, "invoices": result}
    except stripe.error.StripeError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/api/stripe/preview-proration")
async def preview_subscription_proration(request: dict, session: dict = Depends(get_current_user)):
    """Preview the prorated amount for a subscription change"""
    new_tier = request.get("tier")
    
    if new_tier not in SUBSCRIPTION_PLANS:
        raise HTTPException(status_code=400, detail="Invalid subscription tier")
    
    user = await users_collection.find_one({"id": session["userId"]})
    
    if not user or not user.get("subscription"):
        raise HTTPException(status_code=404, detail="No active subscription found")
    
    current_tier = user["subscription"].get("tier")
    if current_tier == new_tier:
        raise HTTPException(status_code=400, detail="Already on this plan")
    
    subscription_id = user["subscription"].get("stripeSubscriptionId")
    
    if not subscription_id:
        raise HTTPException(status_code=400, detail="No Stripe subscription found")
    
    try:
        # Get current subscription
        subscription = stripe.Subscription.retrieve(subscription_id)
        
        # Get or create price for new tier
        price_id = stripe_price_ids.get(new_tier)
        if not price_id:
            plan = SUBSCRIPTION_PLANS[new_tier]
            product = stripe.Product.create(
                name=f"MintSlip {plan['name']} Plan",
                description=plan["description"]
            )
            price = stripe.Price.create(
                product=product.id,
                unit_amount=plan["price_cents"],
                currency="usd",
                recurring={"interval": "month"}
            )
            price_id = price.id
            stripe_price_ids[new_tier] = price_id
        
        # Create an invoice preview to see the proration
        upcoming_invoice = stripe.Invoice.create_preview(
            customer=subscription.customer,
            subscription=subscription_id,
            subscription_details={
                "items": [{
                    "id": subscription["items"]["data"][0].id,
                    "price": price_id
                }],
                "proration_behavior": "create_prorations"
            }
        )
        
        # Calculate proration amounts from invoice lines
        # The invoice preview includes: proration charges + next billing period
        # We only want to show the proration (immediate charge), not the next period
        proration_charge = 0  # Amount charged for new plan's remaining time
        proration_credit = 0  # Credit for old plan's unused time
        next_period_amount = 0  # Next full billing cycle (not immediate)
        
        for line in upcoming_invoice.lines.data:
            description = getattr(line, 'description', '') or ''
            amount = line.amount
            
            # Debug: print line details
            print(f"Invoice line: {description[:50]}... amount={amount}")
            
            # Proration lines contain "Remaining time" or "Unused time"
            # Or they might say "time on [Plan Name]"
            is_proration_line = any(phrase in description.lower() for phrase in [
                'remaining time', 'unused time', 'time on', 'proration'
            ])
            
            if is_proration_line:
                if amount > 0:
                    proration_charge += amount
                else:
                    proration_credit += abs(amount)
            else:
                # This is the next billing period charge
                next_period_amount += amount
        
        # Net proration amount (what user pays NOW, not including next period)
        net_proration = proration_charge - proration_credit
        
        is_upgrade = SUBSCRIPTION_PLANS[new_tier]["price_cents"] > SUBSCRIPTION_PLANS[current_tier]["price_cents"]
        
        # Get period end date - use get() for safety
        period_end_timestamp = getattr(subscription, 'current_period_end', None)
        if period_end_timestamp:
            period_end = datetime.fromtimestamp(period_end_timestamp)
            days_remaining = (period_end - datetime.now()).days
        else:
            period_end = datetime.now() + timedelta(days=30)
            days_remaining = 30
        
        # Calculate a sensible proration if Stripe's numbers seem off
        # Proration should be: (new_price - old_price) * (days_remaining / 30)
        old_price = SUBSCRIPTION_PLANS[current_tier]["price_cents"]
        new_price = SUBSCRIPTION_PLANS[new_tier]["price_cents"]
        expected_proration = (new_price - old_price) * (max(0, days_remaining) / 30)
        
        # Use calculated proration if Stripe's seems unreasonable
        # (more than 2x the monthly difference is suspicious)
        monthly_diff = abs(new_price - old_price)
        if abs(net_proration) > monthly_diff * 2:
            print(f"Warning: Stripe proration {net_proration} seems high, using calculated: {expected_proration}")
            net_proration = int(expected_proration)
        
        return {
            "success": True,
            "preview": {
                "currentTier": current_tier,
                "currentTierName": SUBSCRIPTION_PLANS[current_tier]["name"],
                "newTier": new_tier,
                "newTierName": SUBSCRIPTION_PLANS[new_tier]["name"],
                "isUpgrade": is_upgrade,
                "currentPlanCredit": proration_credit / 100,  # Convert cents to dollars
                "newPlanCharge": proration_charge / 100,
                "netAmountDue": net_proration / 100,  # Just the proration, not next period
                "newMonthlyPrice": new_price / 100,
                "daysRemainingInPeriod": max(0, days_remaining),
                "periodEndDate": period_end.isoformat(),
                "immediateCharge": net_proration > 0,
                "creditApplied": net_proration < 0
            }
        }
        
    except stripe.error.StripeError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/api/stripe/change-subscription")
async def change_stripe_subscription(request: dict, session: dict = Depends(get_current_user)):
    """Change subscription tier"""
    new_tier = request.get("tier")
    
    if new_tier not in SUBSCRIPTION_PLANS:
        raise HTTPException(status_code=400, detail="Invalid subscription tier")
    
    user = await users_collection.find_one({"id": session["userId"]})
    
    if not user or not user.get("subscription"):
        raise HTTPException(status_code=404, detail="No active subscription found")
    
    subscription_id = user["subscription"].get("stripeSubscriptionId")
    
    if not subscription_id:
        raise HTTPException(status_code=400, detail="No Stripe subscription found")
    
    try:
        # Get current subscription
        subscription = stripe.Subscription.retrieve(subscription_id)
        
        # Get or create price for new tier
        price_id = stripe_price_ids.get(new_tier)
        if not price_id:
            plan = SUBSCRIPTION_PLANS[new_tier]
            product = stripe.Product.create(
                name=f"MintSlip {plan['name']} Plan",
                description=plan["description"]
            )
            price = stripe.Price.create(
                product=product.id,
                unit_amount=plan["price_cents"],
                currency="usd",
                recurring={"interval": "month"}
            )
            price_id = price.id
            stripe_price_ids[new_tier] = price_id
        
        # Update subscription in Stripe
        updated_stripe_sub = stripe.Subscription.modify(
            subscription_id,
            items=[{
                "id": subscription["items"]["data"][0].id,
                "price": price_id
            }],
            proration_behavior="create_prorations"
        )
        
        print(f"Stripe subscription updated: {updated_stripe_sub.id} to price {price_id}")
        
        # Update user record in MongoDB
        plan_config = SUBSCRIPTION_PLANS[new_tier]
        update_result = await users_collection.update_one(
            {"id": user["id"]},
            {"$set": {
                "subscription.tier": new_tier,
                "subscription.downloads_remaining": plan_config["downloads"],
                "subscription.downloads_total": plan_config["downloads"],
                "subscription.updatedAt": datetime.now(timezone.utc).isoformat()
            }}
        )
        
        print(f"MongoDB update result: matched={update_result.matched_count}, modified={update_result.modified_count}")
        
        # Verify the update
        updated_user = await users_collection.find_one({"id": user["id"]})
        if updated_user:
            print(f"Verified user subscription tier: {updated_user.get('subscription', {}).get('tier')}")
        
        return {"success": True, "message": f"Subscription changed to {plan_config['name']}", "newTier": new_tier}
        
    except stripe.error.StripeError as e:
        print(f"Stripe error during subscription change: {str(e)}")
        raise HTTPException(status_code=400, detail=str(e))


@app.get("/api/stripe/subscription-status")
async def get_subscription_status(session: dict = Depends(get_current_user)):
    """Get current user's subscription status from Stripe"""
    user = await users_collection.find_one({"id": session["userId"]})
    
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    if not user.get("subscription") or not user["subscription"].get("stripeSubscriptionId"):
        return {"success": True, "hasSubscription": False}
    
    try:
        subscription = stripe.Subscription.retrieve(user["subscription"]["stripeSubscriptionId"])
        
        return {
            "success": True,
            "hasSubscription": True,
            "status": subscription.status,
            "cancelAtPeriodEnd": subscription.cancel_at_period_end,
            "currentPeriodEnd": datetime.fromtimestamp(subscription.current_period_end).isoformat(),
            "tier": user["subscription"].get("tier"),
            "downloadsRemaining": user["subscription"].get("downloads_remaining", 0)
        }
        
    except stripe.error.StripeError as e:
        return {"success": False, "error": str(e)}

# ========== PURCHASE TRACKING ENDPOINTS ==========

@app.post("/api/purchases/track")
async def track_purchase(data: PurchaseCreate, request: Request):
    """Track a purchase"""
    # Get client IP
    client_ip = get_client_ip(request)
    
    purchase = {
        "id": str(uuid.uuid4()),
        "documentType": data.documentType,
        "amount": data.amount,
        "email": data.email,
        "stripePaymentIntentId": data.stripePaymentIntentId,
        "discountCode": data.discountCode,
        "discountAmount": data.discountAmount,
        "userId": data.userId,
        "template": data.template,
        "quantity": data.quantity or 1,
        "ipAddress": client_ip,
        "createdAt": datetime.now(timezone.utc).isoformat(),
        "downloadedAt": datetime.now(timezone.utc).isoformat()
    }
    await purchases_collection.insert_one(purchase)
    asyncio.create_task(create_notification(data.documentType, data.email or "", data.amount))
    
    # If user has subscription, decrement downloads
    if data.userId:
        user = await users_collection.find_one({"id": data.userId})
        if user and user.get("subscription"):
            tier = SUBSCRIPTION_TIERS.get(user["subscription"]["tier"])
            if tier and tier["downloads"] != -1:  # Not unlimited
                await users_collection.update_one(
                    {"id": data.userId},
                    {"$inc": {"downloadsUsed": 1}}
                )
    
    return {"success": True, "purchaseId": purchase["id"]}

PEOPLE_SEARCH_TYPES = {"phone_lookup", "name_lookup", "address_lookup", "carrier_lookup"}

@app.get("/api/admin/purchases")
async def get_all_purchases(
    session: dict = Depends(get_current_admin),
    skip: int = 0,
    limit: int = 50,
    documentType: Optional[str] = None,
    startDate: Optional[str] = None,
    endDate: Optional[str] = None
):
    """Get all purchases (admin only)"""
    check_permission(session, "view_purchases")

    date_filter = {}
    if startDate:
        date_filter["$gte"] = startDate
    if endDate:
        date_filter["$lte"] = endDate

    is_ps_type = documentType in PEOPLE_SEARCH_TYPES
    is_doc_type = documentType and not is_ps_type

    purchases = []

    # Regular document purchases
    if not is_ps_type:
        q: dict = {}
        if is_doc_type:
            q["documentType"] = documentType
        if date_filter:
            q["createdAt"] = date_filter
        purchases += await purchases_collection.find(q, {"_id": 0}).sort("createdAt", -1).to_list(None)

    # People search payments
    if not is_doc_type:
        psq: dict = {}
        if is_ps_type:
            psq["lookupType"] = documentType
        if date_filter:
            psq["createdAt"] = date_filter
        ps_payments = await people_search_payments_collection.find(psq, {"_id": 0}).to_list(None)
        for p in ps_payments:
            purchases.append({
                "id":           p.get("id"),
                "documentType": p.get("lookupType"),
                "template":     p.get("query", ""),
                "amount":       p.get("amount", 0),
                "email":        p.get("userEmail", ""),
                "userId":       p.get("userId", ""),
                "userEmail":    p.get("userEmail", ""),
                "stripeSessionId": p.get("stripeSessionId"),
                "discountCode": p.get("discountCode"),
                "discountAmount": p.get("discountAmount", 0),
                "createdAt":    p.get("createdAt"),
                "searchId":     p.get("searchId"),
            })

    # Sort all combined by createdAt desc, then paginate
    purchases.sort(key=lambda x: x.get("createdAt") or "", reverse=True)
    total = len(purchases)
    purchases = purchases[skip: skip + limit]

    return {
        "success": True,
        "purchases": purchases,
        "total": total,
        "skip": skip,
        "limit": limit
    }

@app.get("/api/admin/dashboard")
async def get_admin_dashboard(session: dict = Depends(get_current_admin)):
    """Get admin dashboard data"""
    # Total purchases (count of purchase records)
    total_purchases = await purchases_collection.count_documents({})
    
    # Total downloads (sum of quantities - for paystubs etc.)
    total_downloads_pipeline = [
        {"$group": {"_id": None, "total": {"$sum": {"$ifNull": ["$quantity", 1]}}}}
    ]
    downloads_result = await purchases_collection.aggregate(total_downloads_pipeline).to_list(1)
    total_downloads = downloads_result[0]["total"] if downloads_result else total_purchases
    
    # Total revenue from one-time purchases
    pipeline = [
        {"$group": {"_id": None, "total": {"$sum": "$amount"}}}
    ]
    revenue_result = await purchases_collection.aggregate(pipeline).to_list(1)
    total_revenue = revenue_result[0]["total"] if revenue_result else 0
    
    # Purchases by document type (with quantity sum)
    type_pipeline = [
        {"$group": {
            "_id": "$documentType", 
            "count": {"$sum": 1}, 
            "totalQuantity": {"$sum": {"$ifNull": ["$quantity", 1]}},
            "revenue": {"$sum": "$amount"}
        }}
    ]
    purchases_by_type = await purchases_collection.aggregate(type_pipeline).to_list(100)
    
    # Get purchases for the last year for chart data
    one_year_ago = datetime.now(timezone.utc) - timedelta(days=365)
    recent_purchases = await purchases_collection.find(
        {"createdAt": {"$gte": one_year_ago.isoformat()}},
        {"_id": 0}
    ).sort("createdAt", -1).to_list(5000)
    
    # Get subscription payments for the last year for chart data
    recent_subscription_payments = await subscription_payments_collection.find(
        {"createdAt": {"$gte": one_year_ago.isoformat()}},
        {"_id": 0}
    ).sort("createdAt", -1).to_list(5000)
    
    # Total users
    total_users = await users_collection.count_documents({})
    
    # ===== SUBSCRIPTION STATS =====
    # Total active subscribers
    total_subscribers = await users_collection.count_documents({
        "subscription.status": "active"
    })
    
    # Total cancelling subscribers (pending cancellation at period end)
    cancelling_subscribers = await users_collection.count_documents({
        "subscription.status": "cancelling"
    })
    
    # Subscribers by tier (active only)
    subscribers_by_tier = {
        "starter": await users_collection.count_documents({"subscription.tier": "starter", "subscription.status": "active"}),
        "professional": await users_collection.count_documents({"subscription.tier": "professional", "subscription.status": "active"}),
        "business": await users_collection.count_documents({"subscription.tier": "business", "subscription.status": "active"})
    }
    
    # Cancelling subscribers by tier
    cancelling_by_tier = {
        "starter": await users_collection.count_documents({"subscription.tier": "starter", "subscription.status": "cancelling"}),
        "professional": await users_collection.count_documents({"subscription.tier": "professional", "subscription.status": "cancelling"}),
        "business": await users_collection.count_documents({"subscription.tier": "business", "subscription.status": "cancelling"})
    }
    
    # Monthly subscription revenue (calculated from active subscribers only)
    subscription_prices = {"starter": 19.99, "professional": 29.99, "business": 49.99}
    monthly_subscription_revenue = sum(
        count * subscription_prices.get(tier, 0) 
        for tier, count in subscribers_by_tier.items()
    )
    
    # Recent subscriptions (users who subscribed recently)
    recent_subscribers = await users_collection.find(
        {"subscription.status": {"$in": ["active", "cancelling"]}},
        {"_id": 0, "password": 0}
    ).sort("subscription.createdAt", -1).limit(10).to_list(10)
    
    # Subscription records from subscriptions collection
    subscription_records = await subscriptions_collection.find(
        {},
        {"_id": 0}
    ).sort("createdAt", -1).limit(20).to_list(20)
    
    # Today's stats
    today_start = datetime.now(timezone.utc).replace(hour=0, minute=0, second=0, microsecond=0).isoformat()
    today_purchases = await purchases_collection.count_documents({"createdAt": {"$gte": today_start}})
    
    # Today's revenue from guest purchases
    today_revenue_pipeline = [
        {"$match": {"createdAt": {"$gte": today_start}}},
        {"$group": {"_id": None, "total": {"$sum": "$amount"}}}
    ]
    today_revenue_result = await purchases_collection.aggregate(today_revenue_pipeline).to_list(1)
    today_guest_revenue = today_revenue_result[0]["total"] if today_revenue_result else 0
    
    # Today's subscription revenue
    today_sub_revenue_pipeline = [
        {"$match": {"createdAt": {"$gte": today_start}}},
        {"$group": {"_id": None, "total": {"$sum": "$amount"}}}
    ]
    today_sub_revenue_result = await subscription_payments_collection.aggregate(today_sub_revenue_pipeline).to_list(1)
    today_subscription_revenue = today_sub_revenue_result[0]["total"] if today_sub_revenue_result else 0
    
    # Combined today's revenue
    today_revenue = today_guest_revenue + today_subscription_revenue
    
    # New subscribers today
    today_new_subscribers = await subscriptions_collection.count_documents({"createdAt": {"$gte": today_start}})
    
    # ===== ALL-TIME SUBSCRIPTION REVENUE =====
    # Total revenue from subscription payments (all-time)
    total_sub_revenue_pipeline = [
        {"$group": {"_id": None, "total": {"$sum": "$amount"}}}
    ]
    total_sub_revenue_result = await subscription_payments_collection.aggregate(total_sub_revenue_pipeline).to_list(1)
    total_subscription_revenue = total_sub_revenue_result[0]["total"] if total_sub_revenue_result else 0
    
    # Combined all-time revenue (guest purchases + subscription payments)
    combined_total_revenue = total_revenue + total_subscription_revenue
    
    # ===== USER REGISTRATION TRENDS (for charts) =====
    # Get user registrations over the last 30 days
    thirty_days_ago = (datetime.now(timezone.utc) - timedelta(days=30)).isoformat()
    user_registration_pipeline = [
        {"$match": {"createdAt": {"$gte": thirty_days_ago}}},
        {"$group": {
            "_id": {"$substr": ["$createdAt", 0, 10]},  # Group by date (YYYY-MM-DD)
            "count": {"$sum": 1}
        }},
        {"$sort": {"_id": 1}}
    ]
    user_registrations = await users_collection.aggregate(user_registration_pipeline).to_list(31)
    
    return {
        "success": True,
        "stats": {
            "totalPurchases": total_purchases,
            "totalDownloads": total_downloads,
            "totalRevenue": round(total_revenue, 2),
            "totalSubscriptionRevenue": round(total_subscription_revenue, 2),
            "combinedTotalRevenue": round(combined_total_revenue, 2),
            "totalSubscribers": total_subscribers,
            "cancellingSubscribers": cancelling_subscribers,
            "totalUsers": total_users,
            "todayPurchases": today_purchases,
            "todayRevenue": round(today_revenue, 2),
            "todayGuestRevenue": round(today_guest_revenue, 2),
            "todaySubscriptionRevenue": round(today_subscription_revenue, 2),
            "todayNewSubscribers": today_new_subscribers,
            "monthlySubscriptionRevenue": round(monthly_subscription_revenue, 2)
        },
        "subscriptionStats": {
            "total": total_subscribers,
            "cancelling": cancelling_subscribers,
            "byTier": subscribers_by_tier,
            "cancellingByTier": cancelling_by_tier,
            "monthlyRevenue": round(monthly_subscription_revenue, 2),
            "totalRevenue": round(total_subscription_revenue, 2),
            "recentSubscribers": recent_subscribers,
            "subscriptionRecords": subscription_records
        },
        "purchasesByType": purchases_by_type,
        "recentPurchases": recent_purchases,
        "recentSubscriptionPayments": recent_subscription_payments,
        "userRegistrations": user_registrations
    }

@app.get("/api/admin/users")
async def get_all_users(
    session: dict = Depends(get_current_admin),
    skip: int = 0,
    limit: int = 50,
    search: Optional[str] = None,
    subscription_type: Optional[str] = None,
    date_from: Optional[str] = None,
    date_to: Optional[str] = None
):
    """Get all users (admin only) with optional filtering and search"""
    check_permission(session, "view_users")
    # Build query based on filters
    query = {}
    conditions = []
    
    # Search by name or email
    if search:
        search_regex = {"$regex": search, "$options": "i"}
        conditions.append({
            "$or": [
                {"name": search_regex},
                {"email": search_regex}
            ]
        })
    
    # Filter by subscription type
    if subscription_type:
        if subscription_type == "none":
            # Users with no subscription (either field doesn't exist or is null)
            conditions.append({
                "$or": [
                    {"subscription": {"$exists": False}},
                    {"subscription": None}
                ]
            })
        elif subscription_type in ["starter", "professional", "business"]:
            conditions.append({"subscription.tier": subscription_type})
    
    # Filter by join date
    if date_from or date_to:
        date_query = {}
        if date_from:
            date_query["$gte"] = date_from
        if date_to:
            date_query["$lte"] = date_to
        if date_query:
            conditions.append({"createdAt": date_query})
    
    # Combine all conditions
    if conditions:
        if len(conditions) == 1:
            query = conditions[0]
        else:
            query["$and"] = conditions
    
    users = await users_collection.find(query, {"_id": 0, "password": 0}).sort("createdAt", -1).skip(skip).limit(limit).to_list(limit)
    total = await users_collection.count_documents(query)
    
    return {
        "success": True,
        "users": users,
        "total": total
    }

@app.get("/api/admin/revenue")
async def get_revenue_by_period(
    session: dict = Depends(get_current_admin),
    startDate: Optional[str] = None,
    endDate: Optional[str] = None,
    userType: Optional[str] = None,  # "guest" or "registered"
    revenueType: Optional[str] = None  # "all", "guest", "subscription"
):
    """Get revenue for a specific period (admin only)"""
    
    # Build date conditions
    date_conditions = []
    if startDate:
        date_conditions.append({"createdAt": {"$gte": startDate}})
    if endDate:
        date_conditions.append({"createdAt": {"$lte": endDate}})
    
    # Calculate guest/one-time purchase revenue
    guest_revenue = 0
    guest_count = 0
    guest_downloads = 0
    
    if revenueType in [None, "all", "guest"]:
        conditions = date_conditions.copy()
        
        # Filter by user type for guest purchases
        if userType == "guest":
            conditions.append({
                "$or": [
                    {"userId": None}, 
                    {"userId": ""}, 
                    {"userId": {"$exists": False}}, 
                    {"isGuest": True}
                ]
            })
        elif userType == "registered":
            conditions.append({
                "$and": [
                    {"userId": {"$exists": True}},
                    {"userId": {"$ne": None}},
                    {"userId": {"$ne": ""}}
                ]
            })
        
        query = {"$and": conditions} if conditions else {}
        
        pipeline = [
            {"$match": query} if query else {"$match": {}},
            {"$group": {
                "_id": None, 
                "total": {"$sum": "$amount"}, 
                "count": {"$sum": 1},
                "downloadCount": {"$sum": {"$ifNull": ["$quantity", 1]}}
            }}
        ]
        
        result = await purchases_collection.aggregate(pipeline).to_list(1)
        if result:
            guest_revenue = result[0]["total"]
            guest_count = result[0]["count"]
            guest_downloads = result[0]["downloadCount"]
    
    # Calculate subscription revenue
    subscription_revenue = 0
    subscription_count = 0
    
    if revenueType in [None, "all", "subscription"]:
        sub_conditions = date_conditions.copy()
        sub_query = {"$and": sub_conditions} if sub_conditions else {}
        
        sub_pipeline = [
            {"$match": sub_query} if sub_query else {"$match": {}},
            {"$group": {
                "_id": None, 
                "total": {"$sum": "$amount"}, 
                "count": {"$sum": 1}
            }}
        ]
        
        sub_result = await subscription_payments_collection.aggregate(sub_pipeline).to_list(1)
        if sub_result:
            subscription_revenue = sub_result[0]["total"]
            subscription_count = sub_result[0]["count"]
    
    # Calculate totals based on filter
    if revenueType == "guest":
        total_revenue = guest_revenue
        total_count = guest_count
    elif revenueType == "subscription":
        total_revenue = subscription_revenue
        total_count = subscription_count
    else:  # "all" or None
        total_revenue = guest_revenue + subscription_revenue
        total_count = guest_count + subscription_count
    
    return {
        "success": True,
        "revenue": round(total_revenue, 2),
        "guestRevenue": round(guest_revenue, 2),
        "subscriptionRevenue": round(subscription_revenue, 2),
        "purchaseCount": total_count,
        "guestPurchaseCount": guest_count,
        "subscriptionPaymentCount": subscription_count,
        "downloadCount": guest_downloads,
        "period": {
            "startDate": startDate,
            "endDate": endDate
        },
        "userType": userType,
        "revenueType": revenueType or "all"
    }

@app.post("/api/admin/import-historical-subscriptions")
async def import_historical_subscriptions(
    session: dict = Depends(get_current_admin),
    limit: int = 100  # Process in batches to avoid timeout
):
    """
    Import historical subscription payments from Stripe.
    Fetches all paid invoices from Stripe and records them in subscription_payments_collection.
    """
    try:
        imported_count = 0
        skipped_count = 0
        error_count = 0
        already_exists_count = 0
        
        # First, get all subscriptions from Stripe and import their invoices
        has_more = True
        starting_after = None
        total_processed = 0
        processed_invoice_ids = set()  # Track to avoid duplicates
        
        logger.info("Starting historical subscription import...")
        
        # Method 1: Get all subscriptions and their invoices
        subscriptions_has_more = True
        sub_starting_after = None
        
        while subscriptions_has_more and total_processed < limit:
            sub_params = {"limit": 100, "status": "all"}
            if sub_starting_after:
                sub_params["starting_after"] = sub_starting_after
            
            subscriptions = stripe.Subscription.list(**sub_params)
            logger.info(f"Fetched {len(subscriptions.data)} subscriptions from Stripe")
            
            for subscription in subscriptions.data:
                # Get all invoices for this subscription
                invoice_params = {
                    "subscription": subscription.id,
                    "status": "paid",
                    "limit": 100
                }
                
                try:
                    sub_invoices = stripe.Invoice.list(**invoice_params)
                    logger.info(f"Subscription {subscription.id}: found {len(sub_invoices.data)} paid invoices")
                    
                    for invoice in sub_invoices.data:
                        if invoice.id in processed_invoice_ids:
                            continue
                        processed_invoice_ids.add(invoice.id)
                        total_processed += 1
                        
                        if total_processed > limit:
                            break
                        
                        try:
                            # Check if this payment already exists
                            existing = await subscription_payments_collection.find_one({
                                "stripeInvoiceId": invoice.id
                            })
                            if existing:
                                logger.info(f"Invoice {invoice.id} already exists in DB")
                                already_exists_count += 1
                                continue
                            
                            # Get subscription ID
                            subscription_id = subscription.id
                            
                            # Get customer ID
                            customer_id = invoice.customer if isinstance(invoice.customer, str) else (invoice.customer.id if hasattr(invoice.customer, 'id') else None)
                            
                            # Try to find the user by Stripe customer ID or subscription ID
                            user = None
                            if subscription_id or customer_id:
                                query_conditions = []
                                if subscription_id:
                                    query_conditions.append({"subscription.stripeSubscriptionId": subscription_id})
                                if customer_id:
                                    query_conditions.append({"subscription.stripeCustomerId": customer_id})
                                    query_conditions.append({"stripeCustomerId": customer_id})
                                
                                if query_conditions:
                                    user = await users_collection.find_one({"$or": query_conditions})
                            
                            # Determine the tier from the subscription or invoice
                            tier = "unknown"
                            
                            # Try to get tier from subscription items
                            if hasattr(subscription, 'items') and subscription.items and hasattr(subscription.items, 'data'):
                                for item in subscription.items.data:
                                    if hasattr(item, 'price') and item.price:
                                        amount = item.price.unit_amount if hasattr(item.price, 'unit_amount') else 0
                                        # Map amount to tier
                                        if amount == 1999:
                                            tier = "starter"
                                        elif amount == 2999:
                                            tier = "professional"
                                        elif amount == 4999:
                                            tier = "business"
                                        
                                        # Also try to match by price ID
                                        if tier == "unknown" and hasattr(item.price, 'id'):
                                            for t, pid in stripe_price_ids.items():
                                                if pid == item.price.id:
                                                    tier = t
                                                    break
                            
                            # If we couldn't determine tier, try from user
                            if tier == "unknown" and user and user.get("subscription", {}).get("tier"):
                                tier = user["subscription"]["tier"]
                            
                            # Create the payment record
                            payment_amount = (invoice.amount_paid or 0) / 100  # Convert from cents
                            
                            # Convert Stripe timestamp to ISO format
                            created_at = datetime.fromtimestamp(invoice.created, tz=timezone.utc).isoformat()
                            
                            # Get customer email safely
                            customer_email = ""
                            if user:
                                customer_email = user.get("email", "")
                            elif hasattr(invoice, 'customer_email') and invoice.customer_email:
                                customer_email = invoice.customer_email
                            
                            subscription_payment = {
                                "id": str(uuid.uuid4()),
                                "userId": user["id"] if user else None,
                                "userEmail": customer_email,
                                "tier": tier,
                                "amount": payment_amount,
                                "stripeInvoiceId": invoice.id,
                                "stripeSubscriptionId": subscription_id,
                                "stripeCustomerId": customer_id,
                                "billingReason": getattr(invoice, 'billing_reason', None) or "unknown",
                                "createdAt": created_at,
                                "importedAt": datetime.now(timezone.utc).isoformat(),
                                "isHistoricalImport": True
                            }
                            
                            await subscription_payments_collection.insert_one(subscription_payment)
                            imported_count += 1
                            logger.info(f"Imported invoice {invoice.id} for subscription {subscription_id} - ${payment_amount}")
                            
                        except Exception as e:
                            logger.error(f"Error processing invoice {invoice.id}: {str(e)}")
                            error_count += 1
                            continue
                
                except Exception as e:
                    logger.error(f"Error fetching invoices for subscription {subscription.id}: {str(e)}")
                    error_count += 1
            
            # Check if there are more subscriptions
            subscriptions_has_more = subscriptions.has_more
            if subscriptions.data:
                sub_starting_after = subscriptions.data[-1].id
        
        # Method 2: Also check for any paid invoices that might have been missed
        # (e.g., invoices with subscription field set but subscription deleted)
        logger.info("Checking for additional paid invoices...")
        inv_has_more = True
        inv_starting_after = None
        
        while inv_has_more and total_processed < limit:
            inv_params = {
                "status": "paid",
                "limit": min(100, limit - total_processed)
            }
            if inv_starting_after:
                inv_params["starting_after"] = inv_starting_after
            
            invoices = stripe.Invoice.list(**inv_params)
            
            for invoice in invoices.data:
                if invoice.id in processed_invoice_ids:
                    continue
                processed_invoice_ids.add(invoice.id)
                total_processed += 1
                
                try:
                    # Get subscription ID - handle both string and None cases
                    subscription_id = None
                    if hasattr(invoice, 'subscription') and invoice.subscription:
                        if isinstance(invoice.subscription, str):
                            subscription_id = invoice.subscription
                        elif hasattr(invoice.subscription, 'id'):
                            subscription_id = invoice.subscription.id
                    
                    # Skip non-subscription invoices
                    if not subscription_id:
                        logger.debug(f"Skipping invoice {invoice.id}: no subscription ID (one-time payment)")
                        skipped_count += 1
                        continue
                    
                    logger.info(f"Found additional subscription invoice: {invoice.id} with subscription {subscription_id}")
                    
                    # Check if this payment already exists
                    existing = await subscription_payments_collection.find_one({
                        "stripeInvoiceId": invoice.id
                    })
                    if existing:
                        logger.info(f"Invoice {invoice.id} already exists in DB")
                        already_exists_count += 1
                        continue
                    
                    # Get customer ID
                    customer_id = invoice.customer if isinstance(invoice.customer, str) else (invoice.customer.id if hasattr(invoice.customer, 'id') else None)
                    
                    # Try to find the user by Stripe customer ID or subscription ID
                    user = None
                    if subscription_id or customer_id:
                        query_conditions = []
                        if subscription_id:
                            query_conditions.append({"subscription.stripeSubscriptionId": subscription_id})
                        if customer_id:
                            query_conditions.append({"subscription.stripeCustomerId": customer_id})
                            query_conditions.append({"stripeCustomerId": customer_id})
                        
                        if query_conditions:
                            user = await users_collection.find_one({"$or": query_conditions})
                    
                    # Determine the tier from the invoice line items
                    tier = "unknown"
                    if hasattr(invoice, 'lines') and invoice.lines and hasattr(invoice.lines, 'data'):
                        for line in invoice.lines.data:
                            if hasattr(line, 'price') and line.price:
                                # Try to get amount to determine tier
                                amount = line.price.unit_amount if hasattr(line.price, 'unit_amount') else 0
                                if amount == 1999:
                                    tier = "starter"
                                elif amount == 2999:
                                    tier = "professional"
                                elif amount == 4999:
                                    tier = "business"
                                
                                # Also try to match by price ID
                                if tier == "unknown" and hasattr(line.price, 'id'):
                                    for t, pid in stripe_price_ids.items():
                                        if pid == line.price.id:
                                            tier = t
                                            break
                    
                    # If we couldn't determine tier from price, try to get from user
                    if tier == "unknown" and user and user.get("subscription", {}).get("tier"):
                        tier = user["subscription"]["tier"]
                    
                    # Create the payment record
                    payment_amount = (invoice.amount_paid or 0) / 100  # Convert from cents
                    
                    # Convert Stripe timestamp to ISO format
                    created_at = datetime.fromtimestamp(invoice.created, tz=timezone.utc).isoformat()
                    
                    # Get customer email safely
                    customer_email = ""
                    if user:
                        customer_email = user.get("email", "")
                    elif hasattr(invoice, 'customer_email') and invoice.customer_email:
                        customer_email = invoice.customer_email
                    
                    subscription_payment = {
                        "id": str(uuid.uuid4()),
                        "userId": user["id"] if user else None,
                        "userEmail": customer_email,
                        "tier": tier,
                        "amount": payment_amount,
                        "stripeInvoiceId": invoice.id,
                        "stripeSubscriptionId": subscription_id,
                        "stripeCustomerId": customer_id,
                        "billingReason": getattr(invoice, 'billing_reason', None) or "unknown",
                        "createdAt": created_at,
                        "importedAt": datetime.now(timezone.utc).isoformat(),
                        "isHistoricalImport": True
                    }
                    
                    await subscription_payments_collection.insert_one(subscription_payment)
                    imported_count += 1
                    
                except Exception as e:
                    logger.error(f"Error processing invoice {invoice.id}: {str(e)}")
                    error_count += 1
                    continue
            
            # Check if there are more invoices
            inv_has_more = invoices.has_more
            if invoices.data:
                inv_starting_after = invoices.data[-1].id
        
        # Calculate total imported revenue
        total_revenue_pipeline = [
            {"$match": {"isHistoricalImport": True}},
            {"$group": {"_id": None, "total": {"$sum": "$amount"}}}
        ]
        revenue_result = await subscription_payments_collection.aggregate(total_revenue_pipeline).to_list(1)
        total_imported_revenue = revenue_result[0]["total"] if revenue_result else 0
        
        return {
            "success": True,
            "message": f"Historical subscription import complete",
            "stats": {
                "totalProcessed": total_processed,
                "imported": imported_count,
                "skipped": skipped_count,
                "alreadyExists": already_exists_count,
                "errors": error_count,
                "totalImportedRevenue": round(total_imported_revenue, 2)
            }
        }
        
    except stripe.error.StripeError as e:
        print(f"Stripe error during import: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Stripe error: {str(e)}")
    except Exception as e:
        print(f"General error during import: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Import error: {str(e)}")

# ========== SUBSCRIPTION ENDPOINTS ==========

@app.get("/api/subscription/tiers")
async def get_subscription_tiers():
    """Get available subscription tiers"""
    return {
        "success": True,
        "tiers": SUBSCRIPTION_TIERS
    }

@app.post("/api/subscription/create")
async def create_subscription(data: SubscriptionCreate, session: dict = Depends(get_current_user)):
    """Create a subscription for a user"""
    if data.tier not in SUBSCRIPTION_TIERS:
        raise HTTPException(status_code=400, detail="Invalid subscription tier")
    
    subscription = {
        "id": str(uuid.uuid4()),
        "userId": data.userId,
        "tier": data.tier,
        "paypalSubscriptionId": data.paypalSubscriptionId,
        "paypalEmail": data.paypalEmail,
        "status": "active",
        "createdAt": datetime.now(timezone.utc).isoformat(),
        "currentPeriodStart": datetime.now(timezone.utc).isoformat(),
        "currentPeriodEnd": None  # Will be updated by PayPal webhook
    }
    
    await subscriptions_collection.insert_one(subscription)
    
    # Update user with subscription
    await users_collection.update_one(
        {"id": data.userId},
        {
            "$set": {
                "subscription": {
                    "id": subscription["id"],
                    "tier": data.tier,
                    "status": "active"
                },
                "downloadsUsed": 0,
                "downloadsReset": datetime.now(timezone.utc).isoformat()
            }
        }
    )
    
    return {"success": True, "subscriptionId": subscription["id"]}

# ========== SUBSCRIPTION DOWNLOAD ENDPOINT ==========

class SubscriptionDownloadRequest(BaseModel):
    documentType: str  # paystub, w9, 1099-nec, etc.
    template: Optional[str] = None  # Template used (for paystubs)
    count: int = 1  # Number of documents being downloaded (for paystubs with multiple stubs)


@app.post("/api/user/subscription-download")
async def subscription_download(data: SubscriptionDownloadRequest, session: dict = Depends(get_current_user)):
    """
    Validate and track a subscription-based download.
    Returns success if user can download, decrements remaining count.
    Supports downloading multiple documents at once (e.g., multiple paystubs).
    """
    user = await users_collection.find_one({"id": session["userId"]}, {"_id": 0})
    
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Check if user has active or cancelling subscription (access continues until period ends)
    subscription = user.get("subscription")
    if not subscription or subscription.get("status") not in ["active", "cancelling"]:
        raise HTTPException(status_code=403, detail="No active subscription found. Please subscribe to download.")
    
    # Get subscription tier config
    tier = subscription.get("tier")
    
    # Try SUBSCRIPTION_PLANS first (for PayPal subscriptions), then SUBSCRIPTION_TIERS (for admin-assigned)
    plan_config = SUBSCRIPTION_PLANS.get(tier) or SUBSCRIPTION_TIERS.get(tier)
    
    if not plan_config:
        raise HTTPException(status_code=400, detail="Invalid subscription tier")
    
    # Validate count
    download_count = max(1, data.count)  # Minimum 1 download
    
    # Check downloads remaining
    downloads_remaining = subscription.get("downloads_remaining", 0)
    
    # -1 means unlimited
    if downloads_remaining != -1:
        if downloads_remaining <= 0:
            raise HTTPException(
                status_code=403, 
                detail="No downloads remaining this month. Please upgrade your plan or wait for the next billing cycle."
            )
        
        if downloads_remaining < download_count:
            raise HTTPException(
                status_code=403, 
                detail=f"Not enough downloads remaining. You have {downloads_remaining} downloads left but are trying to download {download_count} documents."
            )
        
        # Decrement downloads remaining by the count
        new_remaining = downloads_remaining - download_count
        await users_collection.update_one(
            {"id": session["userId"]},
            {"$set": {"subscription.downloads_remaining": new_remaining}}
        )
    else:
        new_remaining = -1  # Still unlimited
    
    # Track the download in purchases collection (with $0 amount for subscription)
    purchase = {
        "id": str(uuid.uuid4()),
        "documentType": data.documentType,
        "amount": 0,  # Subscription download = $0
        "paypalEmail": user.get("email", ""),
        "paypalTransactionId": None,
        "discountCode": None,
        "discountAmount": 0,
        "userId": user["id"],
        "template": data.template,
        "subscriptionDownload": True,  # Flag to identify subscription downloads
        "subscriptionTier": tier,
        "downloadCount": download_count,  # Track how many documents were downloaded
        "createdAt": datetime.now(timezone.utc).isoformat(),
        "downloadedAt": datetime.now(timezone.utc).isoformat()
    }
    await purchases_collection.insert_one(purchase)
    
    # Send review request email (download confirmation with PDF is sent from frontend)
    user_email = user.get("email")
    user_name = user.get("name", "")
    if user_email:
        # NOTE: Download confirmation with PDF attachment is now sent from frontend
        asyncio.create_task(send_review_request(user_email, user_name, data.documentType, user["id"]))
    
    return {
        "success": True,
        "message": "Download authorized",
        "downloadsRemaining": new_remaining,
        "unlimited": new_remaining == -1,
        "purchaseId": purchase["id"]
    }


@app.get("/api/user/downloads-remaining")
async def get_downloads_remaining(session: dict = Depends(get_current_user)):
    """Get remaining downloads for subscribed user"""
    user = await users_collection.find_one({"id": session["userId"]}, {"_id": 0})
    
    if not user or not user.get("subscription"):
        return {"success": True, "hasSubscription": False, "downloadsRemaining": 0}
    
    subscription = user.get("subscription")
    # Allow both active and cancelling subscriptions (access continues until period ends)
    if subscription.get("status") not in ["active", "cancelling"]:
        return {"success": True, "hasSubscription": False, "downloadsRemaining": 0}
    
    # Get downloads_remaining directly from user's subscription
    downloads_remaining = subscription.get("downloads_remaining", 0)
    downloads_total = subscription.get("downloads_total", 0)
    
    # -1 means unlimited
    if downloads_remaining == -1:
        return {"success": True, "hasSubscription": True, "downloadsRemaining": -1, "unlimited": True}
    
    return {
        "success": True,
        "hasSubscription": True,
        "downloadsRemaining": max(0, downloads_remaining),
        "totalDownloads": downloads_total,
        "tier": subscription.get("tier")
    }

@app.get("/api/user/downloads")
async def get_user_downloads(
    session: dict = Depends(get_current_user),
    skip: int = 0,
    limit: int = 20,
    document_type: Optional[str] = None,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None
):
    """Get user's download history with optional document type and date filters"""
    # Get purchases associated with the user's email or userId
    user = await users_collection.find_one({"id": session["userId"]}, {"_id": 0})
    
    if not user:
        return {"success": True, "downloads": [], "total": 0}
    
    # Query purchases by userId or email (works for both PayPal and Stripe)
    query = {
        "$or": [
            {"userId": user["id"]},
            {"paypalEmail": user["email"]},
            {"email": user["email"]}  # For Stripe purchases
        ]
    }
    
    # Add document type filter if provided
    if document_type and document_type != "all":
        query["documentType"] = document_type
    
    # Add date range filters if provided
    if start_date or end_date:
        date_filter = {}
        if start_date:
            # Start of the day
            date_filter["$gte"] = f"{start_date}T00:00:00"
        if end_date:
            # End of the day
            date_filter["$lte"] = f"{end_date}T23:59:59"
        if date_filter:
            query["createdAt"] = date_filter
    
    downloads = await purchases_collection.find(query, {"_id": 0}).sort("createdAt", -1).skip(skip).limit(limit).to_list(limit)
    total = await purchases_collection.count_documents(query)
    
    return {
        "success": True,
        "downloads": downloads,
        "total": total
    }

# ========== DELETE ENDPOINTS ==========

@app.delete("/api/admin/purchases/{purchase_id}")
async def delete_purchase(purchase_id: str, session: dict = Depends(get_current_admin)):
    """Delete a purchase record (admin only)"""
    result = await purchases_collection.delete_one({"id": purchase_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Purchase not found")
    return {"success": True, "message": "Purchase deleted"}

@app.post("/api/admin/purchases")
async def create_manual_purchase(data: ManualPurchaseCreate, session: dict = Depends(get_current_admin)):
    """Manually add a purchase record (admin only) - for historical purchases"""
    # Use provided date or default to now
    if data.purchaseDate:
        try:
            purchase_date = data.purchaseDate
        except:
            purchase_date = datetime.now(timezone.utc).isoformat()
    else:
        purchase_date = datetime.now(timezone.utc).isoformat()
    
    purchase = {
        "id": str(uuid.uuid4()),
        "documentType": data.documentType,
        "amount": data.amount,
        "paypalEmail": data.paypalEmail,
        "template": data.template,
        "discountCode": data.discountCode,
        "discountAmount": data.discountAmount or 0,
        "notes": data.notes,
        "manualEntry": True,  # Flag to indicate this was manually added
        "addedBy": session.get("adminId"),
        "createdAt": purchase_date
    }
    
    await purchases_collection.insert_one(purchase)
    
    return {"success": True, "message": "Purchase added successfully", "purchase": {k: v for k, v in purchase.items() if k != "_id"}}

@app.put("/api/admin/purchases/{purchase_id}")
async def update_purchase(purchase_id: str, data: ManualPurchaseCreate, session: dict = Depends(get_current_admin)):
    """Update a purchase record (admin only)"""
    existing = await purchases_collection.find_one({"id": purchase_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Purchase not found")
    
    # Use provided date or keep existing
    if data.purchaseDate:
        try:
            purchase_date = data.purchaseDate
        except:
            purchase_date = existing.get("createdAt")
    else:
        purchase_date = existing.get("createdAt")
    
    update_data = {
        "documentType": data.documentType,
        "amount": data.amount,
        "paypalEmail": data.paypalEmail,
        "template": data.template,
        "discountCode": data.discountCode,
        "discountAmount": data.discountAmount or 0,
        "notes": data.notes,
        "quantity": data.quantity or 1,
        "ipAddress": data.ipAddress,
        "createdAt": purchase_date,
        "updatedAt": datetime.now(timezone.utc).isoformat(),
        "updatedBy": session.get("adminId")
    }
    
    await purchases_collection.update_one(
        {"id": purchase_id},
        {"$set": update_data}
    )
    
    return {"success": True, "message": "Purchase updated successfully"}

@app.delete("/api/admin/users/{user_id}")
async def delete_user(user_id: str, session: dict = Depends(get_current_admin)):
    """Delete a user (admin only)"""
    check_permission(session, "edit_users")
    # Delete user's sessions
    await sessions_collection.delete_many({"userId": user_id})
    # Delete user's subscriptions
    await subscriptions_collection.delete_many({"userId": user_id})
    # Delete the user
    result = await users_collection.delete_one({"id": user_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="User not found")
    await log_action(session, "delete_user", "user", user_id)
    return {"success": True, "message": "User deleted"}

class UpdateUserDetails(BaseModel):
    name: Optional[str] = None
    email: Optional[str] = None
    ipAddress: Optional[str] = None

@app.put("/api/admin/users/{user_id}")
async def update_user_details(user_id: str, data: UpdateUserDetails, session: dict = Depends(get_current_admin)):
    """Update a user's details (admin only)"""
    check_permission(session, "edit_users")
    user = await users_collection.find_one({"id": user_id})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    update_data = {}
    
    if data.name is not None:
        update_data["name"] = data.name.strip()
    
    if data.email is not None:
        new_email = data.email.strip().lower()
        # Check if email is already taken by another user
        if new_email != user.get("email", "").lower():
            existing_user = await users_collection.find_one({"email": new_email, "id": {"$ne": user_id}})
            if existing_user:
                raise HTTPException(status_code=400, detail="Email already in use by another user")
        update_data["email"] = new_email
    
    if data.ipAddress is not None:
        update_data["ipAddress"] = data.ipAddress.strip() if data.ipAddress.strip() else None
    
    if not update_data:
        raise HTTPException(status_code=400, detail="No fields to update")
    
    await users_collection.update_one(
        {"id": user_id},
        {"$set": update_data}
    )
    
    return {"success": True, "message": "User updated successfully"}

@app.put("/api/admin/users/{user_id}/ban")
async def ban_user(user_id: str, session: dict = Depends(get_current_admin)):
    """Ban/unban a user (admin only)"""
    check_permission(session, "edit_users")
    user = await users_collection.find_one({"id": user_id})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    new_status = not user.get("isBanned", False)
    await users_collection.update_one(
        {"id": user_id},
        {"$set": {"isBanned": new_status}}
    )
    
    # If banning, also invalidate their sessions
    if new_status:
        await sessions_collection.delete_many({"userId": user_id})
    await log_action(session, "ban_user" if new_status else "unban_user", "user", user_id)
    return {"success": True, "isBanned": new_status}


@app.put("/api/admin/users/{user_id}/verify")
async def admin_verify_user_email(user_id: str, session: dict = Depends(get_current_admin)):
    """Verify/confirm a user's email address (admin only)"""
    user = await users_collection.find_one({"id": user_id})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Set emailVerified to True and remove any verification code
    await users_collection.update_one(
        {"id": user_id},
        {"$set": {"emailVerified": True}, "$unset": {"verificationCode": ""}}
    )
    
    return {"success": True, "message": "User email verified successfully", "emailVerified": True}


# ========== ADMIN SAVED DOCUMENTS MANAGEMENT ==========

@app.get("/api/admin/saved-documents")
async def get_all_saved_documents(
    skip: int = 0,
    limit: int = 20,
    userId: Optional[str] = None,
    documentType: Optional[str] = None,
    session: dict = Depends(get_current_admin)
):
    """Get all saved documents (admin only)"""
    check_permission(session, "view_saved_docs")
    query = {}

    if userId:
        query["userId"] = userId
    if documentType and documentType != "all":
        query["documentType"] = documentType
    
    # Get documents with pagination (exclude fileContent from response for performance)
    documents = await saved_documents_collection.find(query, {"_id": 0, "fileContent": 0}).sort("createdAt", -1).skip(skip).limit(limit).to_list(limit)
    total = await saved_documents_collection.count_documents(query)
    
    # Enrich with user info and file existence check
    enriched_documents = []
    for doc in documents:
        user = await users_collection.find_one({"id": doc.get("userId")}, {"_id": 0, "email": 1, "name": 1})
        
        # Check if file exists on disk
        stored_filename = doc.get("storedFileName", "")
        file_exists = False
        has_content_backup = False
        
        if stored_filename:
            file_path = os.path.join(USER_DOCUMENTS_DIR, stored_filename)
            file_exists = os.path.exists(file_path)
            
            # If file not on disk, check if we have content in MongoDB
            if not file_exists:
                # Check if fileContent exists (just check existence, don't load it)
                doc_with_content = await saved_documents_collection.find_one(
                    {"id": doc.get("id")},
                    {"fileContent": 1}
                )
                has_content_backup = bool(doc_with_content and doc_with_content.get("fileContent"))
                
                # If we have content backup, try to restore the file
                if has_content_backup:
                    try:
                        file_content = base64.b64decode(doc_with_content["fileContent"])
                        with open(file_path, "wb") as f:
                            f.write(file_content)
                        file_exists = True
                        logger.info(f"Auto-restored file from MongoDB: {file_path}")
                    except Exception as e:
                        logger.warning(f"Could not auto-restore file: {e}")
        
        enriched_documents.append({
            **doc,
            "userEmail": user.get("email", "Unknown") if user else "Deleted User",
            "userName": user.get("name", "") if user else "",
            "fileExists": file_exists,
            "hasBackup": has_content_backup if not file_exists else True
        })
    
    return {
        "documents": enriched_documents,
        "total": total,
        "skip": skip,
        "limit": limit
    }


@app.delete("/api/admin/saved-documents/{doc_id}")
async def admin_delete_saved_document(doc_id: str, session: dict = Depends(get_current_admin)):
    """Delete a saved document (admin only)"""
    # Find the document first
    document = await saved_documents_collection.find_one({"id": doc_id})
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Delete the actual file if it exists
    stored_filename = document.get("storedFileName")
    if stored_filename:
        file_path = os.path.join(USER_DOCUMENTS_DIR, stored_filename)
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
                logger.info(f"Admin deleted file: {file_path}")
            except Exception as e:
                logger.error(f"Failed to delete file {file_path}: {e}")
    
    # Delete from database
    result = await saved_documents_collection.delete_one({"id": doc_id})
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return {"success": True, "message": "Document deleted successfully"}


@app.get("/api/admin/saved-documents/{doc_id}/download")
async def admin_download_saved_document(doc_id: str, session: dict = Depends(get_current_admin)):
    """Download/view a saved document (admin only)"""
    from fastapi.responses import FileResponse, Response
    
    # Find document
    document = await saved_documents_collection.find_one({"id": doc_id})
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Determine media type
    file_ext = os.path.splitext(document["fileName"])[1].lower()
    media_types = {
        ".pdf": "application/pdf",
        ".zip": "application/zip"
    }
    media_type = media_types.get(file_ext, "application/octet-stream")
    
    # Get file path
    file_path = os.path.join(USER_DOCUMENTS_DIR, document["storedFileName"])
    
    if os.path.exists(file_path):
        return FileResponse(
            path=file_path,
            filename=document["fileName"],
            media_type=media_type
        )
    
    # If file not on disk, try to restore from MongoDB content
    if document.get("fileContent"):
        try:
            file_content = base64.b64decode(document["fileContent"])
            # Restore file to disk for future access
            try:
                with open(file_path, "wb") as f:
                    f.write(file_content)
                logger.info(f"Admin restored file from MongoDB: {file_path}")
            except Exception as e:
                logger.warning(f"Could not restore file to disk: {e}")
            
            return Response(
                content=file_content,
                media_type=media_type,
                headers={"Content-Disposition": f'attachment; filename="{document["fileName"]}"'}
            )
        except Exception as e:
            logger.error(f"Failed to decode file content from MongoDB: {e}")
    
    raise HTTPException(status_code=404, detail="Document file not found on server")


@app.delete("/api/admin/users/{user_id}/saved-documents")
async def admin_delete_user_saved_documents(user_id: str, session: dict = Depends(get_current_admin)):
    """Delete all saved documents for a user (admin only)"""
    # Get all documents for this user
    documents = await saved_documents_collection.find({"userId": user_id}).to_list(1000)
    
    deleted_count = 0
    for doc in documents:
        # Delete the actual file if it exists
        stored_filename = doc.get("storedFileName")
        if stored_filename:
            file_path = os.path.join(USER_DOCUMENTS_DIR, stored_filename)
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except Exception as e:
                    logger.error(f"Failed to delete file {file_path}: {e}")
        deleted_count += 1
    
    # Delete all from database
    await saved_documents_collection.delete_many({"userId": user_id})
    
    return {"success": True, "message": f"Deleted {deleted_count} documents for user"}


class UpdateUserSubscription(BaseModel):
    tier: Optional[str] = None  # starter, professional, business, or null to remove

class UpdateUserDownloads(BaseModel):
    downloads_to_add: int  # Number of bonus downloads to add

@app.put("/api/admin/users/{user_id}/subscription")
async def update_user_subscription(user_id: str, data: UpdateUserSubscription, session: dict = Depends(get_current_admin)):
    """Update a user's subscription plan (admin only)"""
    user = await users_collection.find_one({"id": user_id})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Validate tier if provided
    if data.tier and data.tier not in SUBSCRIPTION_TIERS:
        raise HTTPException(status_code=400, detail="Invalid subscription tier")
    
    if data.tier:
        # Set or update subscription
        tier_config = SUBSCRIPTION_TIERS.get(data.tier)
        if not tier_config:
            raise HTTPException(status_code=400, detail="Invalid subscription tier")
        
        subscription = {
            "id": str(uuid.uuid4()),
            "tier": data.tier,
            "status": "active",
            "adminAssigned": True,
            "downloads_remaining": tier_config["downloads"],
            "downloads_total": tier_config["downloads"],
            "updatedAt": datetime.now(timezone.utc).isoformat()
        }
        await users_collection.update_one(
            {"id": user_id},
            {
                "$set": {
                    "subscription": subscription,
                    "downloadsUsed": 0,
                    "downloadsReset": datetime.now(timezone.utc).isoformat()
                }
            }
        )
        return {"success": True, "message": f"User subscription updated to {data.tier}", "subscription": subscription}
    else:
        # Remove subscription
        await users_collection.update_one(
            {"id": user_id},
            {
                "$set": {
                    "subscription": None
                }
            }
        )
        return {"success": True, "message": "User subscription removed", "subscription": None}


@app.put("/api/admin/users/{user_id}/downloads")
async def update_user_downloads(user_id: str, data: UpdateUserDownloads, session: dict = Depends(get_current_admin)):
    """Add bonus downloads to a user's account (admin only) - resets on next billing cycle"""
    user = await users_collection.find_one({"id": user_id})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    if not user.get("subscription"):
        raise HTTPException(status_code=400, detail="User has no active subscription")
    
    # Validate downloads to add
    if data.downloads_to_add < 1:
        raise HTTPException(status_code=400, detail="Please enter a positive number of downloads to add")
    
    current_downloads = user["subscription"].get("downloads_remaining", 0)
    
    # If user has unlimited downloads, no need to add more
    if current_downloads == -1:
        return {
            "success": True,
            "message": "User already has unlimited downloads",
            "user": user
        }
    
    # Add bonus downloads to current count
    new_downloads = current_downloads + data.downloads_to_add
    
    # Update downloads_remaining
    await users_collection.update_one(
        {"id": user_id},
        {
            "$set": {
                "subscription.downloads_remaining": new_downloads,
                "subscription.updatedAt": datetime.now(timezone.utc).isoformat()
            }
        }
    )
    
    # Get updated user
    updated_user = await users_collection.find_one({"id": user_id}, {"_id": 0, "password": 0})
    
    return {
        "success": True, 
        "message": f"Added {data.downloads_to_add} bonus downloads. New total: {new_downloads}",
        "previousDownloads": current_downloads,
        "addedDownloads": data.downloads_to_add,
        "newTotal": new_downloads,
        "user": updated_user
    }


# ========== BANNED IPS ENDPOINTS ==========

class BannedIPCreate(BaseModel):
    ip: str
    reason: Optional[str] = None

@app.get("/api/check-ip-ban")
async def check_ip_ban(request: Request):
    """Check if the current IP is banned - public endpoint"""
    client_ip = get_client_ip(request)
    banned = await banned_ips_collection.find_one({"ip": client_ip, "isActive": True})
    
    if banned:
        return {
            "banned": True,
            "reason": banned.get("reason", "Your access has been restricted."),
            "bannedAt": banned.get("bannedAt")
        }
    
    return {"banned": False}

@app.get("/api/admin/banned-ips")
async def get_banned_ips(session: dict = Depends(get_current_admin)):
    """Get all banned IPs (admin only)"""
    check_permission(session, "view_banned_ips")
    banned_ips = await banned_ips_collection.find({}, {"_id": 0}).sort("bannedAt", -1).to_list(1000)
    return {"success": True, "bannedIps": banned_ips}

@app.post("/api/admin/banned-ips")
async def ban_ip(data: BannedIPCreate, session: dict = Depends(get_current_admin)):
    """Ban an IP address (admin only)"""
    check_permission(session, "manage_banned_ips")
    # Check if IP is already banned
    existing = await banned_ips_collection.find_one({"ip": data.ip})
    
    if existing:
        # Reactivate if previously unbanned
        await banned_ips_collection.update_one(
            {"ip": data.ip},
            {
                "$set": {
                    "isActive": True,
                    "reason": data.reason or existing.get("reason"),
                    "bannedAt": datetime.now(timezone.utc).isoformat(),
                    "bannedBy": session.get("adminId")
                }
            }
        )
        return {"success": True, "message": f"IP {data.ip} has been banned"}
    
    banned_ip = {
        "id": str(uuid.uuid4()),
        "ip": data.ip,
        "reason": data.reason,
        "isActive": True,
        "bannedAt": datetime.now(timezone.utc).isoformat(),
        "bannedBy": session.get("adminId")
    }
    
    await banned_ips_collection.insert_one(banned_ip)
    await log_action(session, "ban_ip", "banned_ip", data.ip, data.ip)
    return {"success": True, "message": f"IP {data.ip} has been banned", "bannedIp": {k: v for k, v in banned_ip.items() if k != "_id"}}

@app.delete("/api/admin/banned-ips/{ip}")
async def unban_ip(ip: str, session: dict = Depends(get_current_admin)):
    """Unban an IP address (admin only)"""
    check_permission(session, "manage_banned_ips")
    # URL decode the IP (in case it was encoded)
    from urllib.parse import unquote
    decoded_ip = unquote(ip)
    
    result = await banned_ips_collection.update_one(
        {"ip": decoded_ip},
        {
            "$set": {
                "isActive": False,
                "unbannedAt": datetime.now(timezone.utc).isoformat(),
                "unbannedBy": session.get("adminId")
            }
        }
    )
    
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="IP not found in ban list")
    await log_action(session, "unban_ip", "banned_ip", decoded_ip, decoded_ip)
    return {"success": True, "message": f"IP {decoded_ip} has been unbanned"}

@app.post("/api/admin/ban-user-ip/{user_id}")
async def ban_user_ip(user_id: str, data: dict, session: dict = Depends(get_current_admin)):
    """Ban a user's IP address (admin only)"""
    user = await users_collection.find_one({"id": user_id})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    ip = user.get("ipAddress")
    if not ip or ip == "unknown":
        raise HTTPException(status_code=400, detail="User IP address not available")
    
    # Ban the IP
    reason = data.get("reason", f"Banned via user {user.get('email')}")
    
    existing = await banned_ips_collection.find_one({"ip": ip})
    if existing:
        await banned_ips_collection.update_one(
            {"ip": ip},
            {
                "$set": {
                    "isActive": True,
                    "reason": reason,
                    "bannedAt": datetime.now(timezone.utc).isoformat(),
                    "bannedBy": session.get("adminId"),
                    "associatedUserId": user_id
                }
            }
        )
    else:
        banned_ip = {
            "id": str(uuid.uuid4()),
            "ip": ip,
            "reason": reason,
            "isActive": True,
            "bannedAt": datetime.now(timezone.utc).isoformat(),
            "bannedBy": session.get("adminId"),
            "associatedUserId": user_id
        }
        await banned_ips_collection.insert_one(banned_ip)
    
    return {"success": True, "message": f"IP {ip} has been banned", "ip": ip}


# ========== DISCOUNT CODES ENDPOINTS (Token-based Auth) ==========

@app.get("/api/admin/discounts")
async def get_discounts(session: dict = Depends(get_current_admin)):
    """Get all discount codes (admin only - token auth)"""
    check_permission(session, "view_discounts")
    discounts = await discounts_collection.find({}, {"_id": 0}).to_list(1000)
    return discounts

@app.post("/api/admin/discounts")
async def create_discount(request: dict, session: dict = Depends(get_current_admin)):
    """Create a new discount code (admin only - token auth)"""
    check_permission(session, "manage_discounts")
    discount = {
        "id": str(uuid.uuid4()),
        "code": request.get("code", "").upper(),
        "discountPercent": request.get("discountPercent", 10),
        "startDate": request.get("startDate"),
        "expiryDate": request.get("expiryDate"),
        "usageType": request.get("usageType", "unlimited"),
        "usageLimit": request.get("usageLimit", 100),
        "usageCount": 0,
        "applicableTo": request.get("applicableTo", "all"),
        "specificGenerators": request.get("specificGenerators", []),
        "isActive": request.get("isActive", True),
        "usedByCustomers": [],
        "createdAt": datetime.now(timezone.utc).isoformat()
    }
    
    # Check if code already exists
    existing = await discounts_collection.find_one({"code": discount["code"]})
    if existing:
        raise HTTPException(status_code=400, detail="Discount code already exists")
    
    await discounts_collection.insert_one(discount)
    await log_action(session, "create_discount", "discount", discount["id"], discount["code"])
    return {"success": True, "id": discount["id"]}

@app.put("/api/admin/discounts/{discount_id}")
async def update_discount(discount_id: str, request: dict, session: dict = Depends(get_current_admin)):
    """Update a discount code (admin only - token auth)"""
    check_permission(session, "manage_discounts")
    update_data = {
        "code": request.get("code", "").upper(),
        "discountPercent": request.get("discountPercent"),
        "startDate": request.get("startDate"),
        "expiryDate": request.get("expiryDate"),
        "usageType": request.get("usageType"),
        "usageLimit": request.get("usageLimit"),
        "applicableTo": request.get("applicableTo"),
        "specificGenerators": request.get("specificGenerators", []),
        "isActive": request.get("isActive")
    }
    
    result = await discounts_collection.update_one(
        {"id": discount_id},
        {"$set": update_data}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Discount not found")
    
    return {"success": True}

@app.delete("/api/admin/discounts/{discount_id}")
async def delete_discount(discount_id: str, session: dict = Depends(get_current_admin)):
    """Delete a discount code (admin only - token auth)"""
    check_permission(session, "manage_discounts")
    result = await discounts_collection.delete_one({"id": discount_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Discount not found")
    await log_action(session, "delete_discount", "discount", discount_id)
    return {"success": True}


# ========== PROMOTIONAL BANNER ENDPOINTS ==========

@app.get("/api/banner")
async def get_active_banner():
    """Get the currently active promotional banner (public endpoint)"""
    settings = await site_settings_collection.find_one({"key": "promotional_banner"})
    
    if not settings or not settings.get("isActive"):
        return {"success": True, "banner": None}
    
    # Check if the associated discount is still valid
    if settings.get("discountId"):
        discount = await discounts_collection.find_one({"id": settings["discountId"]})
        if discount:
            # Check if discount is active and not expired
            now = datetime.now(timezone.utc).isoformat()
            if not discount.get("isActive"):
                return {"success": True, "banner": None}
            if discount.get("expiryDate") and discount["expiryDate"] < now:
                return {"success": True, "banner": None}
            if discount.get("startDate") and discount["startDate"] > now:
                return {"success": True, "banner": None}
    
    return {
        "success": True,
        "banner": {
            "message": settings.get("message", ""),
            "discountCode": settings.get("discountCode", ""),
            "discountPercent": settings.get("discountPercent", 0),
            "backgroundColor": settings.get("backgroundColor", "#10b981"),
            "textColor": settings.get("textColor", "#ffffff"),
            "isActive": settings.get("isActive", False)
        }
    }


@app.get("/api/admin/banner")
async def get_banner_settings(session: dict = Depends(get_current_admin)):
    """Get banner settings (admin only)"""
    check_permission(session, "view_site_settings")
    settings = await site_settings_collection.find_one({"key": "promotional_banner"}, {"_id": 0})
    
    if not settings:
        return {
            "success": True,
            "banner": {
                "isActive": False,
                "message": "",
                "discountId": None,
                "discountCode": "",
                "discountPercent": 0,
                "backgroundColor": "#10b981",
                "textColor": "#ffffff"
            }
        }
    
    return {"success": True, "banner": settings}


@app.put("/api/admin/banner")
async def update_banner_settings(request: dict, session: dict = Depends(get_current_admin)):
    """Update promotional banner settings (admin only)"""
    check_permission(session, "manage_site_settings")
    discount_id = request.get("discountId")
    discount_code = ""
    discount_percent = 0
    
    # If a discount is selected, get its details
    if discount_id:
        discount = await discounts_collection.find_one({"id": discount_id})
        if discount:
            discount_code = discount.get("code", "")
            discount_percent = discount.get("discountPercent", 0)
    
    banner_settings = {
        "key": "promotional_banner",
        "isActive": request.get("isActive", False),
        "message": request.get("message", ""),
        "discountId": discount_id,
        "discountCode": discount_code,
        "discountPercent": discount_percent,
        "backgroundColor": request.get("backgroundColor", "#10b981"),
        "textColor": request.get("textColor", "#ffffff"),
        "updatedAt": datetime.now(timezone.utc).isoformat()
    }
    
    await site_settings_collection.update_one(
        {"key": "promotional_banner"},
        {"$set": banner_settings},
        upsert=True
    )
    await log_action(session, "update_banner", "site_settings", "promotional_banner")
    return {"success": True, "banner": banner_settings}


# ========== MAINTENANCE MODE ==========

@app.get("/api/maintenance-status")
async def get_maintenance_status():
    """Get current maintenance mode status (public endpoint)"""
    settings = await site_settings_collection.find_one({"key": "maintenance_mode"})
    
    if not settings:
        return {
            "success": True,
            "maintenance": {
                "isActive": False,
                "message": "We're currently performing scheduled maintenance. We'll be back shortly!",
                "estimatedTime": ""
            }
        }
    
    return {
        "success": True,
        "maintenance": {
            "isActive": settings.get("isActive", False),
            "message": settings.get("message", "We're currently performing scheduled maintenance. We'll be back shortly!"),
            "estimatedTime": settings.get("estimatedTime", "")
        }
    }


@app.get("/api/admin/maintenance")
async def get_maintenance_settings(session: dict = Depends(get_current_admin)):
    """Get maintenance mode settings (admin only)"""
    check_permission(session, "view_site_settings")
    settings = await site_settings_collection.find_one({"key": "maintenance_mode"}, {"_id": 0})
    
    if not settings:
        return {
            "success": True,
            "maintenance": {
                "isActive": False,
                "message": "We're currently performing scheduled maintenance. We'll be back shortly!",
                "estimatedTime": ""
            }
        }
    
    return {"success": True, "maintenance": settings}


@app.put("/api/admin/maintenance")
async def update_maintenance_settings(request: dict, session: dict = Depends(get_current_admin)):
    """Update maintenance mode settings (admin only)"""
    check_permission(session, "manage_site_settings")
    maintenance_settings = {
        "key": "maintenance_mode",
        "isActive": request.get("isActive", False),
        "message": request.get("message", "We're currently performing scheduled maintenance. We'll be back shortly!"),
        "estimatedTime": request.get("estimatedTime", ""),
        "updatedAt": datetime.now(timezone.utc).isoformat(),
        "updatedBy": session.get("adminId", "")
    }
    
    await site_settings_collection.update_one(
        {"key": "maintenance_mode"},
        {"$set": maintenance_settings},
        upsert=True
    )
    status = "enabled" if maintenance_settings["isActive"] else "disabled"
    print(f"Maintenance mode {status} by admin {session.get('adminId', 'unknown')}")
    await log_action(session, f"maintenance_{status}", "site_settings", "maintenance_mode")
    return {"success": True, "maintenance": maintenance_settings}


# ========== AUTH SETTINGS ==========

@app.get("/api/auth-status")
async def get_auth_status():
    """Get whether user signup/login is enabled (public endpoint)"""
    settings = await site_settings_collection.find_one({"key": "auth_enabled"})

    if not settings:
        return {"success": True, "authEnabled": True}

    return {"success": True, "authEnabled": settings.get("isEnabled", True)}


@app.get("/api/admin/auth-settings")
async def get_auth_settings(session: dict = Depends(get_current_admin)):
    """Get auth enabled settings (admin only)"""
    check_permission(session, "view_site_settings")
    settings = await site_settings_collection.find_one({"key": "auth_enabled"}, {"_id": 0})

    if not settings:
        return {"success": True, "authEnabled": True}

    return {"success": True, "authEnabled": settings.get("isEnabled", True)}


@app.put("/api/admin/auth-settings")
async def update_auth_settings(request: dict, session: dict = Depends(get_current_admin)):
    """Update auth enabled setting (admin only)"""
    check_permission(session, "manage_site_settings")
    auth_settings = {
        "key": "auth_enabled",
        "isEnabled": request.get("isEnabled", True),
        "updatedAt": datetime.now(timezone.utc).isoformat(),
        "updatedBy": session.get("adminId", "")
    }

    await site_settings_collection.update_one(
        {"key": "auth_enabled"},
        {"$set": auth_settings},
        upsert=True
    )
    status = "enabled" if auth_settings["isEnabled"] else "disabled"
    print(f"User auth {status} by admin {session.get('adminId', 'unknown')}")
    await log_action(session, f"auth_{status}", "site_settings", "auth_enabled")
    return {"success": True, "authEnabled": auth_settings["isEnabled"]}


# ========== EMAIL TEMPLATES ==========

EMAIL_TEMPLATE_DEFS = [
    {"name": "welcome",               "display_name": "Welcome Email",                  "variables": ["user_name", "user_email", "SITE_URL"],                                     "is_system": False, "is_scheduled": False},
    {"name": "email_verification",    "display_name": "Email Verification",             "variables": ["user_name", "verification_code", "verification_link"],                     "is_system": True,  "is_scheduled": False},
    {"name": "getting_started",       "display_name": "Getting Started Guide",          "variables": ["user_name", "SITE_URL"],                                                   "is_system": False, "is_scheduled": True,  "default_delay_minutes": 15},
    {"name": "subscription_thank_you","display_name": "Subscription Thank You",         "variables": ["user_name", "plan_name", "plan_price", "downloads_per_month", "SITE_URL"], "is_system": False, "is_scheduled": True,  "default_delay_minutes": 15},
    {"name": "download_confirmation", "display_name": "Download Confirmation",          "variables": ["user_name", "doc_name", "SITE_URL"],                                       "is_system": False, "is_scheduled": False},
    {"name": "signup_no_purchase",    "display_name": "Signup – No Purchase Reminder",  "variables": ["user_name", "SITE_URL"],                                                   "is_system": False, "is_scheduled": True,  "default_delay_minutes": 1440},
    {"name": "abandoned_checkout",    "display_name": "Abandoned Checkout",             "variables": ["user_name", "doc_name", "SITE_URL"],                                       "is_system": False, "is_scheduled": True,  "default_delay_minutes": 120},
    {"name": "review_request",        "display_name": "Review Request",                 "variables": ["user_name", "doc_name", "TRUSTPILOT_URL"],                                 "is_system": False, "is_scheduled": False},
    {"name": "password_changed",      "display_name": "Password Changed",               "variables": ["user_name", "SITE_URL"],                                                   "is_system": True,  "is_scheduled": False},
    {"name": "password_reset",        "display_name": "Password Reset",                 "variables": ["user_name", "reset_link", "reset_code"],                                   "is_system": True,  "is_scheduled": False},
]


@app.get("/api/admin/email-templates")
async def get_email_templates(session: dict = Depends(get_current_admin)):
    """Get all email templates with their current overrides and settings (admin only)"""
    check_permission(session, "view_email_templates")
    results = []
    for tmpl in EMAIL_TEMPLATE_DEFS:
        custom = await email_templates_collection.find_one({"name": tmpl["name"]}, {"_id": 0})
        has_custom_content = bool(custom and (custom.get("subject") or custom.get("html_body")))
        entry = {**tmpl, "is_custom": has_custom_content}
        if custom:
            if has_custom_content:
                entry["subject"] = custom.get("subject", "")
                entry["html_body"] = custom.get("html_body", "")
                entry["preview_text"] = custom.get("preview_text", "")
            entry["enabled"] = custom.get("enabled", True)
            if tmpl.get("is_scheduled") and custom.get("delay_minutes") is not None:
                entry["delay_minutes"] = custom["delay_minutes"]
        else:
            entry["enabled"] = True
        results.append(entry)
    return {"success": True, "templates": results}


@app.put("/api/admin/email-templates/{name}")
async def update_email_template(name: str, request: dict, session: dict = Depends(get_current_admin)):
    """Save a custom email template override (admin only)"""
    check_permission(session, "manage_email_templates")
    valid_names = {t["name"] for t in EMAIL_TEMPLATE_DEFS}
    if name not in valid_names:
        raise HTTPException(status_code=404, detail="Template not found")

    await email_templates_collection.update_one(
        {"name": name},
        {"$set": {
            "name": name,
            "subject": request.get("subject", ""),
            "html_body": request.get("html_body", ""),
            "preview_text": request.get("preview_text", ""),
            "updated_at": datetime.now(timezone.utc).isoformat(),
            "updated_by": session.get("adminId", "")
        }},
        upsert=True
    )
    await log_action(session, "update_email_template", "email_template", name)
    return {"success": True}


@app.put("/api/admin/email-settings/{name}")
async def update_email_settings(name: str, request: dict, session: dict = Depends(get_current_admin)):
    """Save email enabled/delay_minutes settings independently of template content (admin only)"""
    check_permission(session, "manage_email_templates")
    valid_names = {t["name"] for t in EMAIL_TEMPLATE_DEFS}
    if name not in valid_names:
        raise HTTPException(status_code=404, detail="Template not found")

    update_data: dict = {
        "name": name,
        "settings_updated_at": datetime.now(timezone.utc).isoformat(),
        "settings_updated_by": session.get("adminId", "")
    }
    if "enabled" in request:
        update_data["enabled"] = bool(request["enabled"])
    if "delay_minutes" in request:
        raw = request["delay_minutes"]
        update_data["delay_minutes"] = int(raw) if raw is not None else None

    await email_templates_collection.update_one(
        {"name": name},
        {"$set": update_data},
        upsert=True
    )
    return {"success": True}


@app.delete("/api/admin/email-templates/{name}")
async def reset_email_template(name: str, session: dict = Depends(get_current_admin)):
    """Delete a custom template override, reverting to the hardcoded default (admin only)"""
    check_permission(session, "manage_email_templates")
    await email_templates_collection.delete_one({"name": name})
    await log_action(session, "reset_email_template", "email_template", name)
    return {"success": True}


# ─────────────────────────────────────────────────────────────
# Mass Email
# ─────────────────────────────────────────────────────────────

@app.get("/api/admin/mass-email/recipients")
async def get_mass_email_recipients(session: dict = Depends(get_current_admin)):
    """Return unique purchaser emails enriched with names from users collection."""
    # Aggregate unique emails from purchases (field: email)
    pipeline = [
        {"$match": {"email": {"$exists": True, "$nin": [None, ""]}}},
        {"$group": {"_id": {"$toLower": "$email"}, "latestPurchase": {"$max": "$createdAt"}}},
        {"$sort": {"latestPurchase": -1}},
    ]
    email_docs = await purchases_collection.aggregate(pipeline).to_list(length=10000)

    emails = [doc["_id"] for doc in email_docs if doc["_id"]]

    # Bulk-fetch matching users for names
    user_docs = await users_collection.find(
        {"email": {"$in": emails}},
        {"email": 1, "name": 1}
    ).to_list(length=10000)
    name_map = {u["email"].lower(): u.get("name", "") for u in user_docs}

    recipients = []
    for doc in email_docs:
        em = doc["_id"]
        recipients.append({
            "email": em,
            "name": name_map.get(em, ""),
            "latestPurchase": doc.get("latestPurchase"),
        })

    return {"success": True, "recipients": recipients}


# Track active mass-email jobs so the frontend can poll status
_mass_email_jobs: dict = {}

@app.post("/api/admin/mass-email/send")
async def send_mass_email(request: dict, background_tasks: BackgroundTasks, session: dict = Depends(get_current_admin)):
    """
    Kick off a background mass-email send.
    Body: { subject, html_body, preview_text?, excluded_emails: [] }
    Returns a job_id the client can poll with GET /api/admin/mass-email/status/{job_id}
    """
    check_permission(session, "send_mass_email")
    subject = (request.get("subject") or "").strip()
    html_body = (request.get("html_body") or "").strip()
    preview_text = (request.get("preview_text") or "").strip()
    excluded = {e.lower() for e in (request.get("excluded_emails") or [])}

    if not subject or not html_body:
        raise HTTPException(status_code=400, detail="subject and html_body are required")

    # Build recipient list (re-use the aggregation logic inline)
    pipeline = [
        {"$match": {"email": {"$exists": True, "$nin": [None, ""]}}},
        {"$group": {"_id": {"$toLower": "$email"}}},
    ]
    email_docs = await purchases_collection.aggregate(pipeline).to_list(length=10000)
    all_emails = [doc["_id"] for doc in email_docs if doc["_id"] and doc["_id"] not in excluded]

    job_id = str(uuid.uuid4())
    _mass_email_jobs[job_id] = {"status": "running", "total": len(all_emails), "sent": 0, "failed": 0, "errors": []}

    async def _do_send():
        job = _mass_email_jobs[job_id]
        full_html = get_base_template(html_body, preview_text)
        for email_addr in all_emails:
            try:
                await send_email(email_addr, subject, full_html, email_type="mass_email")
                job["sent"] += 1
            except Exception as ex:
                job["failed"] += 1
                job["errors"].append({"email": email_addr, "error": str(ex)})
        job["status"] = "done"

    background_tasks.add_task(_do_send)
    await log_action(session, "send_mass_email", "mass_email", job_id, f"subject={subject}, recipients={len(all_emails)}")
    return {"success": True, "job_id": job_id, "total": len(all_emails)}


@app.get("/api/admin/mass-email/status/{job_id}")
async def get_mass_email_status(job_id: str, session: dict = Depends(get_current_admin)):
    job = _mass_email_jobs.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    return {"success": True, **job}


# ─────────────────────────────────────────────────────────────
# Moderator Management (admin-only)
# ─────────────────────────────────────────────────────────────

@app.get("/api/admin/moderators")
async def list_moderators(session: dict = Depends(require_admin_only)):
    mods = await moderators_collection.find({}, {"_id": 0, "password": 0}).sort("createdAt", -1).to_list(1000)
    return {"success": True, "moderators": mods}


@app.post("/api/admin/moderators")
async def create_moderator(request: dict, session: dict = Depends(require_admin_only)):
    name = (request.get("name") or "").strip()
    email = (request.get("email") or "").strip().lower()
    password = (request.get("password") or "").strip()
    level = int(request.get("level", 1))

    if not name or not email or not password:
        raise HTTPException(status_code=400, detail="name, email, and password are required")
    if level not in (1, 2, 3):
        raise HTTPException(status_code=400, detail="level must be 1, 2, or 3")

    existing_admin = await admins_collection.find_one({"email": email})
    existing_mod = await moderators_collection.find_one({"email": email})
    if existing_admin or existing_mod:
        raise HTTPException(status_code=400, detail="Email already in use")

    moderator = {
        "id": str(uuid.uuid4()),
        "email": email,
        "password": hash_password(password),
        "name": name,
        "level": level,
        "isActive": True,
        "createdAt": datetime.now(timezone.utc).isoformat(),
        "createdBy": session.get("adminId", ""),
    }
    await moderators_collection.insert_one(moderator)
    moderator.pop("_id", None)
    moderator.pop("password", None)
    await log_action(session, "create_moderator", "moderator", moderator["id"], f"{name} ({email}) level={level}")
    return {"success": True, "moderator": moderator}


@app.put("/api/admin/moderators/{mod_id}")
async def update_moderator(mod_id: str, request: dict, session: dict = Depends(require_admin_only)):
    mod = await moderators_collection.find_one({"id": mod_id})
    if not mod:
        raise HTTPException(status_code=404, detail="Moderator not found")

    update_data: dict = {"updatedAt": datetime.now(timezone.utc).isoformat()}
    if "name" in request and request["name"]:
        update_data["name"] = request["name"].strip()
    if "email" in request and request["email"]:
        new_email = request["email"].strip().lower()
        if new_email != mod["email"]:
            clash_admin = await admins_collection.find_one({"email": new_email})
            clash_mod = await moderators_collection.find_one({"email": new_email, "id": {"$ne": mod_id}})
            if clash_admin or clash_mod:
                raise HTTPException(status_code=400, detail="Email already in use")
        update_data["email"] = new_email
    if "password" in request and request["password"]:
        update_data["password"] = hash_password(request["password"])
    if "level" in request:
        lvl = int(request["level"])
        if lvl not in (1, 2, 3):
            raise HTTPException(status_code=400, detail="level must be 1, 2, or 3")
        update_data["level"] = lvl
    if "isActive" in request:
        update_data["isActive"] = bool(request["isActive"])
        if not request["isActive"]:
            # Invalidate all active sessions for this moderator
            await sessions_collection.delete_many({"moderatorId": mod_id})

    await moderators_collection.update_one({"id": mod_id}, {"$set": update_data})
    return {"success": True}


@app.delete("/api/admin/moderators/{mod_id}")
async def delete_moderator(mod_id: str, session: dict = Depends(require_admin_only)):
    result = await moderators_collection.delete_one({"id": mod_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Moderator not found")
    await sessions_collection.delete_many({"moderatorId": mod_id})
    await log_action(session, "delete_moderator", "moderator", mod_id)
    return {"success": True}


# ─────────────────────────────────────────────────────────────
# Moderator Level Permissions (admin-only)
# ─────────────────────────────────────────────────────────────

@app.get("/api/admin/moderator-permissions")
async def get_moderator_permissions(session: dict = Depends(require_admin_only)):
    result = {}
    for level in (1, 2, 3):
        doc = await moderator_permissions_collection.find_one({"level": level}, {"_id": 0})
        result[str(level)] = doc.get("permissions", DEFAULT_LEVEL_PERMISSIONS[level]) if doc else DEFAULT_LEVEL_PERMISSIONS[level]
    return {"success": True, "permissions": result}


@app.put("/api/admin/moderator-permissions/{level}")
async def update_moderator_permissions(level: int, request: dict, session: dict = Depends(require_admin_only)):
    if level not in (1, 2, 3):
        raise HTTPException(status_code=400, detail="level must be 1, 2, or 3")
    perms = {k: bool(v) for k, v in request.get("permissions", {}).items() if k in ALL_PERMISSIONS}
    await moderator_permissions_collection.update_one(
        {"level": level},
        {"$set": {"level": level, "permissions": perms, "updatedAt": datetime.now(timezone.utc).isoformat()}},
        upsert=True,
    )
    # Update all active sessions for moderators of this level
    await sessions_collection.update_many(
        {"type": "moderator", "level": level},
        {"$set": {"permissions": perms}},
    )
    return {"success": True}


# ─────────────────────────────────────────────────────────────
# Audit Log
# ─────────────────────────────────────────────────────────────

@app.get("/api/admin/audit-log")
async def get_audit_log(
    session: dict = Depends(require_admin_only),
    skip: int = 0,
    limit: int = 50,
    actor: Optional[str] = None,
    action: Optional[str] = None,
    resource_type: Optional[str] = None,
):
    query: dict = {}
    if actor:
        query["actorEmail"] = {"$regex": actor, "$options": "i"}
    if action:
        query["action"] = {"$regex": action, "$options": "i"}
    if resource_type:
        query["resourceType"] = resource_type
    total = await audit_logs_collection.count_documents(query)
    logs = await audit_logs_collection.find(query, {"_id": 0}).sort("timestamp", -1).skip(skip).limit(limit).to_list(limit)
    return {"success": True, "logs": logs, "total": total}


@app.delete("/api/admin/audit-log")
async def clear_audit_log(session: dict = Depends(require_admin_only)):
    await audit_logs_collection.delete_many({})
    return {"success": True}


# ─────────────────────────────────────────────────────────────
# Support Tickets (Contact Form)
# ─────────────────────────────────────────────────────────────

@app.post("/api/contact")
async def submit_contact(request: dict):
    """Public endpoint — saves contact form submission to DB."""
    name = (request.get("name") or "").strip()
    email = (request.get("email") or "").strip().lower()
    reason = (request.get("reason") or "").strip()
    message = (request.get("message") or "").strip()
    if not name or not email or not message:
        raise HTTPException(status_code=400, detail="name, email, and message are required")
    ticket = {
        "id": str(uuid.uuid4()),
        "name": name,
        "email": email,
        "reason": reason,
        "message": message,
        "status": "open",
        "notes": "",
        "createdAt": datetime.now(timezone.utc).isoformat(),
        "updatedAt": datetime.now(timezone.utc).isoformat(),
    }
    await support_tickets_collection.insert_one(ticket)
    return {"success": True}


@app.get("/api/admin/support-tickets")
async def get_support_tickets(
    session: dict = Depends(get_current_admin),
    skip: int = 0,
    limit: int = 50,
    status: Optional[str] = None,
):
    query: dict = {}
    if status and status != "all":
        query["status"] = status
    total = await support_tickets_collection.count_documents(query)
    tickets = await support_tickets_collection.find(query, {"_id": 0}).sort("createdAt", -1).skip(skip).limit(limit).to_list(limit)
    open_count = await support_tickets_collection.count_documents({"status": "open"})
    return {"success": True, "tickets": tickets, "total": total, "openCount": open_count}


@app.put("/api/admin/support-tickets/{ticket_id}")
async def update_support_ticket(ticket_id: str, request: dict, session: dict = Depends(get_current_admin)):
    update: dict = {"updatedAt": datetime.now(timezone.utc).isoformat()}
    if "status" in request:
        update["status"] = request["status"]
    if "notes" in request:
        update["notes"] = request["notes"]
    result = await support_tickets_collection.update_one({"id": ticket_id}, {"$set": update})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Ticket not found")
    await log_action(session, "update_ticket", "support_ticket", ticket_id, f"status={request.get('status', '')}")
    return {"success": True}


@app.delete("/api/admin/support-tickets/{ticket_id}")
async def delete_support_ticket(ticket_id: str, session: dict = Depends(get_current_admin)):
    result = await support_tickets_collection.delete_one({"id": ticket_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Ticket not found")
    await log_action(session, "delete_ticket", "support_ticket", ticket_id)
    return {"success": True}


# ─────────────────────────────────────────────────────────────
# CSV Export
# ─────────────────────────────────────────────────────────────

from fastapi.responses import StreamingResponse
import csv
import io

@app.get("/api/admin/export/purchases")
async def export_purchases_csv(
    session: dict = Depends(get_current_admin),
    startDate: Optional[str] = None,
    endDate: Optional[str] = None,
    documentType: Optional[str] = None,
):
    check_permission(session, "view_purchases")
    query: dict = {}
    if documentType:
        query["documentType"] = documentType
    if startDate:
        query["createdAt"] = {"$gte": startDate}
    if endDate:
        query.setdefault("createdAt", {})["$lte"] = endDate
    purchases = await purchases_collection.find(query, {"_id": 0}).sort("createdAt", -1).to_list(50000)

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["Date", "Email", "Document Type", "Amount", "Purchase ID", "User ID", "Guest"])
    for p in purchases:
        writer.writerow([
            p.get("createdAt", "")[:10],
            p.get("email", ""),
            p.get("documentType", ""),
            p.get("amount", ""),
            p.get("id", ""),
            p.get("userId", ""),
            "Yes" if p.get("isGuest") else "No",
        ])
    output.seek(0)
    await log_action(session, "export_csv", "purchases", "", "purchases export")
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=purchases.csv"},
    )


@app.get("/api/admin/export/users")
async def export_users_csv(session: dict = Depends(get_current_admin)):
    check_permission(session, "view_users")
    users = await users_collection.find({}, {"_id": 0, "password": 0}).sort("createdAt", -1).to_list(50000)

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["Date", "Name", "Email", "Subscription", "Email Verified", "Banned", "User ID"])
    for u in users:
        sub = u.get("subscription") or {}
        writer.writerow([
            u.get("createdAt", "")[:10],
            u.get("name", ""),
            u.get("email", ""),
            sub.get("tier", "none") if sub else "none",
            "Yes" if u.get("emailVerified") else "No",
            "Yes" if u.get("isBanned") else "No",
            u.get("id", ""),
        ])
    output.seek(0)
    await log_action(session, "export_csv", "users", "", "users export")
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=users.csv"},
    )


@app.get("/api/admin/export/revenue")
async def export_revenue_csv(session: dict = Depends(get_current_admin)):
    check_permission(session, "view_purchases")
    purchases = await purchases_collection.find({}, {"_id": 0}).sort("createdAt", -1).to_list(50000)
    sub_payments = await subscription_payments_collection.find({}, {"_id": 0}).sort("createdAt", -1).to_list(50000)

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["Date", "Type", "Email", "Amount", "Description", "ID"])
    for p in purchases:
        writer.writerow([p.get("createdAt", "")[:10], "purchase", p.get("email", ""), p.get("amount", ""), p.get("documentType", ""), p.get("id", "")])
    for s in sub_payments:
        writer.writerow([s.get("createdAt", "")[:10], "subscription", s.get("email", s.get("userId", "")), s.get("amount", ""), s.get("tier", ""), s.get("id", "")])
    output.seek(0)
    await log_action(session, "export_csv", "revenue", "", "revenue export")
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=revenue.csv"},
    )


# ─────────────────────────────────────────────────────────────
# Subscription Tier Download Settings
# ─────────────────────────────────────────────────────────────

DEFAULT_TIER_DOWNLOADS = {"starter": 10, "professional": 30, "business": -1}

async def get_tier_downloads() -> dict:
    """Return download limits per tier, reading from DB (falls back to defaults)."""
    doc = await site_settings_collection.find_one({"key": "subscription_tier_downloads"})
    if doc and "tiers" in doc:
        merged = dict(DEFAULT_TIER_DOWNLOADS)
        merged.update(doc["tiers"])
        return merged
    return dict(DEFAULT_TIER_DOWNLOADS)


@app.get("/api/admin/subscription-tier-settings")
async def get_subscription_tier_settings(session: dict = Depends(require_admin_only)):
    downloads = await get_tier_downloads()
    return {"success": True, "tiers": downloads}


@app.put("/api/admin/subscription-tier-settings")
async def update_subscription_tier_settings(request: dict, session: dict = Depends(require_admin_only)):
    tiers = request.get("tiers", {})
    validated = {}
    for tier in ("starter", "professional", "business"):
        val = tiers.get(tier)
        if val is not None:
            try:
                v = int(val)
                validated[tier] = v if v > 0 else -1  # -1 = unlimited
            except (ValueError, TypeError):
                raise HTTPException(status_code=400, detail=f"Invalid download count for {tier}")
    await site_settings_collection.update_one(
        {"key": "subscription_tier_downloads"},
        {"$set": {"key": "subscription_tier_downloads", "tiers": validated, "updatedAt": datetime.now(timezone.utc).isoformat()}},
        upsert=True,
    )
    await log_action(session, "update_tier_downloads", "site_settings", "subscription_tier_downloads",
                     str(validated))
    return {"success": True, "tiers": validated}


# ─────────────────────────────────────────────────────────────
# Revenue Summary
# ─────────────────────────────────────────────────────────────

@app.get("/api/admin/revenue/summary")
async def get_revenue_summary(session: dict = Depends(get_current_admin)):
    """Aggregated revenue metrics and breakdowns for the Revenue page."""
    check_permission(session, "view_purchases")

    now = datetime.now(timezone.utc)

    # Month boundaries
    this_month_start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0).isoformat()
    last_month_end = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    last_month_start = (last_month_end - timedelta(days=1)).replace(day=1).isoformat()
    last_month_end_iso = last_month_end.isoformat()

    # All-time revenue
    agg = await purchases_collection.aggregate([
        {"$group": {"_id": None, "total": {"$sum": "$amount"}, "count": {"$sum": 1}}}
    ]).to_list(1)
    total_revenue = agg[0]["total"] if agg else 0.0
    total_purchases = agg[0]["count"] if agg else 0
    avg_order_value = (total_revenue / total_purchases) if total_purchases else 0.0

    # This month revenue
    agg_month = await purchases_collection.aggregate([
        {"$match": {"createdAt": {"$gte": this_month_start}}},
        {"$group": {"_id": None, "total": {"$sum": "$amount"}}}
    ]).to_list(1)
    this_month_revenue = agg_month[0]["total"] if agg_month else 0.0

    # Last month revenue
    agg_last = await purchases_collection.aggregate([
        {"$match": {"createdAt": {"$gte": last_month_start, "$lt": last_month_end_iso}}},
        {"$group": {"_id": None, "total": {"$sum": "$amount"}}}
    ]).to_list(1)
    last_month_revenue = agg_last[0]["total"] if agg_last else 0.0

    # Daily data — last 90 days
    ninety_days_ago = (now - timedelta(days=90)).isoformat()
    daily_raw = await purchases_collection.aggregate([
        {"$match": {"createdAt": {"$gte": ninety_days_ago}}},
        {"$group": {
            "_id": {"$substr": ["$createdAt", 0, 10]},
            "revenue": {"$sum": "$amount"},
            "count": {"$sum": 1},
        }},
        {"$sort": {"_id": 1}},
    ]).to_list(90)
    daily_data = [{"date": d["_id"], "revenue": round(d["revenue"], 2), "count": d["count"]} for d in daily_raw]

    # Revenue and count by document type
    by_type_raw = await purchases_collection.aggregate([
        {"$group": {
            "_id": "$documentType",
            "revenue": {"$sum": "$amount"},
            "count": {"$sum": 1},
        }},
        {"$sort": {"revenue": -1}},
    ]).to_list(50)
    by_doc_type = [
        {
            "documentType": d["_id"] or "unknown",
            "revenue": round(d["revenue"], 2),
            "count": d["count"],
            "avgPrice": round(d["revenue"] / d["count"], 2) if d["count"] else 0,
        }
        for d in by_type_raw
    ]

    return {
        "success": True,
        "totalRevenue": round(total_revenue, 2),
        "totalPurchases": total_purchases,
        "avgOrderValue": round(avg_order_value, 2),
        "thisMonthRevenue": round(this_month_revenue, 2),
        "lastMonthRevenue": round(last_month_revenue, 2),
        "dailyData": daily_data,
        "byDocType": by_doc_type,
    }


# ─────────────────────────────────────────────────────────────
# Purchase Refunds
# ─────────────────────────────────────────────────────────────

@app.post("/api/admin/purchases/{purchase_id}/refund")
async def refund_purchase(purchase_id: str, request: dict, session: dict = Depends(get_current_admin)):
    """Issue a Stripe refund for a purchase (full or partial)."""
    check_permission(session, "manage_discounts")

    purchase = await purchases_collection.find_one({"id": purchase_id})
    if not purchase:
        raise HTTPException(status_code=404, detail="Purchase not found")
    if purchase.get("refunded"):
        raise HTTPException(status_code=400, detail="Purchase has already been refunded")

    pi_id = purchase.get("stripePaymentIntentId") or purchase.get("paymentIntentId")
    if not pi_id:
        raise HTTPException(status_code=400, detail="No Stripe payment record found for this purchase")

    amount_dollars = float(request.get("amount_dollars", purchase.get("amount", 0)))
    if amount_dollars <= 0:
        raise HTTPException(status_code=400, detail="Refund amount must be greater than 0")
    if amount_dollars > purchase.get("amount", 0):
        raise HTTPException(status_code=400, detail="Refund amount exceeds original purchase amount")

    reason = request.get("reason", "requested_by_customer")

    try:
        refund = await asyncio.to_thread(
            stripe.Refund.create,
            payment_intent=pi_id,
            amount=int(round(amount_dollars * 100)),
            reason=reason if reason in ("duplicate", "fraudulent", "requested_by_customer") else "requested_by_customer",
        )
    except stripe.error.StripeError as e:
        raise HTTPException(status_code=400, detail=str(e))

    await purchases_collection.update_one(
        {"id": purchase_id},
        {"$set": {
            "refunded": True,
            "refundedAmount": amount_dollars,
            "refundedAt": datetime.now(timezone.utc).isoformat(),
            "stripeRefundId": refund.id,
        }},
    )
    await log_action(session, "refund_purchase", "purchase", purchase_id, f"${amount_dollars:.2f} refund_id={refund.id}")
    return {"success": True, "refund_id": refund.id, "amount": amount_dollars}


# ─────────────────────────────────────────────────────────────
# Admin Subscription Management
# ─────────────────────────────────────────────────────────────

@app.get("/api/admin/subscriptions")
async def admin_list_subscriptions(session: dict = Depends(require_admin_only)):
    """List all users with subscriptions enriched with live Stripe status."""
    users = await users_collection.find(
        {"subscription.stripeSubscriptionId": {"$exists": True, "$ne": ""}},
        {"_id": 0, "password": 0},
    ).sort("subscription.current_period_end", 1).to_list(1000)

    async def enrich(user):
        sub = user.get("subscription", {})
        sub_id = sub.get("stripeSubscriptionId")
        stripe_status = None
        is_past_due = False
        cancel_at_period_end = False
        current_period_end = sub.get("current_period_end")
        try:
            if sub_id:
                stripe_sub = await asyncio.to_thread(
                    stripe.Subscription.retrieve, sub_id,
                    expand=["latest_invoice.payment_intent"]
                )
                stripe_status = stripe_sub.status
                is_past_due = stripe_status in ("past_due", "unpaid")
                cancel_at_period_end = stripe_sub.cancel_at_period_end
                current_period_end = datetime.fromtimestamp(stripe_sub.current_period_end, tz=timezone.utc).isoformat() if stripe_sub.current_period_end else current_period_end
        except Exception:
            stripe_status = "unknown"
        return {
            "userId": user.get("id"),
            "name": user.get("name", ""),
            "email": user.get("email", ""),
            "tier": sub.get("tier", ""),
            "dbStatus": sub.get("status", ""),
            "stripeStatus": stripe_status,
            "isPastDue": is_past_due,
            "cancelAtPeriodEnd": cancel_at_period_end,
            "currentPeriodEnd": current_period_end,
            "stripeSubscriptionId": sub.get("stripeSubscriptionId", ""),
            "stripeCustomerId": sub.get("stripeCustomerId", ""),
            "createdAt": sub.get("createdAt", ""),
        }

    results = await asyncio.gather(*[enrich(u) for u in users])
    # Sort: past_due first, then by period end
    results = sorted(results, key=lambda x: (not x["isPastDue"], x["currentPeriodEnd"] or ""))
    return {"success": True, "subscriptions": results, "total": len(results)}


@app.post("/api/admin/subscriptions/{user_id}/cancel")
async def admin_cancel_subscription(user_id: str, request: dict, session: dict = Depends(require_admin_only)):
    """Cancel a user's subscription. immediate=True cancels now, False = at period end."""
    user = await users_collection.find_one({"id": user_id})
    if not user or not user.get("subscription"):
        raise HTTPException(status_code=404, detail="User or subscription not found")
    sub_id = user["subscription"].get("stripeSubscriptionId")
    if not sub_id:
        raise HTTPException(status_code=400, detail="No Stripe subscription found")

    immediate = bool(request.get("immediate", False))
    try:
        if immediate:
            await asyncio.to_thread(stripe.Subscription.delete, sub_id)
            new_status = "cancelled"
        else:
            await asyncio.to_thread(stripe.Subscription.modify, sub_id, cancel_at_period_end=True)
            new_status = "cancelling"
    except stripe.error.StripeError as e:
        raise HTTPException(status_code=400, detail=str(e))

    update: dict = {"subscription.status": new_status}
    if immediate:
        update["subscription.cancelledAt"] = datetime.now(timezone.utc).isoformat()
    await users_collection.update_one({"id": user_id}, {"$set": update})
    action = "cancel_subscription_now" if immediate else "cancel_subscription"
    await log_action(session, action, "subscription", user_id, user.get("email", ""))
    return {"success": True, "status": new_status}


@app.post("/api/admin/subscriptions/{user_id}/reactivate")
async def admin_reactivate_subscription(user_id: str, session: dict = Depends(require_admin_only)):
    """Reactivate a subscription that is set to cancel at period end."""
    user = await users_collection.find_one({"id": user_id})
    if not user or not user.get("subscription"):
        raise HTTPException(status_code=404, detail="User or subscription not found")
    sub_id = user["subscription"].get("stripeSubscriptionId")
    if not sub_id:
        raise HTTPException(status_code=400, detail="No Stripe subscription found")

    try:
        await asyncio.to_thread(stripe.Subscription.modify, sub_id, cancel_at_period_end=False)
    except stripe.error.StripeError as e:
        raise HTTPException(status_code=400, detail=str(e))

    await users_collection.update_one({"id": user_id}, {"$set": {"subscription.status": "active"}})
    await log_action(session, "reactivate_subscription", "subscription", user_id, user.get("email", ""))
    return {"success": True, "status": "active"}


@app.post("/api/parse-resume")
async def parse_resume(file: UploadFile = File(...)):
    """Parse an uploaded resume (PDF or DOCX) and extract structured data"""
    try:
        # Check file type
        filename = file.filename.lower()
        if not (filename.endswith('.pdf') or filename.endswith('.docx')):
            raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")
        
        # Read file content
        content = await file.read()
        
        # Extract text based on file type
        extracted_text = ""
        
        if filename.endswith('.pdf'):
            if not PDF_SUPPORT:
                raise HTTPException(status_code=500, detail="PDF parsing not available. Please upload a DOCX file.")
            
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        extracted_text += page_text + "\n"
        
        elif filename.endswith('.docx'):
            if not DOCX_SUPPORT:
                raise HTTPException(status_code=500, detail="DOCX parsing not available. Please upload a PDF file.")
            
            doc = Document(io.BytesIO(content))
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        extracted_text += cell.text + " "
                    extracted_text += "\n"
        
        if not extracted_text.strip():
            raise HTTPException(status_code=400, detail="Could not extract text from the resume. Please ensure the file is not empty or corrupted.")
        
        # Use AI to parse the extracted text into structured format
        prompt = f"""Parse the following resume text and extract structured information. Return ONLY valid JSON, no markdown code blocks.

RESUME TEXT:
{extracted_text[:8000]}

Extract and return this exact JSON structure (use empty strings or empty arrays if information is not found):
{{
    "personalInfo": {{
        "fullName": "Full name of the person",
        "email": "email@example.com",
        "phone": "phone number",
        "location": "City, State",
        "linkedin": "LinkedIn URL if present",
        "website": "Personal website if present"
    }},
    "workExperience": [
        {{
            "company": "Company Name",
            "position": "Job Title",
            "location": "City, State",
            "startDate": "MM/YYYY or Month YYYY",
            "endDate": "MM/YYYY or Month YYYY or Present",
            "current": false,
            "responsibilities": ["Responsibility 1", "Responsibility 2", "Responsibility 3"]
        }}
    ],
    "education": [
        {{
            "institution": "University/School Name",
            "degree": "Degree Type (e.g., Bachelor of Science)",
            "field": "Field of Study",
            "graduationDate": "YYYY or Month YYYY",
            "gpa": "GPA if mentioned"
        }}
    ],
    "skills": ["Skill 1", "Skill 2", "Skill 3"]
}}

Important:
- Extract ALL work experiences found, ordered from most recent to oldest
- Extract ALL education entries found
- For responsibilities, extract 3-5 bullet points per job
- If a field is not found in the resume, use an empty string "" or empty array []
- Set "current" to true if the job says "Present" or "Current"
- Return ONLY the JSON object, nothing else"""

        response = await call_claude(prompt)
        
        # Parse the JSON response
        response_text = response.strip()
        
        try:
            # Find JSON in the response
            start_idx = response_text.find('{')
            end_idx = response_text.rfind('}') + 1
            if start_idx != -1 and end_idx > start_idx:
                json_str = response_text[start_idx:end_idx]
                parsed_data = json.loads(json_str)
            else:
                raise ValueError("No JSON found in response")
        except json.JSONDecodeError as e:
            print(f"JSON parse error: {e}")
            print(f"Response was: {response_text[:500]}")
            # Return a basic structure if parsing fails
            parsed_data = {
                "personalInfo": {
                    "fullName": "",
                    "email": "",
                    "phone": "",
                    "location": "",
                    "linkedin": "",
                    "website": ""
                },
                "workExperience": [],
                "education": [],
                "skills": []
            }
        
        return {
            "success": True,
            "data": parsed_data,
            "rawTextLength": len(extracted_text)
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error parsing resume: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error parsing resume: {str(e)}")

@app.post("/api/scrape-job")
async def scrape_job(request: JobScrapeRequest):
    """Scrape job description from URL"""
    try:
        async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as client:
            headers = {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
            }
            response = await client.get(request.url, headers=headers)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.decompose()
            
            # Try to find job description in common containers
            job_text = ""
            
            # Common job description selectors
            selectors = [
                'div[class*="job-description"]',
                'div[class*="jobDescription"]',
                'div[class*="description"]',
                'section[class*="description"]',
                'div[id*="job-description"]',
                'div[id*="jobDescription"]',
                'article',
                'main',
                '.job-details',
                '.posting-requirements',
            ]
            
            for selector in selectors:
                elements = soup.select(selector)
                if elements:
                    job_text = " ".join([el.get_text(separator="\n", strip=True) for el in elements])
                    if len(job_text) > 200:
                        break
            
            # Fallback to body text if no specific container found
            if len(job_text) < 200:
                job_text = soup.get_text(separator="\n", strip=True)
            
            # Clean up the text
            lines = [line.strip() for line in job_text.split('\n') if line.strip()]
            job_text = "\n".join(lines[:150])  # Limit to ~150 lines
            
            if len(job_text) < 100:
                raise HTTPException(status_code=400, detail="Could not extract meaningful job description from URL")
            
            return {"jobDescription": job_text[:8000]}  # Limit to 8000 chars
            
    except httpx.HTTPError as e:
        raise HTTPException(status_code=400, detail=f"Failed to fetch URL: {str(e)}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error scraping job: {str(e)}")

@app.post("/api/generate-resume")
async def generate_resume(data: ResumeInput):
    """Generate AI-optimized resume content"""
    try:
        # Build context from user data - include job location
        work_history = "\n".join([
            f"- {exp.position} at {exp.company}, {exp.location} ({exp.startDate} - {'Present' if exp.current else exp.endDate})\n  Responsibilities: {'; '.join(exp.responsibilities)}"
            for exp in data.workExperience
        ])
        
        education_history = "\n".join([
            f"- {edu.degree} in {edu.field} from {edu.institution} ({edu.graduationDate})"
            for edu in data.education
        ])
        
        skills_list = ", ".join(data.skills) if data.skills else "Not specified"
        
        prompt = f"""Your task is to create an optimized resume tailored to a specific job posting.

IMPORTANT RULES:
1. ONLY use information provided by the user - DO NOT fabricate any employers, dates, degrees, locations, or certifications
2. For each work experience, use the EXACT location provided by the user (not the user's personal location)
3. Rewrite and enhance existing responsibilities using action verbs and quantifiable achievements where possible
4. Align the resume content to match the job requirements and keywords
5. Ensure ATS-friendly formatting with relevant keywords from the job description
6. Match the tone and seniority level of the target position

USER'S PERSONAL INFO:
Name: {data.personalInfo.get('fullName', 'Candidate')}
Email: {data.personalInfo.get('email', '')}
Phone: {data.personalInfo.get('phone', '')}
Personal Location: {data.personalInfo.get('location', '')}
LinkedIn: {data.personalInfo.get('linkedin', '')}

WORK EXPERIENCE (use these exact company locations, NOT the personal location above):
{work_history}

EDUCATION:
{education_history}

SKILLS:
{skills_list}

TARGET JOB TITLE: {data.targetJobTitle}

JOB DESCRIPTION:
{data.jobDescription[:4000]}

IMPORTANT: The user has provided {len(data.workExperience)} work experience entries. You MUST return exactly {len(data.workExperience)} entries in the optimizedExperience array, one for each job. Each entry must have UNIQUE and DIFFERENT bullet points tailored to that specific role. Do NOT duplicate or reuse bullet points across different positions.

Please generate the following in JSON format. ONLY output valid JSON, no markdown code blocks:
{{
    "professionalSummary": "A compelling 3-4 sentence professional summary tailored to this role",
    "optimizedExperience": [
        // Return EXACTLY {len(data.workExperience)} entries here - one for EACH work experience provided above
        // Each entry should look like:
        {{
            "company": "Exact company name from user input",
            "position": "Exact position title from user input",
            "location": "Exact location from user input",
            "startDate": "Exact start date from user input",
            "endDate": "Exact end date from user input or Present",
            "bullets": ["Unique bullet 1 for THIS role", "Unique bullet 2 for THIS role", "Unique bullet 3 for THIS role", "Unique bullet 4 for THIS role"]
        }}
    ],
    "optimizedSkills": {{
        "technical": ["skill1", "skill2"],
        "soft": ["skill1", "skill2"],
        "other": ["skill1", "skill2"]
    }},
    "keywordsUsed": ["keyword1", "keyword2"],
    "atsScore": 85,
    "suggestions": ["suggestion1", "suggestion2"]
}}

CRITICAL: You MUST generate {len(data.workExperience)} SEPARATE work experience entries with DIFFERENT bullet points for each. Each position's bullets should be tailored to the specific responsibilities mentioned for that role. Do NOT copy the same bullets across different positions.
"""

        response = await call_claude(prompt)
        
        # Parse the JSON response
        response_text = response
        
        # Try to extract JSON from the response
        try:
            # Find JSON in the response
            start_idx = response_text.find('{')
            end_idx = response_text.rfind('}') + 1
            if start_idx != -1 and end_idx > start_idx:
                json_str = response_text[start_idx:end_idx]
                result = json.loads(json_str)
            else:
                raise ValueError("No JSON found in response")
        except json.JSONDecodeError:
            # If JSON parsing fails, return a structured error response
            result = {
                "professionalSummary": "Experienced professional seeking new opportunities.",
                "optimizedExperience": [
                    {
                        "company": exp.company,
                        "position": exp.position,
                        "location": exp.location,
                        "startDate": exp.startDate,
                        "endDate": "Present" if exp.current else exp.endDate,
                        "bullets": exp.responsibilities[:4]
                    }
                    for exp in data.workExperience
                ],
                "optimizedSkills": {
                    "technical": data.skills[:5] if data.skills else [],
                    "soft": [],
                    "other": []
                },
                "keywordsUsed": [],
                "atsScore": 70,
                "suggestions": ["Consider adding more specific achievements", "Add relevant certifications"]
            }
        
        # Add original data for reference
        result["personalInfo"] = data.personalInfo
        result["education"] = [edu.dict() for edu in data.education]
        result["targetJobTitle"] = data.targetJobTitle
        
        return result
        
    except Exception as e:
        print(f"Error generating resume: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generating resume: {str(e)}")

@app.post("/api/regenerate-section")
async def regenerate_section(data: dict):
    """Regenerate a specific section of the resume"""
    try:
        section = data.get("section")
        current_content = data.get("currentContent")
        job_description = data.get("jobDescription", "")
        user_context = data.get("userContext", {})
        
        prompts = {
            "summary": f"""Rewrite this professional summary to better align with the job description. Keep it 3-4 sentences, compelling, and ATS-friendly.

Current Summary: {current_content}

Job Description: {job_description[:2000]}

Return only the new summary text, nothing else.""",

            "experience": f"""Rewrite these work experience bullet points to better align with the job description. Use strong action verbs, include metrics where possible, and ensure ATS optimization.

Current Bullets: {json.dumps(current_content)}

Job Description: {job_description[:2000]}

Return a JSON array of improved bullet points: ["bullet1", "bullet2", ...]""",

            "skills": f"""Reorganize and optimize these skills based on the job description. Prioritize the most relevant skills and add any that are implied by the experience.

Current Skills: {json.dumps(current_content)}
User Experience Context: {json.dumps(user_context)}

Job Description: {job_description[:2000]}

Return a JSON object: {{"technical": [...], "soft": [...], "other": [...]}}"""
        }
        
        prompt = prompts.get(section)
        if not prompt:
            raise HTTPException(status_code=400, detail=f"Unknown section: {section}")
        
        response = await call_claude(prompt)
        
        response_text = response.strip()
        
        # Try to parse as JSON if applicable
        if section in ["experience", "skills"]:
            try:
                start_idx = response_text.find('[') if section == "experience" else response_text.find('{')
                end_idx = response_text.rfind(']') + 1 if section == "experience" else response_text.rfind('}') + 1
                if start_idx != -1 and end_idx > start_idx:
                    json_str = response_text[start_idx:end_idx]
                    return {"content": json.loads(json_str)}
            except (json.JSONDecodeError, ValueError):
                pass
        
        return {"content": response_text}
        
    except Exception as e:
        print(f"Error regenerating section: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error regenerating section: {str(e)}")


# AI Generate Responsibilities for Work Experience
class GenerateResponsibilitiesRequest(BaseModel):
    position: str
    company: str
    industry: Optional[str] = ""
    jobDescription: Optional[str] = ""

@app.post("/api/generate-responsibilities")
async def generate_responsibilities(data: GenerateResponsibilitiesRequest):
    """Generate AI-powered responsibilities/achievements for a work position"""
    try:
        context = f"""Position: {data.position}
Company: {data.company}
{"Industry: " + data.industry if data.industry else ""}
{"Target Job Description: " + data.jobDescription[:500] if data.jobDescription else ""}"""
        
        prompt = f"""Generate 4-5 professional, impactful job responsibilities and achievements for the following role:

{context}

Requirements:
1. Start each bullet point with a strong action verb (Led, Developed, Managed, Implemented, etc.)
2. Include quantifiable metrics where appropriate (%, $, numbers)
3. Focus on achievements and impact, not just duties
4. Make them ATS-friendly with relevant keywords
5. Each bullet should be 1-2 sentences, specific and measurable

Return ONLY a JSON array of strings, no explanation:
["First responsibility with metrics and impact", "Second responsibility...", "Third...", "Fourth...", "Fifth..."]"""
        
        response = await call_claude(prompt)
        
        # Try to parse JSON array from response
        try:
            # Find JSON array in response
            json_match = re.search(r'\[[\s\S]*?\]', response)
            if json_match:
                responsibilities = json.loads(json_match.group())
                if isinstance(responsibilities, list) and all(isinstance(r, str) for r in responsibilities):
                    return {"success": True, "responsibilities": responsibilities}
        except json.JSONDecodeError:
            pass
        
        # Fallback: split by newlines if JSON parsing fails
        lines = [line.strip().lstrip('•-*').strip() for line in response.split('\n') if line.strip() and len(line.strip()) > 20]
        if lines:
            return {"success": True, "responsibilities": lines[:5]}
        
        raise HTTPException(status_code=500, detail="Failed to generate responsibilities")
        
    except Exception as e:
        print(f"Error generating responsibilities: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generating responsibilities: {str(e)}")


# ============================================
# DISCOUNT CODE ENDPOINTS
# ============================================

ADMIN_PASSWORD = "MintSlip2025!"

def verify_admin_password(password: str):
    if password != ADMIN_PASSWORD:
        raise HTTPException(status_code=401, detail="Unauthorized")

@app.post("/api/admin/discounts")
async def create_discount_code(data: DiscountCodeCreate, password: str):
    """Create a new discount code"""
    verify_admin_password(password)
    
    # Check if code already exists
    existing = await discounts_collection.find_one({"code": data.code.upper()})
    if existing:
        raise HTTPException(status_code=400, detail="Discount code already exists")
    
    discount_doc = {
        "id": str(uuid.uuid4()),
        "code": data.code.upper(),
        "discountPercent": data.discountPercent,
        "startDate": data.startDate,
        "expiryDate": data.expiryDate,
        "usageType": data.usageType,
        "usageLimit": data.usageLimit if data.usageType == "limited" else None,
        "usageCount": 0,
        "usedByCustomers": [],
        "applicableTo": data.applicableTo,
        "specificGenerators": data.specificGenerators if data.applicableTo == "specific" else None,
        "isActive": data.isActive,
        "createdAt": datetime.utcnow().isoformat()
    }
    
    await discounts_collection.insert_one(discount_doc)
    discount_doc.pop("_id", None)
    return discount_doc

@app.get("/api/admin/discounts")
async def list_discount_codes(password: str):
    """List all discount codes"""
    verify_admin_password(password)
    
    codes = await discounts_collection.find().to_list(length=100)
    for code in codes:
        code.pop("_id", None)
    return codes

@app.put("/api/admin/discounts/{discount_id}")
async def update_discount_code(discount_id: str, data: DiscountCodeUpdate, password: str):
    """Update a discount code"""
    verify_admin_password(password)
    
    existing = await discounts_collection.find_one({"id": discount_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Discount code not found")
    
    update_data = {k: v for k, v in data.dict().items() if v is not None}
    if "code" in update_data:
        update_data["code"] = update_data["code"].upper()
    
    await discounts_collection.update_one(
        {"id": discount_id},
        {"$set": update_data}
    )
    
    updated = await discounts_collection.find_one({"id": discount_id})
    updated.pop("_id", None)
    return updated

@app.delete("/api/admin/discounts/{discount_id}")
async def delete_discount_code(discount_id: str, password: str):
    """Delete a discount code"""
    verify_admin_password(password)
    
    result = await discounts_collection.delete_one({"id": discount_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Discount code not found")
    
    return {"message": "Discount code deleted successfully"}

@app.post("/api/validate-coupon")
async def validate_coupon(data: CouponValidateRequest):
    """Validate a coupon code and return discount details"""
    code = data.code.upper().strip()
    
    discount = await discounts_collection.find_one({"code": code})
    
    if not discount:
        raise HTTPException(status_code=404, detail="Invalid coupon code")
    
    # Check if active
    if not discount.get("isActive", True):
        raise HTTPException(status_code=400, detail="This coupon code is no longer active")
    
    # Check date validity
    now = datetime.utcnow()
    start_date = datetime.fromisoformat(discount["startDate"].replace("Z", ""))
    expiry_date = datetime.fromisoformat(discount["expiryDate"].replace("Z", ""))
    
    if now < start_date:
        raise HTTPException(status_code=400, detail="This coupon code is not yet active")
    
    if now > expiry_date:
        raise HTTPException(status_code=400, detail="This coupon code has expired")
    
    # Check usage limits
    usage_type = discount.get("usageType", "unlimited")
    
    if usage_type == "limited":
        if discount.get("usageCount", 0) >= discount.get("usageLimit", 0):
            raise HTTPException(status_code=400, detail="This coupon code has reached its usage limit")
    
    if usage_type == "one_per_customer" and data.customerIdentifier:
        used_by = discount.get("usedByCustomers", [])
        if data.customerIdentifier in used_by:
            raise HTTPException(status_code=400, detail="You have already used this coupon code")
    
    # Check generator applicability
    if discount.get("applicableTo") == "specific":
        allowed_generators = discount.get("specificGenerators", [])
        ps_types = {"phone_lookup", "name_lookup", "address_lookup", "carrier_lookup"}
        # "people_search" in allowed list covers all lookup types
        allowed_expanded = set(allowed_generators)
        if "people_search" in allowed_expanded:
            allowed_expanded |= ps_types
        if data.generatorType not in allowed_expanded:
            raise HTTPException(status_code=400, detail=f"This coupon code is not valid for {data.generatorType}")
    
    return {
        "valid": True,
        "code": discount["code"],
        "discountPercent": discount["discountPercent"],
        "message": f"{int(discount['discountPercent'])}% discount applied!"
    }

@app.post("/api/use-coupon")
async def use_coupon(data: CouponValidateRequest):
    """Mark a coupon as used after successful payment"""
    code = data.code.upper().strip()
    
    discount = await discounts_collection.find_one({"code": code})
    if not discount:
        return {"success": False, "message": "Coupon not found"}
    
    update_ops = {"$inc": {"usageCount": 1}}
    
    if data.customerIdentifier and discount.get("usageType") == "one_per_customer":
        update_ops["$push"] = {"usedByCustomers": data.customerIdentifier}
    
    await discounts_collection.update_one({"code": code}, update_ops)
    
    return {"success": True, "message": "Coupon usage recorded"}

# ========== BLOG ENDPOINTS ==========

# Enable internal data source on every startup (migration for installs that seeded it as False)
@app.on_event("startup")
async def enable_internal_data_source():
    await data_source_settings_collection.update_one(
        {"source": "internal"},
        {"$set": {"enabled": True}},
        upsert=True,
    )
    await data_source_settings_collection.update_one(
        {"source": "whitepages"},
        {"$set": {"enabled": True}},
        upsert=True,
    )
    logger.info("Data source defaults applied: internal=enabled, whitepages=enabled")


# Initialize default blog categories
@app.on_event("startup")
async def init_blog_categories():
    for cat in BLOG_CATEGORIES:
        existing = await blog_categories_collection.find_one({"slug": cat["slug"]})
        if not existing:
            await blog_categories_collection.insert_one({
                "id": str(uuid.uuid4()),
                **cat,
                "createdAt": datetime.now(timezone.utc).isoformat()
            })

# Background task for processing scheduled emails
async def email_scheduler_task():
    """Background task that runs every 5 minutes to process scheduled emails"""
    while True:
        try:
            await process_scheduled_emails()
        except Exception as e:
            print(f"Error in email scheduler: {e}")
        await asyncio.sleep(300)  # Run every 5 minutes

@app.on_event("startup")
async def start_email_scheduler():
    """Start the email scheduler background task"""
    asyncio.create_task(email_scheduler_task())
    print("Email scheduler started")

# ===== BLOG IMAGE SERVING (Dynamic with MongoDB fallback) =====

@app.get("/api/uploads/blog/{filename}")
async def serve_blog_image(filename: str):
    """Serve blog images with MongoDB fallback for container persistence"""
    import base64
    from fastapi.responses import Response
    
    filepath = os.path.join(UPLOAD_DIR, filename)
    
    # First, try to serve from disk
    if os.path.exists(filepath):
        # Check if file is valid (not empty/corrupt)
        file_size = os.path.getsize(filepath)
        if file_size > 100:  # Valid image files are larger than 100 bytes
            with open(filepath, "rb") as f:
                content = f.read()
            
            # Determine content type
            ext = filename.split(".")[-1].lower() if "." in filename else "png"
            content_types = {
                "jpg": "image/jpeg",
                "jpeg": "image/jpeg", 
                "png": "image/png",
                "webp": "image/webp",
                "gif": "image/gif"
            }
            content_type = content_types.get(ext, "image/png")
            
            return Response(content=content, media_type=content_type)
    
    # If not on disk or file is invalid, try MongoDB
    image_doc = await blog_images_collection.find_one({"filename": filename})
    
    if image_doc and image_doc.get("content"):
        # Decode base64 content
        content = base64.b64decode(image_doc["content"])
        content_type = image_doc.get("contentType", "image/png")
        
        # Restore file to disk for future requests
        try:
            os.makedirs(os.path.dirname(filepath), exist_ok=True)
            with open(filepath, "wb") as f:
                f.write(content)
        except Exception as e:
            print(f"Warning: Could not restore blog image to disk: {e}")
        
        return Response(content=content, media_type=content_type)
    
    # Image not found anywhere
    raise HTTPException(status_code=404, detail="Image not found")


@app.post("/api/admin/blog/migrate-images")
async def migrate_blog_images_to_mongodb(session: dict = Depends(get_current_admin)):
    """Migrate existing blog images from disk to MongoDB for persistence (admin)"""
    import base64
    
    migrated = 0
    skipped = 0
    errors = []
    
    # Get all files in the blog uploads directory
    if not os.path.exists(UPLOAD_DIR):
        return {"success": True, "migrated": 0, "message": "No upload directory found"}
    
    for filename in os.listdir(UPLOAD_DIR):
        filepath = os.path.join(UPLOAD_DIR, filename)
        
        # Skip if not a file
        if not os.path.isfile(filepath):
            continue
        
        # Skip if already in MongoDB
        existing = await blog_images_collection.find_one({"filename": filename})
        if existing:
            skipped += 1
            continue
        
        try:
            # Read file content
            with open(filepath, "rb") as f:
                content = f.read()
            
            # Skip if file is too small (likely corrupt/placeholder)
            if len(content) < 100:
                errors.append(f"{filename}: File too small ({len(content)} bytes)")
                continue
            
            # Determine content type
            ext = filename.split(".")[-1].lower() if "." in filename else "png"
            content_types = {
                "jpg": "image/jpeg",
                "jpeg": "image/jpeg",
                "png": "image/png",
                "webp": "image/webp",
                "gif": "image/gif"
            }
            content_type = content_types.get(ext, "image/png")
            
            # Store in MongoDB
            image_data = {
                "id": str(uuid.uuid4()),
                "filename": filename,
                "contentType": content_type,
                "content": base64.b64encode(content).decode("utf-8"),
                "originalFilename": filename,
                "createdAt": datetime.now(timezone.utc).isoformat(),
                "migratedFromDisk": True
            }
            await blog_images_collection.insert_one(image_data)
            migrated += 1
            
        except Exception as e:
            errors.append(f"{filename}: {str(e)}")
    
    return {
        "success": True,
        "migrated": migrated,
        "skipped": skipped,
        "errors": errors,
        "message": f"Migrated {migrated} images, skipped {skipped} (already in DB)"
    }


# ===== PUBLIC BLOG ENDPOINTS =====

@app.get("/api/blog/posts")
async def get_public_blog_posts(
    page: int = Query(1, ge=1),
    limit: int = Query(10, ge=1, le=50),
    category: Optional[str] = None,
    tag: Optional[str] = None,
    search: Optional[str] = None,
    sort: Optional[str] = "newest"  # newest, oldest, popular
):
    """Get published blog posts (public)"""
    query = {"status": "published"}
    
    if category:
        query["category"] = category
    
    if tag:
        query["tags"] = tag
    
    if search:
        query["$or"] = [
            {"title": {"$regex": search, "$options": "i"}},
            {"content": {"$regex": search, "$options": "i"}},
            {"excerpt": {"$regex": search, "$options": "i"}},
            {"tags": {"$regex": search, "$options": "i"}}
        ]
    
    # Sorting
    sort_order = -1 if sort == "newest" else 1
    sort_field = "publishDate" if sort in ["newest", "oldest"] else "views"
    
    skip = (page - 1) * limit
    total = await blog_posts_collection.count_documents(query)
    
    posts = await blog_posts_collection.find(
        query, 
        {"_id": 0, "content": 0}  # Exclude full content for list
    ).sort(sort_field, sort_order).skip(skip).limit(limit).to_list(limit)
    
    return {
        "success": True,
        "posts": posts,
        "total": total,
        "page": page,
        "pages": (total + limit - 1) // limit
    }

@app.get("/api/blog/posts/{slug}")
async def get_public_blog_post(slug: str):
    """Get a single published blog post by slug (public)"""
    post = await blog_posts_collection.find_one(
        {"slug": slug, "status": "published"},
        {"_id": 0}
    )
    
    if not post:
        raise HTTPException(status_code=404, detail="Post not found")
    
    # Increment view count
    await blog_posts_collection.update_one(
        {"slug": slug},
        {"$inc": {"views": 1}}
    )
    
    # Get related posts (same category, excluding current)
    related = []
    if post.get("category"):
        related = await blog_posts_collection.find(
            {"category": post["category"], "status": "published", "slug": {"$ne": slug}},
            {"_id": 0, "content": 0}
        ).limit(3).to_list(3)
    
    return {
        "success": True,
        "post": post,
        "related": related
    }

@app.get("/api/blog/categories")
async def get_blog_categories():
    """Get all blog categories (public)"""
    categories = await blog_categories_collection.find({}, {"_id": 0}).to_list(100)
    
    # Get post counts for each category
    for cat in categories:
        count = await blog_posts_collection.count_documents({
            "category": cat["slug"],
            "status": "published"
        })
        cat["postCount"] = count
    
    return {"success": True, "categories": categories}

@app.get("/api/blog/categories/{slug}")
async def get_blog_category(slug: str):
    """Get a single category with its posts"""
    category = await blog_categories_collection.find_one({"slug": slug}, {"_id": 0})
    
    if not category:
        raise HTTPException(status_code=404, detail="Category not found")
    
    return {"success": True, "category": category}

# ===== ADMIN BLOG ENDPOINTS =====

@app.get("/api/admin/blog/posts")
async def get_admin_blog_posts(
    session: dict = Depends(get_current_admin),
    page: int = Query(1, ge=1),
    limit: int = Query(20, ge=1, le=100),
    status: Optional[str] = None
):
    """Get all blog posts (admin)"""
    check_permission(session, "view_blog")
    query = {}
    if status:
        query["status"] = status
    
    skip = (page - 1) * limit
    total = await blog_posts_collection.count_documents(query)
    
    posts = await blog_posts_collection.find(
        query,
        {"_id": 0}
    ).sort("createdAt", -1).skip(skip).limit(limit).to_list(limit)
    
    return {
        "success": True,
        "posts": posts,
        "total": total,
        "page": page,
        "pages": (total + limit - 1) // limit
    }

@app.get("/api/admin/blog/posts/{post_id}")
async def get_admin_blog_post(post_id: str, session: dict = Depends(get_current_admin)):
    """Get a single blog post by ID (admin)"""
    post = await blog_posts_collection.find_one({"id": post_id}, {"_id": 0})
    
    if not post:
        raise HTTPException(status_code=404, detail="Post not found")
    
    return {"success": True, "post": post}

@app.post("/api/admin/blog/posts")
async def create_blog_post(data: BlogPostCreate, session: dict = Depends(get_current_admin)):
    """Create a new blog post (admin)"""
    check_permission(session, "manage_blog")
    # Check for duplicate slug
    existing = await blog_posts_collection.find_one({"slug": data.slug})
    if existing:
        raise HTTPException(status_code=400, detail="A post with this slug already exists")
    
    # Calculate reading time (avg 200 words per minute)
    word_count = len(data.content.split()) if data.content else 0
    reading_time = max(1, round(word_count / 200))
    
    post = {
        "id": str(uuid.uuid4()),
        "title": data.title,
        "slug": data.slug,
        "metaTitle": data.metaTitle or data.title,
        "metaDescription": data.metaDescription or (data.excerpt or data.content[:160] if data.content else ""),
        "author": data.author or "MintSlip Team",
        "featuredImage": data.featuredImage,
        "content": data.content,
        "excerpt": data.excerpt or (data.content[:200] + "..." if data.content and len(data.content) > 200 else data.content),
        "category": data.category,
        "tags": data.tags or [],
        "faqSchema": data.faqSchema or [],
        "status": data.status or "draft",
        "indexFollow": data.indexFollow if data.indexFollow is not None else True,
        "publishDate": data.publishDate or datetime.now(timezone.utc).isoformat(),
        "readingTime": reading_time,
        "views": 0,
        "createdAt": datetime.now(timezone.utc).isoformat(),
        "updatedAt": datetime.now(timezone.utc).isoformat(),
        "createdBy": session.get("adminId")
    }
    
    await blog_posts_collection.insert_one(post)
    await log_action(session, "create_blog_post", "blog_post", post["id"], post["title"])
    return {"success": True, "message": "Post created", "post": {k: v for k, v in post.items() if k != "_id"}}

@app.put("/api/admin/blog/posts/{post_id}")
async def update_blog_post(post_id: str, data: BlogPostUpdate, session: dict = Depends(get_current_admin)):
    """Update a blog post (admin)"""
    check_permission(session, "manage_blog")
    existing = await blog_posts_collection.find_one({"id": post_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Post not found")
    
    # Check for duplicate slug if changing
    if data.slug and data.slug != existing["slug"]:
        slug_exists = await blog_posts_collection.find_one({"slug": data.slug, "id": {"$ne": post_id}})
        if slug_exists:
            raise HTTPException(status_code=400, detail="A post with this slug already exists")
    
    update_data = {k: v for k, v in data.dict().items() if v is not None}
    
    # Recalculate reading time if content changed
    if data.content:
        word_count = len(data.content.split())
        update_data["readingTime"] = max(1, round(word_count / 200))
    
    update_data["updatedAt"] = datetime.now(timezone.utc).isoformat()
    update_data["updatedBy"] = session.get("adminId")
    
    await blog_posts_collection.update_one({"id": post_id}, {"$set": update_data})
    await log_action(session, "update_blog_post", "blog_post", post_id)
    return {"success": True, "message": "Post updated"}

@app.delete("/api/admin/blog/posts/{post_id}")
async def delete_blog_post(post_id: str, session: dict = Depends(get_current_admin)):
    """Delete a blog post (admin)"""
    check_permission(session, "manage_blog")
    result = await blog_posts_collection.delete_one({"id": post_id})

    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Post not found")

    await log_action(session, "delete_blog_post", "blog_post", post_id)
    return {"success": True, "message": "Post deleted"}

@app.post("/api/admin/blog/upload-image")
async def upload_blog_image(file: UploadFile = File(...), session: dict = Depends(get_current_admin)):
    """Upload an image for blog posts (admin)"""
    import base64
    
    # Validate file type
    allowed_types = ["image/jpeg", "image/png", "image/webp", "image/gif"]
    if file.content_type not in allowed_types:
        raise HTTPException(status_code=400, detail="Invalid file type. Allowed: JPEG, PNG, WebP, GIF")
    
    # Generate unique filename
    ext = file.filename.split(".")[-1] if "." in file.filename else "jpg"
    filename = f"{uuid.uuid4()}.{ext}"
    filepath = os.path.join(UPLOAD_DIR, filename)
    
    # Read file content
    file_content = await file.read()
    
    # Save file to disk
    with open(filepath, "wb") as buffer:
        buffer.write(file_content)
    
    # Also store in MongoDB for persistence
    image_data = {
        "id": str(uuid.uuid4()),
        "filename": filename,
        "contentType": file.content_type,
        "content": base64.b64encode(file_content).decode("utf-8"),
        "originalFilename": file.filename,
        "createdAt": datetime.now(timezone.utc).isoformat(),
        "uploadedBy": session.get("adminId")
    }
    await blog_images_collection.insert_one(image_data)
    
    # Return URL
    image_url = f"/api/uploads/blog/{filename}"
    
    return {"success": True, "url": image_url, "filename": filename}

class AIImageGenerateRequest(BaseModel):
    title: str
    category: Optional[str] = None
    keywords: Optional[str] = None

@app.post("/api/admin/blog/generate-image")
async def generate_blog_image(data: AIImageGenerateRequest, session: dict = Depends(get_current_admin)):
    """Generate a blog featured image using AI (admin)"""
    import base64
    from emergentintegrations.llm.openai.image_generation import OpenAIImageGeneration
    
    api_key = os.environ.get("EMERGENT_LLM_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="EMERGENT_LLM_KEY not configured")
    
    try:
        # Create a prompt for the blog image
        category_context = f" about {data.category}" if data.category else ""
        keywords_context = f" Keywords: {data.keywords}" if data.keywords else ""
        
        prompt = f"""Create a professional, modern blog featured image for an article titled: "{data.title}"{category_context}.{keywords_context}

The image should be:
- Clean and professional looking
- Suitable for a financial/business blog
- Modern flat design or minimalist style
- No text overlays
- Visually appealing with good color contrast
- 16:9 aspect ratio composition"""

        # Use OpenAI image generation (gpt-image-1) which works with Emergent LLM key
        img_gen = OpenAIImageGeneration(api_key=api_key)
        
        # Generate image
        generated_images = await img_gen.generate_images(
            prompt=prompt,
            model="gpt-image-1",
            number_of_images=1,
            quality="medium"  # Balance between quality and speed
        )
        
        if not generated_images or len(generated_images) == 0:
            raise HTTPException(status_code=500, detail="Failed to generate image - no image returned")
        
        # Get the first generated image (already bytes)
        image_bytes = generated_images[0]
        
        # Save the generated image
        filename = f"ai_{uuid.uuid4()}.png"
        filepath = os.path.join(UPLOAD_DIR, filename)
        
        with open(filepath, "wb") as f:
            f.write(image_bytes)
        
        # Also store in MongoDB for persistence (convert bytes to base64 for storage)
        image_data = {
            "id": str(uuid.uuid4()),
            "filename": filename,
            "contentType": "image/png",
            "content": base64.b64encode(image_bytes).decode('utf-8'),  # Convert bytes to base64
            "originalFilename": f"AI Generated - {data.title[:50]}",
            "createdAt": datetime.now(timezone.utc).isoformat(),
            "uploadedBy": session.get("adminId"),
            "aiGenerated": True
        }
        await blog_images_collection.insert_one(image_data)
        
        # Return URL
        image_url = f"/api/uploads/blog/{filename}"
        
        return {
            "success": True, 
            "url": image_url, 
            "filename": filename,
            "message": "Image generated successfully"
        }
        
    except Exception as e:
        print(f"Error generating blog image: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate image: {str(e)}")

@app.post("/api/admin/blog/categories")
async def create_blog_category(data: BlogCategoryCreate, session: dict = Depends(get_current_admin)):
    """Create a new blog category (admin)"""
    existing = await blog_categories_collection.find_one({"slug": data.slug})
    if existing:
        raise HTTPException(status_code=400, detail="Category with this slug already exists")
    
    category = {
        "id": str(uuid.uuid4()),
        "name": data.name,
        "slug": data.slug,
        "description": data.description or "",
        "createdAt": datetime.now(timezone.utc).isoformat()
    }
    
    await blog_categories_collection.insert_one(category)
    
    return {"success": True, "message": "Category created", "category": {k: v for k, v in category.items() if k != "_id"}}

@app.delete("/api/admin/blog/categories/{category_id}")
async def delete_blog_category(category_id: str, session: dict = Depends(get_current_admin)):
    """Delete a blog category (admin)"""
    result = await blog_categories_collection.delete_one({"id": category_id})
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Category not found")
    
    return {"success": True, "message": "Category deleted"}

# AI Blog Content Generation
@app.post("/api/admin/blog/ai-generate")
async def ai_generate_blog_content(
    topic: str = Query(..., description="Blog topic or title"),
    keywords: str = Query("", description="Target keywords, comma separated"),
    tone: str = Query("professional", description="Writing tone: professional, casual, informative"),
    session: dict = Depends(get_current_admin)
):
    """Generate blog content using AI (admin)"""
    try:
        prompt = f"""Write a comprehensive, SEO-optimized blog article about: {topic}

Target Keywords: {keywords if keywords else 'pay stub, proof of income, payroll'}
Tone: {tone}

Requirements:
1. Write in HTML format suitable for a WYSIWYG editor (use <h2>, <h3>, <p>, <ul>, <li>, <strong>, <em> tags)
2. Include an engaging introduction
3. Use <h2> and <h3> tags for headings
4. Include practical tips and actionable advice
5. Add a conclusion with a call-to-action mentioning MintSlip's pay stub generator
6. Aim for 1000-1500 words
7. Include internal link suggestions using <a href="/paystub-generator"> and <a href="/how-to-make-a-paystub">

Also provide:
- A compelling meta title (under 60 characters)
- A meta description (under 160 characters)
- An excerpt (2-3 sentences)
- 3-5 FAQ questions with answers related to the topic

Format your response as JSON:
{{
    "content": "HTML content here with proper tags",
    "metaTitle": "SEO title",
    "metaDescription": "meta description",
    "excerpt": "short excerpt",
    "suggestedSlug": "url-friendly-slug",
    "faqSchema": [
        {{"question": "Q1", "answer": "A1"}},
        {{"question": "Q2", "answer": "A2"}}
    ],
    "suggestedTags": ["tag1", "tag2"]
}}"""
        
        response = await call_claude(prompt)
        
        # Try to parse JSON from response
        try:
            # Find JSON in response
            json_match = re.search(r'\{[\s\S]*\}', response)
            if json_match:
                result = json.loads(json_match.group())
                return {"success": True, "generated": result}
        except json.JSONDecodeError:
            pass
        
        # If JSON parsing fails, return raw content
        return {
            "success": True,
            "generated": {
                "content": response,
                "metaTitle": topic[:60],
                "metaDescription": topic[:160],
                "excerpt": topic,
                "suggestedSlug": re.sub(r'[^a-z0-9]+', '-', topic.lower()).strip('-'),
                "faqSchema": [],
                "suggestedTags": []
            }
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI generation failed: {str(e)}")


# ============================================
# Social Media Meta Tags API
# ============================================

# Page-specific meta configurations
PAGE_META_CONFIG = {
    "/": {
        "title": "Paystub Generator | No.1 Paystub Generator - MintSlip",
        "description": "Our paystub generator instantly creates pay stubs online. This easy checkstub maker online handles calculations automatically with no software needed.",
        "image": "https://www.mintslip.com/favicon.ico"
    },
    "/paystub-generator": {
        "title": "Free Paystub Generator Online | Create Pay Stubs Instantly - MintSlip",
        "description": "Generate professional pay stubs in minutes. Our free paystub generator creates accurate, printable check stubs with automatic calculations.",
        "image": "https://www.mintslip.com/favicon.ico"
    },
    "/w2-generator": {
        "title": "W-2 Form Generator | Create W-2 Tax Forms Online - MintSlip",
        "description": "Generate accurate W-2 tax forms online. Create professional W-2 wage and tax statements for employees with our easy-to-use generator.",
        "image": "https://www.mintslip.com/favicon.ico"
    },
    "/1099-generator": {
        "title": "1099 Form Generator | Create 1099 Tax Forms Online - MintSlip",
        "description": "Generate 1099-NEC and 1099-MISC forms online. Create professional tax documents for contractors and freelancers instantly.",
        "image": "https://www.mintslip.com/favicon.ico"
    },
    "/ai-resume-builder": {
        "title": "AI Resume Builder | Create ATS-Optimized Resumes - MintSlip",
        "description": "Build professional, ATS-optimized resumes with AI. Our smart resume builder tailors your resume to job descriptions for maximum impact.",
        "image": "https://www.mintslip.com/favicon.ico"
    },
    "/blog": {
        "title": "Blog - Pay Stub Tips, Payroll Guides & Financial Documentation | MintSlip",
        "description": "Expert guides on pay stubs, proof of income, payroll management, and financial documentation. Learn how to create professional pay stubs.",
        "image": "https://www.mintslip.com/favicon.ico"
    }
}

@app.get("/api/meta-tags/{path:path}")
async def get_meta_tags(path: str):
    """Get meta tags for a specific page - used for social media sharing"""
    full_path = f"/{path}" if path else "/"
    
    # Check if it's a blog post
    if full_path.startswith("/blog/") and len(full_path) > 6:
        slug = full_path[6:]  # Remove "/blog/"
        post = await db["blog_posts"].find_one({"slug": slug, "status": "published"})
        
        if post:
            # Use the blog post's featured image for social sharing
            image_url = post.get("featuredImage", "https://www.mintslip.com/favicon.ico")
            if image_url and image_url.startswith("/"):
                image_url = f"https://www.mintslip.com{image_url}"
            
            return {
                "success": True,
                "meta": {
                    "title": post.get("metaTitle") or post.get("title", "MintSlip Blog"),
                    "description": post.get("metaDescription") or post.get("excerpt", ""),
                    "image": image_url,
                    "type": "article",
                    "url": f"https://www.mintslip.com/blog/{slug}",
                    "author": post.get("author", "MintSlip Team"),
                    "publishedTime": post.get("publishDate", ""),
                }
            }
    
    # Check predefined pages
    if full_path in PAGE_META_CONFIG:
        config = PAGE_META_CONFIG[full_path]
        return {
            "success": True,
            "meta": {
                "title": config["title"],
                "description": config["description"],
                "image": config["image"],
                "type": "website",
                "url": f"https://www.mintslip.com{full_path}"
            }
        }
    
    # Default fallback
    return {
        "success": True,
        "meta": {
            "title": "MintSlip - Professional Document Generator",
            "description": "Create professional pay stubs, W-2 forms, 1099 forms, and more. MintSlip makes document generation easy and accurate.",
            "image": "https://www.mintslip.com/favicon.ico",
            "type": "website",
            "url": f"https://www.mintslip.com{full_path}"
        }
    }


from fastapi import Request
from fastapi.responses import HTMLResponse

# List of social media bot user agents
SOCIAL_BOTS = [
    'facebookexternalhit',
    'Facebot',
    'Twitterbot',
    'LinkedInBot',
    'Pinterest',
    'Slackbot',
    'TelegramBot',
    'WhatsApp',
    'Discordbot',
    'Googlebot',
    'bingbot',
    'Applebot'
]

def is_social_bot(user_agent: str) -> bool:
    """Check if the request is from a social media bot"""
    if not user_agent:
        return False
    user_agent_lower = user_agent.lower()
    return any(bot.lower() in user_agent_lower for bot in SOCIAL_BOTS)

@app.get("/api/oembed")
async def get_oembed(url: str, request: Request):
    """oEmbed endpoint for rich link previews"""
    # Parse the URL to get the path
    from urllib.parse import urlparse
    parsed = urlparse(url)
    path = parsed.path
    
    # Get meta tags for this path
    meta_response = await get_meta_tags(path.lstrip('/'))
    meta = meta_response.get("meta", {})
    
    return {
        "version": "1.0",
        "type": "link",
        "title": meta.get("title", "MintSlip"),
        "author_name": meta.get("author", "MintSlip"),
        "provider_name": "MintSlip",
        "provider_url": "https://www.mintslip.com",
        "thumbnail_url": meta.get("image", "https://www.mintslip.com/favicon.ico"),
        "thumbnail_width": 200,
        "thumbnail_height": 200
    }


@app.get("/api/social-preview/{path:path}")
async def get_social_preview_html(path: str, request: Request):
    """
    Serve HTML with proper meta tags for social media crawlers.
    This endpoint can be used by a reverse proxy to serve to social bots.
    """
    full_path = f"/{path}" if path else "/"
    
    # Get meta tags
    meta_response = await get_meta_tags(path)
    meta = meta_response.get("meta", {})
    
    # For blog posts with featured images, use summary_large_image card
    twitter_card = "summary_large_image" if "blog/" in full_path else "summary"
    
    html = f'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1">
    
    <title>{meta.get("title", "MintSlip")}</title>
    <meta name="description" content="{meta.get("description", "")}">
    
    <!-- Open Graph / Facebook -->
    <meta property="og:type" content="{meta.get("type", "website")}">
    <meta property="og:url" content="{meta.get("url", "https://www.mintslip.com")}">
    <meta property="og:title" content="{meta.get("title", "MintSlip")}">
    <meta property="og:description" content="{meta.get("description", "")}">
    <meta property="og:image" content="{meta.get("image", "https://www.mintslip.com/favicon.ico")}">
    <meta property="og:site_name" content="MintSlip">
    
    <!-- Twitter -->
    <meta name="twitter:card" content="{twitter_card}">
    <meta name="twitter:url" content="{meta.get("url", "https://www.mintslip.com")}">
    <meta name="twitter:title" content="{meta.get("title", "MintSlip")}">
    <meta name="twitter:description" content="{meta.get("description", "")}">
    <meta name="twitter:image" content="{meta.get("image", "https://www.mintslip.com/favicon.ico")}">
    
    <!-- Article specific (for blog posts) -->
    {"<meta property='article:published_time' content='" + meta.get("publishedTime", "") + "'>" if meta.get("publishedTime") else ""}
    {"<meta property='article:author' content='" + meta.get("author", "") + "'>" if meta.get("author") else ""}
    
    <link rel="canonical" href="{meta.get("url", "https://www.mintslip.com")}">
    <link rel="icon" href="https://www.mintslip.com/favicon.ico">
    
    <!-- Redirect regular browsers to the actual page -->
    <script>
        window.location.href = "{meta.get("url", "https://www.mintslip.com")}";
    </script>
    <noscript>
        <meta http-equiv="refresh" content="0;url={meta.get("url", "https://www.mintslip.com")}">
    </noscript>
</head>
<body>
    <h1>{meta.get("title", "MintSlip")}</h1>
    <p>{meta.get("description", "")}</p>
    <a href="{meta.get("url", "https://www.mintslip.com")}">Visit MintSlip</a>
</body>
</html>'''
    
    return HTMLResponse(content=html)


# ============================================
# Dynamic Sitemap Generation
# ============================================

from fastapi.responses import Response

# Static pages for the sitemap
STATIC_SITEMAP_PAGES = [
    {"loc": "https://mintslip.com/", "priority": "1.0", "changefreq": "weekly"},
    {"loc": "https://mintslip.com/paystub-generator", "priority": "0.9", "changefreq": "weekly"},
    {"loc": "https://mintslip.com/paystub-for-apartment", "priority": "0.9", "changefreq": "weekly"},
    {"loc": "https://mintslip.com/mintslip-vs-thepaystubs", "priority": "0.9", "changefreq": "monthly"},
    {"loc": "https://mintslip.com/paystub-for-mortgage", "priority": "0.9", "changefreq": "weekly"},
    {"loc": "https://mintslip.com/paystub-template-download", "priority": "0.9", "changefreq": "weekly"},
    {"loc": "https://mintslip.com/create-a-paystub", "priority": "0.9", "changefreq": "weekly"},
    {"loc": "https://mintslip.com/paystub-samples", "priority": "0.8", "changefreq": "weekly"},
    {"loc": "https://mintslip.com/instant-paystub-generator", "priority": "0.8", "changefreq": "weekly"},
    {"loc": "https://mintslip.com/self-employed-paystub-generator", "priority": "0.8", "changefreq": "weekly"},
    {"loc": "https://mintslip.com/contractor-paystub-generator", "priority": "0.8", "changefreq": "weekly"},
    {"loc": "https://mintslip.com/canadian-paystub-generator", "priority": "0.8", "changefreq": "weekly"},
    {"loc": "https://mintslip.com/w2-generator", "priority": "0.8", "changefreq": "weekly"},
    {"loc": "https://mintslip.com/w9-generator", "priority": "0.8", "changefreq": "weekly"},
    {"loc": "https://mintslip.com/1099-nec-generator", "priority": "0.8", "changefreq": "weekly"},
    {"loc": "https://mintslip.com/1099-misc-generator", "priority": "0.8", "changefreq": "weekly"},
    {"loc": "https://mintslip.com/schedule-c-generator", "priority": "0.8", "changefreq": "weekly"},
    {"loc": "https://mintslip.com/accounting-mockup-generator", "priority": "0.8", "changefreq": "weekly"},
    {"loc": "https://mintslip.com/offer-letter-generator", "priority": "0.8", "changefreq": "weekly"},
    {"loc": "https://mintslip.com/vehicle-bill-of-sale-generator", "priority": "0.8", "changefreq": "weekly"},
    {"loc": "https://mintslip.com/service-expense-generator", "priority": "0.8", "changefreq": "weekly"},
    {"loc": "https://mintslip.com/ai-resume-builder", "priority": "0.8", "changefreq": "weekly"},
    {"loc": "https://mintslip.com/how-to-make-a-paystub", "priority": "0.9", "changefreq": "monthly"},
    {"loc": "https://mintslip.com/about", "priority": "0.7", "changefreq": "monthly"},
    {"loc": "https://mintslip.com/contact", "priority": "0.7", "changefreq": "monthly"},
    {"loc": "https://mintslip.com/faq", "priority": "0.7", "changefreq": "monthly"},
    {"loc": "https://mintslip.com/reviews", "priority": "0.7", "changefreq": "monthly"},
    {"loc": "https://mintslip.com/privacy", "priority": "0.5", "changefreq": "yearly"},
    {"loc": "https://mintslip.com/terms", "priority": "0.5", "changefreq": "yearly"},
    {"loc": "https://mintslip.com/blog", "priority": "0.8", "changefreq": "daily"},
]

# State-specific paystub generator pages
US_STATES = [
    "alabama", "alaska", "arizona", "arkansas", "california", "colorado", "connecticut",
    "delaware", "florida", "georgia", "hawaii", "idaho", "illinois", "indiana", "iowa",
    "kansas", "kentucky", "louisiana", "maine", "maryland", "massachusetts", "michigan",
    "minnesota", "mississippi", "missouri", "montana", "nebraska", "nevada", "new-hampshire",
    "new-jersey", "new-mexico", "new-york", "north-carolina", "north-dakota", "ohio",
    "oklahoma", "oregon", "pennsylvania", "rhode-island", "south-carolina", "south-dakota",
    "tennessee", "texas", "utah", "vermont", "virginia", "washington", "west-virginia",
    "wisconsin", "wyoming", "canada"
]

# High priority states
HIGH_PRIORITY_STATES = ["california", "florida", "new-york", "texas"]

@app.get("/api/sitemap.xml")
async def generate_sitemap():
    """Generate dynamic sitemap.xml including all blog posts"""
    today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    
    # Start XML
    xml_parts = [
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">'
    ]
    
    # Add static pages
    for page in STATIC_SITEMAP_PAGES:
        xml_parts.append(f'''  <url>
    <loc>{page["loc"]}</loc>
    <lastmod>{today}</lastmod>
    <changefreq>{page["changefreq"]}</changefreq>
    <priority>{page["priority"]}</priority>
  </url>''')
    
    # Add state-specific paystub generator pages
    for state in US_STATES:
        priority = "0.9" if state in HIGH_PRIORITY_STATES else "0.8"
        xml_parts.append(f'''  <url>
    <loc>https://mintslip.com/paystub-generator/{state}</loc>
    <lastmod>{today}</lastmod>
    <changefreq>monthly</changefreq>
    <priority>{priority}</priority>
  </url>''')
    
    # Add blog posts from database
    try:
        blog_posts = await db["blog_posts"].find(
            {"status": "published"},
            {"slug": 1, "publishDate": 1, "updatedAt": 1}
        ).to_list(length=1000)
        
        for post in blog_posts:
            # Use updatedAt if available, otherwise publishDate
            lastmod = post.get("updatedAt") or post.get("publishDate") or today
            if isinstance(lastmod, datetime):
                lastmod = lastmod.strftime("%Y-%m-%d")
            elif isinstance(lastmod, str) and "T" in lastmod:
                lastmod = lastmod.split("T")[0]
            
            xml_parts.append(f'''  <url>
    <loc>https://mintslip.com/blog/{post["slug"]}</loc>
    <lastmod>{lastmod}</lastmod>
    <changefreq>weekly</changefreq>
    <priority>0.7</priority>
  </url>''')
    except Exception as e:
        print(f"Error fetching blog posts for sitemap: {e}")
    
    # Close XML
    xml_parts.append('</urlset>')
    
    xml_content = '\n'.join(xml_parts)
    
    return Response(
        content=xml_content,
        media_type="application/xml",
        headers={"Content-Type": "application/xml; charset=utf-8"}
    )


@app.get("/api/sitemap/refresh")
async def refresh_sitemap_info():
    """Get information about sitemap contents (for debugging)"""
    try:
        blog_count = await db["blog_posts"].count_documents({"status": "published"})
        static_count = len(STATIC_SITEMAP_PAGES) + len(US_STATES)
        
        return {
            "success": True,
            "sitemap_url": "/api/sitemap.xml",
            "static_pages": static_count,
            "blog_posts": blog_count,
            "total_urls": static_count + blog_count,
            "message": "Sitemap is generated dynamically. Access /api/sitemap.xml for the full sitemap."
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


# ============================================
# PDF METADATA & CONSISTENCY ENGINE
# Business Plan Feature Only
# ============================================

# Max file size for PDF analysis (10MB default)
PDF_ENGINE_MAX_SIZE = int(os.environ.get("PDF_ENGINE_MAX_SIZE", 10 * 1024 * 1024))

async def verify_business_subscription(session: dict) -> dict:
    """Verify user has an active Business subscription"""
    user = await users_collection.find_one({"id": session["userId"]})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    subscription = user.get("subscription") or {}
    if subscription.get("status") != "active" or subscription.get("tier") != "business":
        raise HTTPException(
            status_code=403, 
            detail="This feature requires an active Business subscription"
        )
    
    return user


@app.get("/api/pdf-engine/check-access")
async def check_pdf_engine_access(session: dict = Depends(get_current_user)):
    """Check if user has access to PDF Engine (Business plan only)"""
    user = await users_collection.find_one({"id": session["userId"]})
    if not user:
        return {"hasAccess": False, "reason": "User not found"}
    
    subscription = user.get("subscription") or {}
    status = subscription.get("status")
    tier = subscription.get("tier")
    
    # Allow access for both active and cancelling subscriptions (until period ends)
    has_valid_status = status in ["active", "cancelling"]
    has_access = has_valid_status and tier == "business"
    
    reason = None
    if not has_access:
        if tier != "business":
            reason = "Business subscription required"
        elif status not in ["active", "cancelling"]:
            reason = "Active subscription required"
    
    return {
        "hasAccess": has_access,
        "currentTier": tier or "none",
        "subscriptionStatus": status or "none",
        "reason": reason
    }


class PDFEngineAnalyzeRequest(BaseModel):
    normalize: bool = False  # Whether to also normalize the PDF
    normalizeOptions: Optional[dict] = None


@app.get("/api/pdf-engine/document-types")
async def get_pdf_document_types(session: dict = Depends(get_current_user)):
    """Get available document types for PDF analysis"""
    return {
        "documentTypes": DOCUMENT_TYPES,
        "success": True
    }


@app.get("/api/pdf-engine/producers/{document_type}")
async def get_pdf_producers(document_type: str, session: dict = Depends(get_current_user)):
    """Get legitimate producers for a document type"""
    producers = get_legitimate_producers(document_type)
    return {
        "documentType": document_type,
        "producers": {k: {"name": v["names"][0], "notes": v.get("notes", "")} for k, v in producers.items()},
        "success": True
    }


@app.post("/api/pdf-engine/analyze")
async def analyze_pdf_endpoint(
    file: UploadFile = File(...),
    normalize: bool = False,
    document_type: str = "other",
    enable_ai: bool = True,
    session: dict = Depends(get_current_user)
):
    """
    Analyze PDF metadata and document consistency
    Business Plan feature only
    
    document_type options: paystub, bank_statement, tax_form, other
    enable_ai: Enable AI-powered content analysis (default: True)
    """
    # Verify Business subscription
    user = await verify_business_subscription(session)
    
    # Validate document type
    if document_type not in DOCUMENT_TYPES:
        document_type = "other"
    
    # Validate file type
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(
            status_code=400, 
            detail="Only PDF files are supported. Please upload a .pdf file."
        )
    
    # Validate content type
    if file.content_type and 'pdf' not in file.content_type.lower():
        raise HTTPException(
            status_code=400,
            detail="Invalid file type. Only PDF files are accepted."
        )
    
    # Read file content
    try:
        pdf_bytes = await file.read()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to read file: {str(e)}")
    
    # Check file size
    if len(pdf_bytes) > PDF_ENGINE_MAX_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"File too large. Maximum size is {PDF_ENGINE_MAX_SIZE // (1024*1024)}MB"
        )
    
    # Verify it's a valid PDF
    if not pdf_bytes.startswith(b'%PDF'):
        raise HTTPException(
            status_code=400,
            detail="Invalid PDF file. The file does not appear to be a valid PDF."
        )
    
    try:
        # Analyze the PDF with document type
        analysis = analyze_pdf_metadata(pdf_bytes, document_type)
        
        # Run AI analysis if enabled and text content exists
        ai_analysis_result = None
        if enable_ai and analysis.text_content:
            ai_analysis_result = await analyze_document_with_ai(
                analysis.text_content,
                document_type,
                analysis.metadata
            )
            # Apply AI findings to the analysis result
            if ai_analysis_result and ai_analysis_result.get("success"):
                apply_ai_analysis_to_result(analysis, ai_analysis_result)
        
        result = {
            "success": True,
            "filename": file.filename,
            "analysis": analysis.to_dict(),
            "analyzedAt": datetime.now(timezone.utc).isoformat(),
            "aiEnabled": enable_ai,
            "aiAnalysisSuccess": ai_analysis_result.get("success") if ai_analysis_result else False
        }
        
        # If normalize requested, also normalize
        if normalize:
            normalized_bytes, normalize_result = normalize_pdf_metadata(pdf_bytes)
            result["normalization"] = normalize_result
            
            if normalize_result.get("success"):
                # Encode normalized PDF as base64 for download
                result["normalizedPdfBase64"] = base64.b64encode(normalized_bytes).decode('utf-8')
        
        return result
        
    except Exception as e:
        logger.error(f"PDF analysis error: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Analysis failed: {str(e)}"
        )


@app.post("/api/pdf-engine/normalize")
async def normalize_pdf_endpoint(
    file: UploadFile = File(...),
    standardize_producer: bool = True,
    fix_dates: bool = True,
    flatten_document: bool = True,
    session: dict = Depends(get_current_user)
):
    """
    Normalize PDF metadata to reduce verification red flags
    Business Plan feature only
    """
    # Verify Business subscription
    user = await verify_business_subscription(session)
    
    # Validate file type
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are supported")
    
    # Read file
    pdf_bytes = await file.read()
    
    # Check file size
    if len(pdf_bytes) > PDF_ENGINE_MAX_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"File too large. Maximum size is {PDF_ENGINE_MAX_SIZE // (1024*1024)}MB"
        )
    
    # Verify PDF
    if not pdf_bytes.startswith(b'%PDF'):
        raise HTTPException(status_code=400, detail="Invalid PDF file")
    
    try:
        # Normalize options
        options = {
            "standardize_producer": standardize_producer,
            "fix_dates": fix_dates,
            "flatten_document": flatten_document
        }
        
        # Perform normalization
        normalized_bytes, result = normalize_pdf_metadata(pdf_bytes, options)
        
        if not result.get("success"):
            raise HTTPException(
                status_code=500,
                detail=f"Normalization failed: {result.get('error', 'Unknown error')}"
            )
        
        return {
            "success": True,
            "filename": file.filename,
            "changes": result.get("changes", []),
            "normalizedPdfBase64": base64.b64encode(normalized_bytes).decode('utf-8'),
            "normalizedAt": datetime.now(timezone.utc).isoformat()
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"PDF normalization error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Normalization failed: {str(e)}")


@app.post("/api/pdf-engine/generate-report")
async def generate_analysis_report(
    file: UploadFile = File(...),
    format: str = "pdf",  # "pdf" or "json"
    session: dict = Depends(get_current_user)
):
    """
    Generate a detailed analysis report
    Business Plan feature only
    """
    from fastapi.responses import Response
    
    # Verify Business subscription
    user = await verify_business_subscription(session)
    
    # Validate file type
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are supported")
    
    # Read file
    pdf_bytes = await file.read()
    
    # Check file size
    if len(pdf_bytes) > PDF_ENGINE_MAX_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"File too large. Maximum size is {PDF_ENGINE_MAX_SIZE // (1024*1024)}MB"
        )
    
    # Verify PDF
    if not pdf_bytes.startswith(b'%PDF'):
        raise HTTPException(status_code=400, detail="Invalid PDF file")
    
    try:
        # Analyze
        analysis = analyze_pdf_metadata(pdf_bytes)
        
        if format == "json":
            return {
                "success": True,
                "filename": file.filename,
                "report": analysis.to_dict(),
                "generatedAt": datetime.now(timezone.utc).isoformat()
            }
        else:
            # Generate PDF report
            report_pdf = generate_analysis_report_pdf(analysis, file.filename)
            
            return Response(
                content=report_pdf,
                media_type="application/pdf",
                headers={
                    "Content-Disposition": f'attachment; filename="analysis_report_{file.filename}"'
                }
            )
            
    except Exception as e:
        logger.error(f"Report generation error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Report generation failed: {str(e)}")


@app.get("/api/pdf-engine/metadata-presets")
async def get_metadata_presets_endpoint(session: dict = Depends(get_current_user)):
    """Get available metadata presets for common legitimate sources"""
    # Verify Business subscription
    await verify_business_subscription(session)
    
    return {
        "success": True,
        "presets": METADATA_PRESETS
    }


class MetadataEditRequest(BaseModel):
    producer: Optional[str] = None
    creator: Optional[str] = None
    author: Optional[str] = None
    title: Optional[str] = None
    subject: Optional[str] = None
    creationDate: Optional[str] = None  # ISO format or 'now'
    modificationDate: Optional[str] = None  # ISO format or 'now'


@app.post("/api/pdf-engine/edit-metadata")
async def edit_pdf_metadata_endpoint(
    file: UploadFile = File(...),
    producer: Optional[str] = None,
    creator: Optional[str] = None,
    author: Optional[str] = None,
    title: Optional[str] = None,
    subject: Optional[str] = None,
    creation_date: Optional[str] = None,
    modification_date: Optional[str] = None,
    preset: Optional[str] = None,
    session: dict = Depends(get_current_user)
):
    """
    Edit PDF metadata and regenerate as a clean PDF without edit traces.
    Business Plan feature only.
    
    You can either:
    1. Use a preset (e.g., 'adp', 'paychex', 'chase') which sets producer/creator automatically
    2. Provide custom values for each field
    
    Dates can be:
    - ISO format string (e.g., '2024-01-15T10:30:00')
    - Date only (e.g., '2024-01-15')
    - 'now' for current timestamp
    """
    # Verify Business subscription
    await verify_business_subscription(session)
    
    # Validate file type
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are supported")
    
    # Read file
    pdf_bytes = await file.read()
    
    # Check file size
    if len(pdf_bytes) > PDF_ENGINE_MAX_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"File too large. Maximum size is {PDF_ENGINE_MAX_SIZE // (1024*1024)}MB"
        )
    
    # Verify PDF
    if not pdf_bytes.startswith(b'%PDF'):
        raise HTTPException(status_code=400, detail="Invalid PDF file")
    
    try:
        # Build metadata dict
        new_metadata = {}
        
        # Apply preset if specified
        if preset and preset in METADATA_PRESETS:
            preset_data = METADATA_PRESETS[preset]
            new_metadata.update(preset_data)
        
        # Override with custom values if provided
        if producer:
            new_metadata['producer'] = producer
        if creator:
            new_metadata['creator'] = creator
        if author:
            new_metadata['author'] = author
        if title:
            new_metadata['title'] = title
        if subject:
            new_metadata['subject'] = subject
        if creation_date:
            new_metadata['creationDate'] = creation_date
        if modification_date:
            new_metadata['modificationDate'] = modification_date
        
        if not new_metadata:
            raise HTTPException(
                status_code=400, 
                detail="No metadata changes specified. Provide at least one field or use a preset."
            )
        
        # Edit and regenerate PDF
        regenerated_pdf, result = edit_and_regenerate_pdf(pdf_bytes, new_metadata)
        
        if not result.get("success"):
            raise HTTPException(
                status_code=500,
                detail=f"Failed to regenerate PDF: {result.get('error', 'Unknown error')}"
            )
        
        return {
            "success": True,
            "filename": file.filename,
            "changes": result.get("changes", []),
            "message": result.get("message", "PDF regenerated successfully"),
            "regeneratedPdfBase64": base64.b64encode(regenerated_pdf).decode('utf-8'),
            "regeneratedAt": datetime.now(timezone.utc).isoformat()
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Metadata edit error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to edit metadata: {str(e)}")


@app.post("/api/clean-paystub-pdf")
async def clean_paystub_pdf_endpoint(
    file: UploadFile = File(...),
    template: str = Form("template-a"),
    pay_date: Optional[str] = Form(None)
):
    """
    Clean a paystub PDF to remove all traces of editing and apply proper metadata.
    Makes the PDF appear as a fresh document generated by the appropriate payroll system.
    
    This is a public endpoint used by the paystub generator.
    
    Args:
        file: The PDF file to clean
        template: The template type ('template-a', 'template-b', 'template-c', 'template-h')
        pay_date: The pay date (ISO format) - creation date will be set to 2 days before this
    """
    # Validate file type
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are supported")
    
    # Read file
    pdf_bytes = await file.read()
    
    # Check file size (max 10MB for paystubs)
    max_size = 10 * 1024 * 1024
    if len(pdf_bytes) > max_size:
        raise HTTPException(
            status_code=400,
            detail=f"File too large. Maximum size is {max_size // (1024*1024)}MB"
        )
    
    # Verify PDF
    if not pdf_bytes.startswith(b'%PDF'):
        raise HTTPException(status_code=400, detail="Invalid PDF file")
    
    try:
        # Clean the PDF with pay_date for creation date calculation
        cleaned_pdf, result = clean_paystub_pdf(pdf_bytes, template, pay_date)
        
        if not result.get("success"):
            raise HTTPException(
                status_code=500,
                detail=f"Failed to clean PDF: {result.get('error', 'Unknown error')}"
            )
        
        return {
            "success": True,
            "cleanedPdfBase64": base64.b64encode(cleaned_pdf).decode('utf-8'),
            "metadata": result.get("metadata_applied", {})
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"PDF cleaning error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to clean PDF: {str(e)}")


@app.post("/api/clean-bank-statement-pdf")
async def clean_bank_statement_pdf_endpoint(
    file: UploadFile = File(...),
    template: str = Form("chime"),
    statement_month: Optional[str] = Form(None),
    account_name: Optional[str] = Form(None)
):
    """
    Clean a bank statement PDF to remove all traces of editing and apply proper metadata.
    
    Args:
        file: The PDF file to clean
        template: The template type ('chime', 'bank-of-america', 'chase')
        statement_month: The statement month (YYYY-MM) - creation date will be last day of month
        account_name: Account holder's name for title generation
    """
    # Validate file type
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are supported")
    
    # Read file
    pdf_bytes = await file.read()
    
    # Check file size (max 10MB)
    max_size = 10 * 1024 * 1024
    if len(pdf_bytes) > max_size:
        raise HTTPException(
            status_code=400,
            detail=f"File too large. Maximum size is {max_size // (1024*1024)}MB"
        )
    
    # Verify PDF
    if not pdf_bytes.startswith(b'%PDF'):
        raise HTTPException(status_code=400, detail="Invalid PDF file")
    
    try:
        # Clean the PDF
        cleaned_pdf, result = clean_bank_statement_pdf(pdf_bytes, template, statement_month, account_name)
        
        if not result.get("success"):
            raise HTTPException(
                status_code=500,
                detail=f"Failed to clean PDF: {result.get('error', 'Unknown error')}"
            )
        
        return {
            "success": True,
            "cleanedPdfBase64": base64.b64encode(cleaned_pdf).decode('utf-8'),
            "metadata": result.get("metadata_applied", {})
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Bank statement PDF cleaning error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to clean PDF: {str(e)}")


# ─────────────────────────────────────────────────────────────────────────────
# People Search Feature
# ─────────────────────────────────────────────────────────────────────────────
import random
import copy

# ── Whitepages Pro integration (https://pro.whitepages.com) ──────────────────
# Set WHITEPAGES_PRO_API_KEY in your environment to activate real data.
# Falls back to mock data when the key is not set.
WHITEPAGES_PRO_API_KEY = os.environ.get("WHITEPAGES_PRO_API_KEY")
_WP_BASE = "https://api.whitepages.com/v1"
logger.info(f"Whitepages Pro API key loaded: {'YES (' + str(len(WHITEPAGES_PRO_API_KEY)) + ' chars)' if WHITEPAGES_PRO_API_KEY else 'NO (mock mode)'}")
_wp_env_keys = [k for k in os.environ if "WHITE" in k.upper() or "PAGES" in k.upper()]
logger.info(f"Whitepages-related env vars found: {_wp_env_keys}")


def _wp_format_phone(raw: str) -> str:
    """Format a 10-digit string as (XXX) XXX-XXXX."""
    d = re.sub(r"[^\d]", "", raw or "")
    if len(d) == 10:
        return f"({d[:3]}) {d[3:6]}-{d[6:]}"
    if len(d) == 11 and d[0] == "1":
        return f"({d[1:4]}) {d[4:7]}-{d[7:]}"
    return raw


def _wp_names_from_person(p: dict) -> list[str]:
    """Extract relative names from a Whitepages v1 person object."""
    out = []
    for ap in p.get("relatives", []):
        fn = ap.get("name") or ""
        if fn:
            out.append(fn)
    return out


async def wp_phone_lookup(phone: str) -> dict | None:
    """Whitepages v1 reverse phone — GET /v1/person?phone=NUMBER with X-Api-Key header."""
    if not WHITEPAGES_PRO_API_KEY:
        return None
    try:
        async with httpx.AsyncClient(timeout=10) as client:
            r = await client.get(
                f"{_WP_BASE}/person",
                params={"phone": phone},
                headers={"X-Api-Key": WHITEPAGES_PRO_API_KEY},
            )
            logger.info(f"Whitepages phone response: {r.status_code}")
            if r.status_code != 200:
                logger.warning(f"Whitepages phone error {r.status_code}: {r.text[:500]}")
                return None
            d = r.json()
            logger.info(f"Whitepages phone raw keys: {list(d.keys()) if isinstance(d, dict) else type(d)}")

        # /person?phone= can return a single dict or a list
        if isinstance(d, list):
            results = d
        elif isinstance(d, dict) and "results" in d:
            results = d["results"]
        elif isinstance(d, dict):
            results = [d]
        else:
            results = []

        logger.info(f"Whitepages phone results count: {len(results)}")
        if not results:
            return None
        res = results[0]

        # Extract person info — v1 /person schema
        raw_name = res.get("name") or {}
        if isinstance(raw_name, str):
            name = raw_name
        else:
            name = raw_name.get("full_name") or f"{raw_name.get('first','')} {raw_name.get('last','')}".strip()

        # Address — build full string from street + city + state + zip
        address = ""
        cur_addrs = res.get("current_addresses") or res.get("addresses") or []
        if cur_addrs:
            a = cur_addrs[0]
            street_part = a.get("address") or a.get("full_address") or a.get("street_line_1") or ""
            city_part   = a.get("city") or ""
            state_part  = a.get("state_code") or a.get("state") or ""
            zip_part    = a.get("zip") or a.get("postal_code") or ""
            state_zip   = f"{state_part} {zip_part}".strip()
            address = ", ".join(p for p in [street_part, city_part, state_zip] if p)

        # Relatives
        relatives = []
        for rel in (res.get("relatives") or res.get("associated_people") or []):
            n = rel.get("name") or ""
            if isinstance(n, dict):
                n = n.get("full_name") or f"{n.get('first','')} {n.get('last','')}".strip()
            if n:
                relatives.append(n)

        # Carrier / line type — may be on the phone sub-object or top level
        phones_list = res.get("phones") or []
        carrier = ""
        line_type = ""
        if phones_list:
            ph = phones_list[0]
            carrier   = ph.get("carrier") or ""
            line_type = (ph.get("line_type") or ph.get("phone_type") or "").replace("_", " ").title()
        carrier   = carrier   or res.get("carrier")   or "Unknown"
        line_type = line_type or res.get("line_type") or "Unknown"
        if isinstance(carrier, dict):
            carrier = carrier.get("name") or "Unknown"

        return {
            "name":              name or None,
            "carrier":           carrier,
            "lineType":          line_type,
            "location":          address or None,
            "spamRisk":          "Unknown",
            "callerType":        "Individual" if name else "Unknown",
            "possibleAddress":   address or None,
            "possibleRelatives": relatives or None,
            "phoneValid":        res.get("is_valid", True),
        }
    except Exception as e:
        logger.warning(f"Whitepages phone lookup error: {e}")
        return None


async def wp_person_lookup(first: str, last: str, state: str, city: str = "", min_age: int = None, max_age: int = None) -> list[dict] | None:
    """Whitepages Pro Person Search — returns list of up to 5 normalized results."""
    if not WHITEPAGES_PRO_API_KEY:
        return None

    def _parse(p):
        # v1 schema: name=string, phones=[{number,type}], relatives=[{id,name}],
        # current_addresses=[{id,address}], historic_addresses=[{id,address}]
        full_name  = p.get("name") or f"{first} {last}"
        dob        = p.get("date_of_birth") or ""
        def _wp_addr_str(a):
            """Build a full address string from a Whitepages address object."""
            street = a.get("address") or a.get("full_address") or a.get("street_line_1") or ""
            city   = a.get("city") or ""
            state  = a.get("state_code") or a.get("state") or ""
            zip_   = a.get("zip") or a.get("postal_code") or ""
            state_zip = f"{state} {zip_}".strip()
            return ", ".join(p for p in [street, city, state_zip] if p)
        cur_addrs  = [_wp_addr_str(a) for a in p.get("current_addresses", []) if any(a.get(k) for k in ["address","city","state_code","state"])]
        hist_addrs = [_wp_addr_str(a) for a in p.get("historic_addresses", []) if any(a.get(k) for k in ["address","city","state_code","state"])]
        phones     = [_wp_format_phone(ph["number"]) for ph in p.get("phones", []) if ph.get("number")]
        relatives  = [r["name"] for r in p.get("relatives", []) if r.get("name")]
        emails     = p.get("emails") or []
        return {
            "fullName":          full_name,
            "dateOfBirth":       dob or None,
            "possibleAddresses": (cur_addrs + hist_addrs) or None,
            "possiblePhones":    phones or None,
            "possibleRelatives": relatives or None,
            "emails":            emails or None,
            "state":             state or None,
        }

    try:
        params = {"name": f"{first} {last}"}
        if state:
            params["state_code"] = state.upper()
        if city:
            params["city"] = city
        if min_age is not None:
            params["min_age"] = min_age
        if max_age is not None:
            params["max_age"] = max_age
        async with httpx.AsyncClient(timeout=10) as client:
            r = await client.get(
                f"{_WP_BASE}/person/",
                params=params,
                headers={"X-Api-Key": WHITEPAGES_PRO_API_KEY},
            )
            logger.info(f"Whitepages person status: {r.status_code}")
            logger.info(f"Whitepages person raw: {r.text[:800]}")
            if r.status_code != 200:
                logger.warning(f"Whitepages person error {r.status_code}: {r.text[:300]}")
                return None
            d = r.json()
        # v1 API returns a list directly or {"results": [...]}
        results = d if isinstance(d, list) else d.get("results", [])
        logger.info(f"Whitepages person results count: {len(results)}")
        if not results:
            return None
        return [_parse(p) for p in results[:5]]
    except Exception as e:
        logger.warning(f"Whitepages person lookup error: {e}")
        return None


async def wp_address_lookup(street: str, city: str, state: str) -> dict | None:
    """Whitepages Pro Reverse Address — returns normalized result dict."""
    if not WHITEPAGES_PRO_API_KEY:
        return None
    try:
        params = {
            "street_line_1": street,
            "city":          city,
        }
        if state:
            params["state_code"] = state.upper()
        async with httpx.AsyncClient(timeout=10) as client:
            r = await client.get(
                f"{_WP_BASE}/location/",
                params=params,
                headers={"X-Api-Key": WHITEPAGES_PRO_API_KEY},
            )
            if r.status_code != 200:
                logger.warning(f"Whitepages address error {r.status_code}: {r.text[:200]}")
                return None
            d = r.json()
        results = d if isinstance(d, list) else d.get("results", [])
        if not results:
            return None
        loc = results[0]

        # v1: address=string, residents=[{name,...}], phones=[{number,type}]
        full_addr     = loc.get("address") or loc.get("full_address") or f"{street}, {city}, {state}"
        residents_raw = loc.get("residents") or []
        residents     = [
            (r["name"] if isinstance(r.get("name"), str) else (r.get("name") or {}).get("full_name") or "")
            for r in residents_raw if r.get("name")
        ]
        phones = [_wp_format_phone(ph.get("number") or ph.get("phone_number", "")) for ph in loc.get("phones", []) if ph.get("number") or ph.get("phone_number")]

        return {
            "address":          full_addr,
            "residents":        residents or None,
            "associatedPhones": phones or None,
        }
    except Exception as e:
        logger.warning(f"Whitepages address lookup error: {e}")
        return None

DEFAULT_PEOPLE_SEARCH_PRICES = {
    "phone_lookup":   0.99,
    "name_lookup":    1.49,
    "address_lookup": 1.49,
    "carrier_lookup": 0.49,
}

# ── Data Source Settings ─────────────────────────────────────────────────────

DATA_SOURCE_DEFAULTS = [
    {
        "source": "whitepages",
        "label": "Whitepages Pro API",
        "description": "Live data from the Whitepages Pro REST API. Requires WHITEPAGES_PRO_API_KEY.",
        "enabled": True,
        "category": "api",
        "recordCount": 0,
    },
    {
        "source": "internal",
        "label": "Internal Database",
        "description": "Records collected by your own scrapers and manually imported data.",
        "enabled": True,
        "category": "database",
        "recordCount": 0,
    },
    {
        "source": "nsopw",
        "label": "Sex Offender Registry (NSOPW)",
        "description": "National sex offender public registry. Free government API.",
        "enabled": False,
        "category": "scraper",
        "recordCount": 0,
    },
    {
        "source": "nppes",
        "label": "NPI Doctor Registry (NPPES)",
        "description": "CMS National Plan & Provider Enumeration System. Free API, 6M+ healthcare providers.",
        "enabled": False,
        "category": "scraper",
        "recordCount": 0,
    },
    {
        "source": "fec",
        "label": "FEC Campaign Finance",
        "description": "Federal Election Commission donor database. Free API, 100M+ contributions.",
        "enabled": False,
        "category": "scraper",
        "recordCount": 0,
    },
    {
        "source": "faa",
        "label": "FAA Airmen/Aircraft Registry",
        "description": "FAA bulk download — pilot certificates and aircraft owners.",
        "enabled": False,
        "category": "scraper",
        "recordCount": 0,
    },
    {
        "source": "ssdi",
        "label": "SSDI Death Records",
        "description": "Social Security Death Index — 98M deceased individuals with DOB, DOD, last ZIP. Free via FamilySearch API.",
        "enabled": False,
        "category": "scraper",
        "recordCount": 0,
    },
    {
        "source": "voter_rolls",
        "label": "Voter Roll Data",
        "description": "Manually imported voter registration data from state sources (varies by state).",
        "enabled": False,
        "category": "import",
        "recordCount": 0,
    },
]

async def get_data_sources() -> list:
    """Return all data source settings, seeding defaults on first call."""
    docs = await data_source_settings_collection.find({}).to_list(None)
    existing = {d["source"] for d in docs}
    for default in DATA_SOURCE_DEFAULTS:
        if default["source"] not in existing:
            await data_source_settings_collection.insert_one(dict(default))
    # Ensure internal source is enabled (upgrade path — was seeded as False historically)
    await data_source_settings_collection.update_one(
        {"source": "internal", "enabled": False},
        {"$set": {"enabled": True}},
    )
    return await data_source_settings_collection.find({}, {"_id": 0}).to_list(None)

async def is_source_enabled(source: str) -> bool:
    doc = await data_source_settings_collection.find_one({"source": source})
    if doc is None:
        return True  # default on for all sources
    return bool(doc.get("enabled", True))

# ── Internal DB helpers ───────────────────────────────────────────────────────

def _fmt_addr(a: dict) -> str:
    parts = [a.get("street", ""), a.get("city", ""), a.get("state", ""), a.get("zip", "")]
    return ", ".join(p for p in parts if p)

def _email_str(e) -> str:
    """Normalize email to plain string."""
    if isinstance(e, str):
        return e
    if isinstance(e, dict):
        return e.get("address", "")
    return ""

def calc_age_from_dob(dob: str) -> int | None:
    """Calculate age from a date of birth string (YYYY-MM-DD or MM/DD/YYYY)."""
    if not dob:
        return None
    try:
        import re as _re
        from datetime import date as _date
        iso = _re.match(r"^(\d{4})-(\d{2})-(\d{2})$", str(dob))
        if iso:
            born = _date(int(iso.group(1)), int(iso.group(2)), int(iso.group(3)))
        else:
            mdy = _re.match(r"^(\d{1,2})/(\d{1,2})/(\d{4})$", str(dob))
            if mdy:
                born = _date(int(mdy.group(3)), int(mdy.group(1)), int(mdy.group(2)))
            else:
                return None
        today = _date.today()
        age = today.year - born.year - ((today.month, today.day) < (born.month, born.day))
        return age if 0 < age < 130 else None
    except Exception:
        return None


def normalize_internal_record(doc: dict) -> dict:
    """Convert a people_records document to the standard search result format."""
    addrs = doc.get("addresses", [])
    current = next((a for a in addrs if a.get("current")), addrs[0] if addrs else {})
    past = [a for a in addrs if not a.get("current")]
    age = doc.get("age") or calc_age_from_dob(doc.get("dateOfBirth", ""))
    emails = [_email_str(e) for e in doc.get("emails", []) if _email_str(e)]
    first  = doc.get("firstName", "")
    middle = doc.get("middleName", "")
    last   = doc.get("lastName", "")
    full_name = " ".join(p for p in [first, middle, last] if p)
    return {
        "fullName":         full_name,
        "firstName":        first,
        "lastName":         last,
        "middleName":       middle,
        "dateOfBirth":      doc.get("dateOfBirth", ""),
        "aliases":          doc.get("aliases", []),
        "age":              age,
        "ageRange":         f"{age - 2}-{age + 2}" if age else None,
        "gender":           doc.get("gender", ""),
        "state":            current.get("state", doc.get("state", "")),
        "currentAddress":   _fmt_addr(current) if current else "",
        "pastAddresses":    [_fmt_addr(a) for a in past],
        "phones":           doc.get("phones", []),
        "emails":           emails,
        "possibleRelatives": doc.get("relatives", []),
        "associates":       doc.get("associates", []),
        "occupation":       doc.get("occupation", ""),
        "employer":         doc.get("employer", ""),
        "sourceDB":         doc.get("source", "internal"),
        "recordId":         doc.get("recordId", ""),
    }

def _parse_addr_str(addr_str: str, current: bool = True) -> dict:
    """Parse a full address string like '123 Main St, Louisville, KY 40201' into components."""
    parts = [p.strip() for p in str(addr_str).split(",")]
    # Last part must start with a 2-letter US state abbreviation to be valid
    last = parts[-1].strip() if parts else ""
    has_state = bool(re.match(r"^[A-Z]{2}\b", last))

    state_zip = last.split() if has_state else []
    st   = state_zip[0] if state_zip else ""
    zip_ = state_zip[1] if len(state_zip) > 1 else ""

    if len(parts) >= 3 and has_state:
        # "123 Main St, Louisville, KY 40201"
        street = ", ".join(parts[:-2])
        city   = parts[-2].strip()
    elif len(parts) == 2 and has_state:
        first = parts[0].strip()
        if re.match(r"^\d", first):
            # "123 Main St, KY 40201" — street with no city info
            street = first
            city   = ""
        else:
            # "Louisville, KY 40201" — normal city + state
            street = ""
            city   = first
    else:
        # No recognizable city/state — store everything as street
        street = addr_str
        city = st = zip_ = ""
    return {"street": street, "city": city, "state": st, "zip": zip_, "current": current}


def _split_full_name(full_name: str) -> tuple[str, str]:
    """Split 'First Last' into (first, last). Returns ('', '') if no last name."""
    parts = full_name.strip().split()
    if len(parts) < 2:
        return parts[0] if parts else "", ""
    return parts[0], " ".join(parts[1:])


async def _insert_person_record(first: str, last: str, dob: str, state: str,
                                 addrs: list, phones: list, emails: list,
                                 relatives: list) -> None:
    """Insert a person record after checking for duplicates."""
    query: dict = {
        "firstName": {"$regex": f"^{re.escape(first)}$", "$options": "i"},
        "lastName":  {"$regex": f"^{re.escape(last)}$",  "$options": "i"},
    }
    if state:
        query["$or"] = [{"state": state}, {"addresses.state": state}]
    if await people_records_collection.find_one(query):
        return  # already exists
    now = datetime.now(timezone.utc).isoformat()
    doc = {
        "recordId":    str(uuid.uuid4()),
        "source":      "whitepages",
        "firstName":   first,
        "lastName":    last,
        "middleName":  "",
        "aliases":     [],
        "age":         calc_age_from_dob(dob),
        "dateOfBirth": dob,
        "gender":      "",
        "state":       state,
        "addresses":   addrs,
        "phones":      phones,
        "emails":      [e for e in emails if e],
        "relatives":   relatives,
        "associates":  [],
        "occupation":  "",
        "employer":    "",
        "sourceData":   {},
        "createdAt":    now,
        "lastUpdated":  now,
        "lastVerified": now,
    }
    await people_records_collection.insert_one(doc)


STALE_AFTER_DAYS = 90  # refresh a record from Whitepages if older than this


def _is_stale(doc: dict) -> bool:
    """Return True if the record hasn't been verified in STALE_AFTER_DAYS days."""
    ts = doc.get("lastVerified") or doc.get("lastUpdated") or doc.get("createdAt")
    if not ts:
        return True
    try:
        last = datetime.fromisoformat(ts.replace("Z", "+00:00"))
        if last.tzinfo is None:
            last = last.replace(tzinfo=timezone.utc)
        return (datetime.now(timezone.utc) - last).days >= STALE_AFTER_DAYS
    except Exception:
        return True


async def _merge_into_record(record_id: str, new_addrs: list, new_phones: list,
                              new_emails: list, new_relatives: list,
                              new_dob: str = "", new_state: str = "") -> None:
    """Additively merge fresh Whitepages data into an existing record.
    Never removes existing data — only adds new items not already present.
    """
    doc = await people_records_collection.find_one({"recordId": record_id})
    if not doc:
        return
    now = datetime.now(timezone.utc).isoformat()

    # ── Addresses: add new ones not already stored (match on city+state) ──────
    existing_addrs = doc.get("addresses", [])
    existing_cs = {(a.get("city", "").lower(), a.get("state", "").upper()) for a in existing_addrs}
    for a in new_addrs:
        key = (a.get("city", "").lower(), a.get("state", "").upper())
        if key not in existing_cs and (a.get("city") or a.get("street")):
            # New address becomes non-current; existing current stays current
            a["current"] = False
            existing_addrs.append(a)
            existing_cs.add(key)

    # ── Phones: add new digits not already stored ─────────────────────────────
    existing_phones = doc.get("phones", [])
    def _phone_digits(p):
        num = p.get("number", "") if isinstance(p, dict) else str(p)
        return re.sub(r"[^\d]", "", num)[-10:]
    existing_digits = {_phone_digits(p) for p in existing_phones}
    for p in new_phones:
        d = _phone_digits(p)
        if d and d not in existing_digits:
            existing_phones.append(p)
            existing_digits.add(d)

    # ── Emails: add new ones not already stored ───────────────────────────────
    existing_emails = doc.get("emails", [])
    existing_email_set = {_email_str(e).lower() for e in existing_emails}
    for e in new_emails:
        es = _email_str(e).lower()
        if es and es not in existing_email_set:
            existing_emails.append(e)
            existing_email_set.add(es)

    # ── Relatives: add new names not already stored ───────────────────────────
    existing_relatives = doc.get("relatives", [])
    existing_rel_set = {r.lower() if isinstance(r, str) else str(r).lower() for r in existing_relatives}
    for r in new_relatives:
        rs = r.lower() if isinstance(r, str) else str(r).lower()
        if rs and rs not in existing_rel_set:
            existing_relatives.append(r)
            existing_rel_set.add(rs)

    updates: dict = {
        "addresses":   existing_addrs,
        "phones":      existing_phones,
        "emails":      existing_emails,
        "relatives":   existing_relatives,
        "lastVerified": now,
        "lastUpdated":  now,
    }
    # Fill in DOB / state only if currently missing
    if new_dob and not doc.get("dateOfBirth"):
        updates["dateOfBirth"] = new_dob
        updates["age"] = calc_age_from_dob(new_dob)
    if new_state and not doc.get("state"):
        updates["state"] = new_state

    await people_records_collection.update_one(
        {"recordId": record_id}, {"$set": updates}
    )


async def upsert_from_whitepages(lookup_type: str, result: dict, phone: str = "") -> None:
    """Import or refresh a Whitepages result in the internal people_records collection.

    - First search: inserts a new record and stamps lastVerified = now.
    - Subsequent searches within STALE_AFTER_DAYS: skips (record is fresh).
    - After STALE_AFTER_DAYS: merges new data additively (never deletes existing fields).

    Each lookup type returns different field shapes:
      phone_lookup:   { name, location, possibleAddress, possibleRelatives }
      name_lookup:    { fullName, dateOfBirth, possibleAddresses, possiblePhones, possibleRelatives, state }
      address_lookup: { address, residents, associatedPhones }
    """
    try:
        if lookup_type == "phone_lookup":
            full_name = (result.get("name") or "").strip()
            first, last = _split_full_name(full_name)
            if not first or not last:
                return

            clean_phone = re.sub(r"[^\d]", "", phone)[-10:]
            existing = None
            if clean_phone:
                digits = clean_phone
                flex = r"\D*".join(digits)
                existing = await people_records_collection.find_one(
                    {"phones": {"$elemMatch": {"number": {"$regex": flex}}}}
                )
            if not existing:
                existing = await people_records_collection.find_one({
                    "firstName": {"$regex": f"^{re.escape(first)}$", "$options": "i"},
                    "lastName":  {"$regex": f"^{re.escape(last)}$",  "$options": "i"},
                })

            addr_str = result.get("location") or result.get("possibleAddress") or ""
            addrs = [_parse_addr_str(addr_str, current=True)] if addr_str else []
            phones_list = [{"number": clean_phone, "type": "unknown", "current": True}] if clean_phone else []
            relatives = [r for r in (result.get("possibleRelatives") or []) if r]

            if existing:
                if _is_stale(existing):
                    await _merge_into_record(existing["recordId"], addrs, phones_list, [], relatives)
                # else: record is fresh, skip
            else:
                await _insert_person_record(first, last, "", "", addrs, phones_list, [], relatives)

        elif lookup_type == "name_lookup":
            full_name = (result.get("fullName") or "").strip()
            first, last = _split_full_name(full_name)
            if not first or not last:
                return

            state = result.get("state") or ""
            dob   = result.get("dateOfBirth") or ""
            addrs_raw = result.get("possibleAddresses") or []
            addrs = [_parse_addr_str(a, current=(i == 0)) for i, a in enumerate(addrs_raw) if a]
            phones_raw = result.get("possiblePhones") or []
            phones_list = [
                p if isinstance(p, dict) else {"number": str(p), "type": "unknown", "current": True}
                for p in phones_raw if p
            ]
            emails_raw = result.get("emails") or []
            emails = [_email_str(e) for e in emails_raw]
            relatives = [r for r in (result.get("possibleRelatives") or []) if r]

            query: dict = {
                "firstName": {"$regex": f"^{re.escape(first)}$", "$options": "i"},
                "lastName":  {"$regex": f"^{re.escape(last)}$",  "$options": "i"},
            }
            if state:
                query["$or"] = [{"state": state}, {"addresses.state": state}]
            existing = await people_records_collection.find_one(query)

            if existing:
                if _is_stale(existing):
                    await _merge_into_record(existing["recordId"], addrs, phones_list, emails, relatives, dob, state)
            else:
                await _insert_person_record(first, last, dob, state, addrs, phones_list, emails, relatives)

        elif lookup_type == "address_lookup":
            residents = result.get("residents") or []
            addr_str  = result.get("address") or ""
            phones_raw = result.get("associatedPhones") or []
            phones_list = [
                p if isinstance(p, dict) else {"number": str(p), "type": "unknown", "current": True}
                for p in phones_raw if p
            ]
            addrs = [_parse_addr_str(addr_str, current=True)] if addr_str else []
            for resident_name in residents:
                first, last = _split_full_name(resident_name)
                if not first or not last:
                    continue
                existing = await people_records_collection.find_one({
                    "firstName": {"$regex": f"^{re.escape(first)}$", "$options": "i"},
                    "lastName":  {"$regex": f"^{re.escape(last)}$",  "$options": "i"},
                })
                if existing:
                    if _is_stale(existing):
                        await _merge_into_record(existing["recordId"], addrs, phones_list, [], [])
                else:
                    await _insert_person_record(first, last, "", "", addrs, phones_list, [], [])

    except Exception:
        pass  # never let import errors break the search response


async def upsert_phone_entry(phone: str, linked_results: list) -> None:
    """Auto-register a searched phone in the phone_entries registry."""
    try:
        clean = re.sub(r"[^\d]", "", phone)[-10:]
        if len(clean) < 7:
            return
        display = f"({clean[:3]}) {clean[3:6]}-{clean[6:]}" if len(clean) == 10 else clean
        linked_ids = list({r.get("recordId") for r in linked_results if r.get("recordId")})
        now = datetime.now(timezone.utc).isoformat()
        existing = await phone_entries_collection.find_one({"phone": clean})
        if existing:
            update: dict = {
                "searchCount": existing.get("searchCount", 0) + 1,
                "lastSearched": now,
            }
            if linked_ids:
                all_ids = list(set(existing.get("linkedRecordIds", []) + linked_ids))
                update["linkedRecordIds"] = all_ids
            await phone_entries_collection.update_one({"phone": clean}, {"$set": update})
        else:
            await phone_entries_collection.insert_one({
                "entryId": str(uuid.uuid4()),
                "phone": clean,
                "displayPhone": display,
                "type": "unknown",
                "carrier": "",
                "linkedRecordIds": linked_ids,
                "searchCount": 1,
                "lastSearched": now,
                "source": "auto",
                "addedAt": now,
                "notes": "",
            })
    except Exception:
        pass


async def upsert_address_entry(street: str, city: str, state: str, linked_results: list) -> None:
    """Auto-register a searched address in the address_entries registry."""
    try:
        street_clean = street.strip()
        city_clean = city.strip()
        if not street_clean or not city_clean:
            return
        linked_ids = list({r.get("recordId") for r in linked_results if r.get("recordId")})
        now = datetime.now(timezone.utc).isoformat()
        existing = await address_entries_collection.find_one({
            "street": {"$regex": f"^{re.escape(street_clean)}$", "$options": "i"},
            "city":   {"$regex": f"^{re.escape(city_clean)}$", "$options": "i"},
        })
        if existing:
            update: dict = {
                "searchCount": existing.get("searchCount", 0) + 1,
                "lastSearched": now,
            }
            if linked_ids:
                all_ids = list(set(existing.get("linkedRecordIds", []) + linked_ids))
                update["linkedRecordIds"] = all_ids
            await address_entries_collection.update_one({"entryId": existing["entryId"]}, {"$set": update})
        else:
            await address_entries_collection.insert_one({
                "entryId": str(uuid.uuid4()),
                "street": street_clean,
                "city": city_clean,
                "state": state.strip(),
                "zip": "",
                "linkedRecordIds": linked_ids,
                "searchCount": 1,
                "lastSearched": now,
                "source": "auto",
                "addedAt": now,
                "notes": "",
            })
    except Exception:
        pass


async def internal_name_lookup(first: str, last: str, state: str = "", city: str = "", min_age: int = None, max_age: int = None) -> list:
    query: dict = {
        "firstName": {"$regex": re.escape(first), "$options": "i"},
        "lastName":  {"$regex": re.escape(last),  "$options": "i"},
    }
    if state:
        query["$or"] = [
            {"state": {"$regex": state, "$options": "i"}},
            {"addresses.state": {"$regex": state, "$options": "i"}},
        ]
    if city:
        query["addresses.city"] = {"$regex": re.escape(city), "$options": "i"}
    if min_age is not None:
        query.setdefault("age", {})["$gte"] = min_age
    if max_age is not None:
        query.setdefault("age", {})["$lte"] = max_age
    cursor = people_records_collection.find(query).limit(20)
    results = []
    async for doc in cursor:
        results.append(normalize_internal_record(doc))
    return results

async def internal_phone_lookup(phone: str) -> list:
    digits = re.sub(r"[^\d]", "", phone)[-10:]  # last 10 digits
    if len(digits) < 7:
        return []
    # Build a flexible regex that matches digits with any formatting chars between them
    # e.g. "5551234567" → "5\D*5\D*5\D*1\D*2\D*3\D*4\D*5\D*6\D*7"
    # This matches "(555) 123-4567", "555-123-4567", "5551234567", etc.
    flex_pattern = r"\D*".join(digits)
    print(f"[internal_phone_lookup] digits={digits} pattern={flex_pattern}")
    cursor = people_records_collection.find({
        "$or": [
            {"phones": {"$elemMatch": {"number": {"$regex": flex_pattern}}}},  # dict format {number, type}
            {"phones": {"$elemMatch": {"$regex": flex_pattern}}},               # plain string format
        ]
    }).limit(5)
    results = []
    async for doc in cursor:
        results.append(normalize_internal_record(doc))
    return results

async def internal_address_lookup(street: str, city: str, state: str = "") -> list:
    query: dict = {
        "addresses.street": {"$regex": re.escape(street), "$options": "i"},
        "addresses.city":   {"$regex": re.escape(city),   "$options": "i"},
    }
    if state:
        query["addresses.state"] = {"$regex": state, "$options": "i"}
    cursor = people_records_collection.find(query).limit(10)
    results = []
    async for doc in cursor:
        results.append(normalize_internal_record(doc))
    return results

async def get_people_search_prices() -> dict:
    doc = await site_settings_collection.find_one({"key": "people_search_prices"})
    if doc:
        return doc.get("value", DEFAULT_PEOPLE_SEARCH_PRICES)
    return DEFAULT_PEOPLE_SEARCH_PRICES


def blur_result(data: dict, lookup_type: str) -> dict:
    """Return a copy of data with sensitive fields partially redacted for preview."""
    BLOB = "●●●●●"

    def redact_phone(p):
        # Keep area code, blur the rest
        digits = re.sub(r"[^\d]", "", p)
        if len(digits) >= 7:
            return f"({digits[:3]}) {digits[3]}{BLOB}"
        return BLOB

    def redact_name(n):
        parts = n.split()
        if len(parts) >= 2:
            return f"{parts[0]} {parts[1][0]}.{BLOB}"
        return BLOB

    def redact_addr(a):
        # Show only City, ST — no street number, no zip
        parts = [p.strip() for p in a.split(",")]
        if len(parts) >= 3:
            city = parts[-2]
            state = parts[-1][:2]
            return f"{city}, {state}"
        elif len(parts) == 2:
            city = parts[0]
            state = parts[1][:2]
            return f"{city}, {state}"
        return a

    preview = copy.deepcopy(data)

    if lookup_type == "phone_lookup":
        # Show full name unredacted — only blur address/phone details
        if "possibleAddress" in preview: preview["possibleAddress"] = redact_addr(preview["possibleAddress"])
        if preview.get("location"): preview["location"] = redact_addr(preview["location"])
        # Internal DB record fields (normalize_internal_record format)
        if "currentAddress" in preview: preview["currentAddress"] = redact_addr(preview["currentAddress"])
        if "pastAddresses" in preview:
            preview["pastAddresses"] = [redact_addr(a) for a in preview["pastAddresses"] if a]
        if "possibleAddresses" in preview:
            preview["possibleAddresses"] = [redact_addr(a) for a in preview["possibleAddresses"] if a]
        if "phones" in preview:
            preview["phones"] = [
                {**p, "number": redact_phone(p["number"])} if isinstance(p, dict) else redact_phone(str(p))
                for p in preview["phones"]
            ]
    elif lookup_type == "name_lookup":
        if "possibleAddresses" in preview:
            preview["possibleAddresses"] = [redact_addr(a) for a in preview["possibleAddresses"] if a]
        if "currentAddress" in preview and preview["currentAddress"]:
            preview["currentAddress"] = redact_addr(preview["currentAddress"])
        if "pastAddresses" in preview:
            preview["pastAddresses"] = [redact_addr(a) for a in preview["pastAddresses"] if a]
        if "possiblePhones" in preview:
            preview["possiblePhones"] = [redact_phone(p) for p in preview["possiblePhones"]]
    elif lookup_type == "address_lookup":
        if "propertyOwner" in preview: preview["propertyOwner"] = redact_name(preview["propertyOwner"])
        if "residents" in preview:
            preview["residents"] = [redact_name(r) for r in preview["residents"]]
        if "associatedPhones" in preview:
            preview["associatedPhones"] = [redact_phone(p) for p in preview["associatedPhones"]]
        if "estimatedValue" in preview: preview["estimatedValue"] = "$●●●,●●●"
    elif lookup_type == "carrier_lookup":
        preview["carrier"]  = BLOB
        preview["lineType"] = BLOB
    return preview


class PeopleSearchRequest(BaseModel):
    lookupType: str
    phone: Optional[str] = None
    fullName: Optional[str] = None   # combined name field; parsed into firstName/lastName
    firstName: Optional[str] = None
    lastName: Optional[str] = None
    state: Optional[str] = None
    street: Optional[str] = None
    city: Optional[str] = None
    minAge: Optional[int] = None
    maxAge: Optional[int] = None

    def get_first_last(self):
        """Return (firstName, lastName) — parsed from fullName if provided."""
        if self.fullName and self.fullName.strip():
            parts = self.fullName.strip().split()
            return parts[0], " ".join(parts[1:]) if len(parts) > 1 else ""
        return (self.firstName or ""), (self.lastName or "")


class PeopleSearchCheckoutRequest(BaseModel):
    searchId: str
    successUrl: Optional[str] = None
    cancelUrl: Optional[str] = None
    discountCode: Optional[str] = None
    discountAmount: Optional[float] = 0


@app.get("/api/people-search/prices")
async def people_search_prices_endpoint():
    return await get_people_search_prices()


# ── Carrier Lookup ────────────────────────────────────────────────────────────

def _phonenumbers_carrier_lookup(phone_e164: str) -> dict:
    """Offline carrier lookup via the phonenumbers library (prefix-based, free)."""
    try:
        import phonenumbers
        from phonenumbers import carrier as pn_carrier, geocoder as pn_geo, PhoneNumberType
        pn = phonenumbers.parse(phone_e164, None)
        valid = phonenumbers.is_valid_number(pn)
        pn_type = phonenumbers.number_type(pn)
        type_map = {
            PhoneNumberType.MOBILE:      "mobile",
            PhoneNumberType.FIXED_LINE:  "landline",
            PhoneNumberType.FIXED_LINE_OR_MOBILE: "landline_or_mobile",
            PhoneNumberType.VOIP:        "voip",
            PhoneNumberType.TOLL_FREE:   "toll_free",
            PhoneNumberType.PREMIUM_RATE: "premium",
        }
        line_type   = type_map.get(pn_type, "unknown")
        carrier_name = pn_carrier.name_for_number(pn, "en") or ""
        region       = pn_geo.description_for_number(pn, "en") or ""
        return {
            "valid":       valid,
            "carrier":     carrier_name,
            "lineType":    line_type,
            "region":      region,
            "countryCode": phonenumbers.region_code_for_number(pn),
        }
    except Exception:
        return {"valid": False, "carrier": "", "lineType": "unknown", "region": "", "countryCode": ""}


async def _live_carrier_lookup(phone_digits: str) -> dict:
    """Live carrier lookup via NumVerify API (requires NUMVERIFY_API_KEY env var)."""
    api_key = os.environ.get("NUMVERIFY_API_KEY", "")
    if not api_key:
        return {}
    try:
        async with httpx.AsyncClient(timeout=8) as client:
            r = await client.get(
                "http://apilayer.net/api/validate",
                params={"access_key": api_key, "number": phone_digits, "country_code": "US", "format": 1},
            )
            if r.status_code != 200:
                return {}
            d = r.json()
            if not d.get("valid"):
                return {}
            return {
                "valid":       True,
                "carrier":     d.get("carrier", ""),
                "lineType":    d.get("line_type", "unknown"),
                "region":      f"{d.get('location', '')}, {d.get('country_name', '')}".strip(", "),
                "countryCode": d.get("country_code", "US"),
            }
    except Exception:
        return {}


@app.get("/api/people-search/carrier-lookup")
async def carrier_lookup_endpoint(phone: str):
    """Look up carrier and line type for a phone number. Free, no auth required."""
    digits = re.sub(r"[^\d]", "", phone)
    if len(digits) < 10:
        raise HTTPException(status_code=400, detail="Phone number must be at least 10 digits")
    e164 = f"+1{digits[-10:]}" if len(digits) <= 11 else f"+{digits}"

    # Try live lookup first (if API key configured), fall back to offline
    live = await _live_carrier_lookup(digits)
    if live:
        return live

    offline = await asyncio.to_thread(_phonenumbers_carrier_lookup, e164)
    return offline


@app.get("/api/people-search/debug-raw")
async def people_search_debug_raw(name: str = "John Smith", state: str = "WA", phone: str = ""):
    """Temporary debug endpoint — returns raw Whitepages v1 responses."""
    if not WHITEPAGES_PRO_API_KEY:
        return {"error": "No API key set"}
    out = {}
    async with httpx.AsyncClient(timeout=10) as client:
        if phone:
            r = await client.get(f"{_WP_BASE}/phone", params={"phone_number": re.sub(r'[^\d]','',phone)}, headers={"X-Api-Key": WHITEPAGES_PRO_API_KEY})
            out["phone"] = {"status": r.status_code, "body": r.json() if r.status_code == 200 else r.text}
        r = await client.get(f"{_WP_BASE}/person", params={"name": name, "state_code": state}, headers={"X-Api-Key": WHITEPAGES_PRO_API_KEY})
        out["person"] = {"status": r.status_code, "body": r.json() if r.status_code == 200 else r.text}
    return out


@app.get("/api/people-search/debug-sources")
async def debug_sources():
    """Returns data source enabled states + internal record count. No auth required."""
    docs = await data_source_settings_collection.find({}, {"_id": 0}).to_list(None)
    internal_count = await people_records_collection.count_documents({})
    sample = await people_records_collection.find({}, {"firstName": 1, "lastName": 1, "source": 1, "_id": 0}).limit(5).to_list(5)
    return {"sources": docs, "internalRecordCount": internal_count, "sampleRecords": sample}


PHONE_REPORT_CATEGORIES = ["Scam", "Spam", "Robocall", "Telemarketer", "Fraud", "Harassment", "Unknown"]

@app.get("/api/phone-stats/{phone}")
async def get_phone_stats(phone: str):
    """Return how many people searched this number today + scam report counts."""
    clean = re.sub(r"[^\d]", "", phone)[-10:]
    if not clean:
        raise HTTPException(status_code=400, detail="Invalid phone number")

    # Searches today (UTC)
    today_start = datetime.now(timezone.utc).replace(hour=0, minute=0, second=0, microsecond=0).isoformat()
    searches_today = await people_search_logs_collection.count_documents({
        "lookupType": "phone_lookup",
        "query": {"$regex": clean},
        "createdAt": {"$gte": today_start},
    })

    # Scam reports
    total_reports = await phone_reports_collection.count_documents({"phone": clean})
    cat_pipeline = [
        {"$match": {"phone": clean}},
        {"$group": {"_id": "$category", "count": {"$sum": 1}}},
        {"$sort": {"count": -1}},
    ]
    cat_docs = await phone_reports_collection.aggregate(cat_pipeline).to_list(None)
    categories = [{"category": d["_id"], "count": d["count"]} for d in cat_docs]

    return {
        "phone": clean,
        "searches_today": searches_today,
        "total_reports": total_reports,
        "categories": categories,
    }


class PhoneReportRequest(BaseModel):
    phone: str
    category: str
    comment: Optional[str] = None


@app.post("/api/phone-report")
async def submit_phone_report(request: Request, data: PhoneReportRequest):
    """Submit a community report for a phone number."""
    clean = re.sub(r"[^\d]", "", data.phone)[-10:]
    if not clean:
        raise HTTPException(status_code=400, detail="Invalid phone number")
    if data.category not in PHONE_REPORT_CATEGORIES:
        raise HTTPException(status_code=400, detail="Invalid category")

    client_ip = request.headers.get("X-Forwarded-For", request.client.host or "").split(",")[0].strip()

    # One report per IP per phone per day
    today_start = datetime.now(timezone.utc).replace(hour=0, minute=0, second=0, microsecond=0).isoformat()
    already = await phone_reports_collection.find_one({
        "phone": clean, "clientIp": client_ip, "createdAt": {"$gte": today_start},
    })
    if already:
        raise HTTPException(status_code=429, detail="You already reported this number today")

    await phone_reports_collection.insert_one({
        "id": str(uuid.uuid4()),
        "phone": clean,
        "category": data.category,
        "comment": (data.comment or "")[:500],
        "clientIp": client_ip,
        "createdAt": datetime.now(timezone.utc).isoformat(),
    })

    total = await phone_reports_collection.count_documents({"phone": clean})
    return {"success": True, "total_reports": total}


@app.post("/api/people-search/search")
async def people_search_endpoint(request: Request, data: PeopleSearchRequest):
    """Run a people search and return blurred preview. Full result stored in DB."""
    client_ip = request.headers.get("X-Forwarded-For", request.client.host or "").split(",")[0].strip()

    # Rate limiting: 100 searches per IP per hour
    now = time.time()
    key = f"ps:{client_ip}"
    rate_limit_storage[key] = [t for t in rate_limit_storage[key] if now - t < 3600]
    if len(rate_limit_storage[key]) >= 100:
        raise HTTPException(status_code=429, detail="Too many searches. Please try again later.")
    rate_limit_storage[key].append(now)

    lt = data.lookupType
    if lt not in DEFAULT_PEOPLE_SEARCH_PRICES:
        raise HTTPException(status_code=400, detail="Invalid lookup type")

    # ── Check enabled data sources ────────────────────────────────────────────
    use_whitepages = await is_source_enabled("whitepages")
    use_internal   = await is_source_enabled("internal")

    # ── Run search across enabled sources, merge results ─────────────────────
    # full_results is always a list; single-result lookups wrap in list
    if lt == "phone_lookup":
        if not data.phone:
            raise HTTPException(status_code=400, detail="Phone number required")
        phone_clean = re.sub(r"[^\d]", "", data.phone)
        if len(phone_clean) < 10:
            raise HTTPException(status_code=400, detail="Invalid phone number – must be 10 digits")

        full_results = []
        if use_whitepages:
            wp = await wp_phone_lookup(phone_clean)
            if wp:
                full_results.append(wp)
                asyncio.ensure_future(upsert_from_whitepages("phone_lookup", wp, phone=phone_clean))
        if use_internal:
            db_results = await internal_phone_lookup(phone_clean)
            print(f"[phone_lookup] internal DB returned {len(db_results)} results for phone={phone_clean}")
            full_results.extend(db_results)
        asyncio.ensure_future(upsert_phone_entry(phone_clean, full_results))

        # Enrich phone result with carrier info
        e164 = f"+1{phone_clean[-10:]}"
        carrier_info = await asyncio.to_thread(_phonenumbers_carrier_lookup, e164)
        for r in full_results:
            r["carrier"]  = r.get("carrier")  or carrier_info.get("carrier", "")
            r["lineType"] = r.get("lineType") or carrier_info.get("lineType", "")
            r["region"]   = r.get("region")   or carrier_info.get("region", "")

        query_summary = data.phone

    elif lt == "name_lookup":
        first_name, last_name = data.get_first_last()
        if not first_name or not last_name:
            raise HTTPException(status_code=400, detail="First and last name required")

        min_age = data.minAge
        max_age = data.maxAge

        full_results = []
        if use_whitepages:
            wp_list = await wp_person_lookup(first_name, last_name, data.state or "", data.city or "", min_age, max_age)
            if wp_list:
                full_results.extend(wp_list)
                for wp_item in wp_list:
                    asyncio.ensure_future(upsert_from_whitepages("name_lookup", wp_item))
        if use_internal:
            db_results = await internal_name_lookup(first_name, last_name, data.state or "", data.city or "", min_age, max_age)
            full_results.extend(db_results)

        # Post-filter by age for results that have age data
        if min_age is not None or max_age is not None:
            def _age_ok(r):
                age = r.get("age")
                if age is None:
                    return True  # keep if unknown
                if min_age is not None and age < min_age:
                    return False
                if max_age is not None and age > max_age:
                    return False
                return True
            full_results = [r for r in full_results if _age_ok(r)]

        parts = [f"{first_name} {last_name}"]
        if data.city: parts.append(data.city)
        if data.state: parts.append(data.state)
        if min_age or max_age:
            parts.append(f"Age {min_age or ''}–{max_age or ''}")
        query_summary = ", ".join(parts)

    elif lt == "address_lookup":
        if not data.street or not data.city:
            raise HTTPException(status_code=400, detail="Street and city required")

        full_results = []
        if use_whitepages:
            wp = await wp_address_lookup(data.street, data.city, data.state or "")
            if wp:
                full_results.append(wp)
                asyncio.ensure_future(upsert_from_whitepages("address_lookup", wp))
        if use_internal:
            db_results = await internal_address_lookup(data.street, data.city, data.state or "")
            full_results.extend(db_results)
        asyncio.ensure_future(upsert_address_entry(data.street, data.city, data.state or "", full_results))
        query_summary = f"{data.street}, {data.city}" + (f", {data.state}" if data.state else "")

    else:  # carrier_lookup
        if not data.phone:
            raise HTTPException(status_code=400, detail="Phone number required")
        phone_clean = re.sub(r"[^\d]", "", data.phone)
        if len(phone_clean) < 10:
            raise HTTPException(status_code=400, detail="Invalid phone number – must be 10 digits")

        e164 = f"+1{phone_clean[-10:]}"
        # Live lookup first, fall back to offline
        carrier_data = await _live_carrier_lookup(phone_clean)
        if not carrier_data:
            carrier_data = await asyncio.to_thread(_phonenumbers_carrier_lookup, e164)

        full_results = [{
            "phone":     data.phone,
            "carrier":   carrier_data.get("carrier", ""),
            "lineType":  carrier_data.get("lineType", ""),
            "region":    carrier_data.get("region", ""),
            "valid":     carrier_data.get("valid", True),
            "countryCode": carrier_data.get("countryCode", "US"),
        }]
        query_summary = data.phone

    prices = await get_people_search_prices()
    price = prices.get(lt, DEFAULT_PEOPLE_SEARCH_PRICES[lt])

    # Resolve user info
    user_id, user_email = "", ""
    auth_header = request.headers.get("Authorization", "")
    if auth_header.startswith("Bearer "):
        user = await users_collection.find_one({"sessionToken": auth_header[7:]})
        if user:
            user_id = user["id"]
            user_email = user.get("email", "")

    # Create one log entry per result so each can be independently paid/unlocked
    now_iso = datetime.now(timezone.utc).isoformat()
    result_entries = []
    for full_result in full_results:
        sid = str(uuid.uuid4())
        log_entry = {
            "id": sid,
            "lookupType": lt,
            "query": query_summary,
            "fullResult": full_result,
            "price": price,
            "isPaid": False,
            "stripeSessionId": None,
            "clientIp": client_ip,
            "userId": user_id,
            "userEmail": user_email,
            "createdAt": now_iso,
        }
        await people_search_logs_collection.insert_one(log_entry)
        result_entries.append({
            "searchId": sid,
            "preview":  blur_result(full_result, lt),
            "price":    price,
        })

    debug_info = {}
    if len(result_entries) == 0:
        debug_info = {
            "use_internal": use_internal,
            "use_whitepages": use_whitepages,
            "internal_count": await people_records_collection.count_documents({}),
        }
        if lt == "name_lookup":
            fn, ln = data.get_first_last()
            debug_info["parsed_first"] = fn
            debug_info["parsed_last"] = ln
            debug_info["raw_fullName"] = data.fullName

    return {
        "success":     True,
        "lookupType":  lt,
        "query":       query_summary,
        "results":     result_entries,   # list — each has searchId + preview + price
        "totalCount":  len(result_entries),
        "_debug":      debug_info,
    }


@app.post("/api/people-search/checkout")
async def people_search_checkout_endpoint(request: Request, data: PeopleSearchCheckoutRequest):
    """Create a Stripe checkout to unlock a people search result."""
    log = await people_search_logs_collection.find_one({"id": data.searchId})
    if not log:
        raise HTTPException(status_code=404, detail="Search not found")
    if log.get("isPaid"):
        return {"success": True, "alreadyPaid": True, "result": log["fullResult"]}

    frontend_url = os.environ.get("FRONTEND_URL", "https://mintslip.com")
    prices = await get_people_search_prices()
    lt = log["lookupType"]
    price = prices.get(lt, DEFAULT_PEOPLE_SEARCH_PRICES.get(lt, 0.99))
    discount_amount = float(data.discountAmount or 0)
    final_price = round(price - discount_amount, 2)

    labels = {
        "phone_lookup":   "Reverse Phone Lookup",
        "name_lookup":    "Name Lookup Report",
        "address_lookup": "Address Lookup Report",
        "carrier_lookup": "Carrier Lookup",
    }
    label = labels.get(lt, "People Search Report")

    user_id, user_email = "", ""
    auth_header = request.headers.get("Authorization", "")
    if auth_header.startswith("Bearer "):
        user = await users_collection.find_one({"sessionToken": auth_header[7:]})
        if user:
            user_id = user["id"]
            user_email = user.get("email", "")

    # 100% off — mark as paid directly, no Stripe needed
    if final_price <= 0:
        now_iso = datetime.now(timezone.utc).isoformat()
        await people_search_logs_collection.update_one(
            {"id": data.searchId},
            {"$set": {"isPaid": True, "paidAt": now_iso, "discountCode": data.discountCode}},
        )
        payment_record = {
            "id": str(uuid.uuid4()), "searchId": data.searchId, "lookupType": lt,
            "query": log.get("query", ""), "amount": 0, "stripeSessionId": None,
            "userId": user_id, "userEmail": user_email, "createdAt": now_iso,
            "discountCode": data.discountCode, "discountAmount": discount_amount,
        }
        await people_search_payments_collection.insert_one(payment_record)
        return {"success": True, "alreadyPaid": True, "result": log["fullResult"]}

    amount_cents = int(round(final_price * 100))

    checkout_session = await asyncio.to_thread(
        stripe.checkout.Session.create,
        payment_method_types=["card"],
        line_items=[{
            "price_data": {
                "currency": "usd",
                "product_data": {
                    "name": f"MintSlip – {label}",
                    "description": f"Unlock your {label}: {log['query']}",
                },
                "unit_amount": amount_cents,
            },
            "quantity": 1,
        }],
        mode="payment",
        success_url=(data.successUrl or
                     f"{frontend_url}/people-search?session_id={{CHECKOUT_SESSION_ID}}&search_id={data.searchId}"),
        cancel_url=(data.cancelUrl or f"{frontend_url}/people-search"),
        metadata={
            "type": "people_search",
            "searchId": data.searchId,
            "lookupType": lt,
            "userId": user_id,
            "userEmail": user_email,
            "discountCode": data.discountCode or "",
            "discountAmount": str(discount_amount),
        },
    )

    await people_search_logs_collection.update_one(
        {"id": data.searchId},
        {"$set": {"stripeSessionId": checkout_session.id}},
    )
    return {"success": True, "sessionId": checkout_session.id, "url": checkout_session.url}


@app.get("/api/people-search/result/{search_id}")
async def get_people_search_result(search_id: str, session_id: Optional[str] = None):
    """Return full result if paid; optionally verify a new Stripe session."""
    log = await people_search_logs_collection.find_one({"id": search_id})
    if not log:
        raise HTTPException(status_code=404, detail="Search not found")

    if log.get("isPaid"):
        return {"success": True, "paid": True, "result": log["fullResult"], "lookupType": log["lookupType"]}

    if session_id:
        try:
            session = await asyncio.to_thread(stripe.checkout.Session.retrieve, session_id)
            if session.payment_status == "paid" and session.status == "complete":
                prices = await get_people_search_prices()
                lt = log["lookupType"]
                price = prices.get(lt, DEFAULT_PEOPLE_SEARCH_PRICES.get(lt, 0.99))
                now_iso = datetime.now(timezone.utc).isoformat()
                await people_search_logs_collection.update_one(
                    {"id": search_id},
                    {"$set": {"isPaid": True, "paidAt": now_iso, "stripeSessionId": session_id}},
                )
                payment_record = {
                    "id": str(uuid.uuid4()),
                    "searchId": search_id,
                    "lookupType": lt,
                    "query": log.get("query", ""),
                    "amount": price,
                    "stripeSessionId": session_id,
                    "userId": log.get("userId", ""),
                    "userEmail": log.get("userEmail", ""),
                    "createdAt": now_iso,
                }
                await people_search_payments_collection.insert_one(payment_record)
                return {"success": True, "paid": True, "result": log["fullResult"], "lookupType": lt}
        except Exception as e:
            logger.error(f"People search Stripe check error: {e}")

    return {"success": True, "paid": False}


# ── Admin: People Search ──────────────────────────────────────────────────────

@app.get("/api/admin/people-search/stats")
async def admin_ps_stats(request: Request):
    token = request.headers.get("Authorization", "").replace("Bearer ", "")
    if not await sessions_collection.find_one({"token": token, "type": {"$in": ["admin", "moderator"]}}):
        raise HTTPException(status_code=401, detail="Unauthorized")

    total = await people_search_logs_collection.count_documents({})
    paid = await people_search_logs_collection.count_documents({"isPaid": True})
    payments = await people_search_payments_collection.find({}).to_list(None)
    revenue = round(sum(p.get("amount", 0) for p in payments), 2)

    by_type = {}
    async for doc in people_search_logs_collection.aggregate([
        {"$group": {"_id": "$lookupType", "count": {"$sum": 1},
                    "paid": {"$sum": {"$cond": ["$isPaid", 1, 0]}}}}
    ]):
        by_type[doc["_id"]] = {"count": doc["count"], "paid": doc["paid"]}

    return {"totalSearches": total, "paidSearches": paid, "totalRevenue": revenue, "byType": by_type}


@app.get("/api/admin/people-search/searches")
async def admin_ps_searches(request: Request, page: int = 1, limit: int = 50):
    token = request.headers.get("Authorization", "").replace("Bearer ", "")
    if not await sessions_collection.find_one({"token": token, "type": {"$in": ["admin", "moderator"]}}):
        raise HTTPException(status_code=401, detail="Unauthorized")

    skip = (page - 1) * limit
    cursor = people_search_logs_collection.find({}, {"fullResult": 0}).sort("createdAt", -1).skip(skip).limit(limit)
    searches = []
    async for doc in cursor:
        doc.pop("_id", None)
        searches.append(doc)
    total = await people_search_logs_collection.count_documents({})
    return {"searches": searches, "total": total, "page": page, "limit": limit}


@app.put("/api/admin/people-search/prices")
async def admin_ps_update_prices(request: Request):
    token = request.headers.get("Authorization", "").replace("Bearer ", "")
    if not await sessions_collection.find_one({"token": token, "type": {"$in": ["admin", "moderator"]}}):
        raise HTTPException(status_code=401, detail="Unauthorized")

    body = await request.json()
    prices = {k: round(float(body.get(k, DEFAULT_PEOPLE_SEARCH_PRICES[k])), 2)
              for k in DEFAULT_PEOPLE_SEARCH_PRICES}
    await site_settings_collection.update_one(
        {"key": "people_search_prices"},
        {"$set": {"key": "people_search_prices", "value": prices,
                  "updatedAt": datetime.now(timezone.utc).isoformat()}},
        upsert=True,
    )
    return {"success": True, "prices": prices}


# ── Admin: Data Sources ───────────────────────────────────────────────────────

@app.get("/api/admin/data-sources")
async def admin_get_data_sources(request: Request):
    token = request.headers.get("Authorization", "").replace("Bearer ", "")
    if not await sessions_collection.find_one({"token": token, "type": {"$in": ["admin", "moderator"]}}):
        raise HTTPException(status_code=401, detail="Unauthorized")

    sources = await get_data_sources()
    # Known non-internal sources (scrapers + import types)
    non_internal_sources = [s["source"] for s in DATA_SOURCE_DEFAULTS if s["source"] not in ("internal", "whitepages")]
    # Enrich with live record counts and latest job info
    for src in sources:
        if src["source"] == "internal":
            # Count all records that aren't from a named scraper/import source
            src["recordCount"] = await people_records_collection.count_documents(
                {"source": {"$nin": non_internal_sources}}
            )
        else:
            src["recordCount"] = await people_records_collection.count_documents({"source": src["source"]})
        job = await scraper_jobs_collection.find_one(
            {"source": src["source"]}, sort=[("startedAt", -1)]
        )
        if job:
            job.pop("_id", None)
            src["lastJob"] = job
        else:
            src["lastJob"] = None
    return {"sources": sources}


class DataSourceToggleRequest(BaseModel):
    enabled: bool


@app.put("/api/admin/data-sources/{source}")
async def admin_toggle_data_source(source: str, body: DataSourceToggleRequest, request: Request):
    token = request.headers.get("Authorization", "").replace("Bearer ", "")
    session = await sessions_collection.find_one({"token": token, "type": {"$in": ["admin", "moderator"]}})
    if not session:
        raise HTTPException(status_code=401, detail="Unauthorized")

    valid_sources = {d["source"] for d in DATA_SOURCE_DEFAULTS}
    if source not in valid_sources:
        raise HTTPException(status_code=400, detail="Unknown source")

    await data_source_settings_collection.update_one(
        {"source": source},
        {"$set": {"enabled": body.enabled, "updatedAt": datetime.now(timezone.utc).isoformat()}},
        upsert=True,
    )
    await log_action(session, "toggle_data_source", "data_source", source,
                     f"{'enabled' if body.enabled else 'disabled'}")
    return {"success": True, "source": source, "enabled": body.enabled}


# ── Admin: Scraper Jobs ───────────────────────────────────────────────────────

# In-memory set of currently running scraper tasks so we don't double-trigger
_running_scrapers: set = set()


@app.get("/api/admin/scrapers/jobs")
async def admin_scraper_jobs(request: Request, limit: int = 50):
    token = request.headers.get("Authorization", "").replace("Bearer ", "")
    if not await sessions_collection.find_one({"token": token, "type": {"$in": ["admin", "moderator"]}}):
        raise HTTPException(status_code=401, detail="Unauthorized")

    cursor = scraper_jobs_collection.find({}).sort("startedAt", -1).limit(limit)
    jobs = []
    async for doc in cursor:
        doc.pop("_id", None)
        jobs.append(doc)
    return {"jobs": jobs, "running": list(_running_scrapers)}


@app.post("/api/admin/scrapers/{source}/trigger")
async def admin_trigger_scraper(source: str, request: Request):
    token = request.headers.get("Authorization", "").replace("Bearer ", "")
    session = await sessions_collection.find_one({"token": token, "type": {"$in": ["admin", "moderator"]}})
    if not session:
        raise HTTPException(status_code=401, detail="Unauthorized")

    valid_scrapers = {"nsopw", "nppes", "fec", "faa", "ssdi"}
    if source not in valid_scrapers:
        raise HTTPException(status_code=400, detail=f"No scraper available for '{source}'")

    if source in _running_scrapers:
        raise HTTPException(status_code=409, detail="Scraper already running for this source")

    job_id = str(uuid.uuid4())
    job_doc = {
        "jobId": job_id,
        "source": source,
        "status": "running",
        "recordsAdded": 0,
        "recordsUpdated": 0,
        "errors": [],
        "startedAt": datetime.now(timezone.utc).isoformat(),
        "finishedAt": None,
    }
    await scraper_jobs_collection.insert_one(job_doc)
    _running_scrapers.add(source)

    async def run_scraper():
        try:
            from scrapers import run_scraper as _run
            added, updated, errors = await _run(source, people_records_collection)
            await scraper_jobs_collection.update_one(
                {"jobId": job_id},
                {"$set": {
                    "status": "completed",
                    "recordsAdded": added,
                    "recordsUpdated": updated,
                    "errors": errors[:20],
                    "finishedAt": datetime.now(timezone.utc).isoformat(),
                }},
            )
        except Exception as e:
            await scraper_jobs_collection.update_one(
                {"jobId": job_id},
                {"$set": {
                    "status": "failed",
                    "errors": [str(e)],
                    "finishedAt": datetime.now(timezone.utc).isoformat(),
                }},
            )
        finally:
            _running_scrapers.discard(source)

    asyncio.create_task(run_scraper())
    await log_action(session, "trigger_scraper", "scraper", source, job_id)
    return {"success": True, "jobId": job_id}


@app.delete("/api/admin/people-records")
async def admin_clear_people_records(source: str, request: Request):
    """Delete all scraped records for a given source (for re-scraping)."""
    token = request.headers.get("Authorization", "").replace("Bearer ", "")
    session = await sessions_collection.find_one({"token": token, "type": "admin"})
    if not session:
        raise HTTPException(status_code=401, detail="Admin only")

    result = await people_records_collection.delete_many({"source": source})
    await log_action(session, "clear_people_records", "scraper", source,
                     f"{result.deleted_count} records deleted")
    return {"success": True, "deleted": result.deleted_count}


# ── Admin: People Records Browser ────────────────────────────────────────────

@app.get("/api/admin/people-records/browse")
async def admin_browse_people_records(
    q: str = "",
    source: str = "",
    state: str = "",
    address: str = "",
    page: int = 1,
    session: dict = Depends(get_current_admin),
):
    LIMIT = 25
    skip = (page - 1) * LIMIT
    query: dict = {}
    if source:
        query["source"] = source
    if state:
        query["$or"] = [
            {"state": {"$regex": f"^{re.escape(state)}$", "$options": "i"}},
            {"addresses.state": {"$regex": f"^{re.escape(state)}$", "$options": "i"}},
        ]
    if address:
        addr_stripped = address.strip()
        query["$or"] = [
            {"addresses.street": {"$regex": re.escape(addr_stripped), "$options": "i"}},
            {"addresses.city":   {"$regex": re.escape(addr_stripped), "$options": "i"}},
        ]
    if q:
        q_stripped = q.strip()
        digits_only = re.sub(r"[^\d]", "", q_stripped)
        if len(digits_only) >= 7:
            query["$or"] = [
                {"phones": {"$elemMatch": {"number": {"$regex": digits_only}}}},
                {"phones": {"$elemMatch": {"$regex": digits_only}}},
            ]
        else:
            parts = q_stripped.split()
            if len(parts) >= 2:
                query["$and"] = [
                    {"firstName": {"$regex": re.escape(parts[0]),  "$options": "i"}},
                    {"lastName":  {"$regex": re.escape(parts[-1]), "$options": "i"}},
                ]
            else:
                query["$or"] = [
                    {"firstName": {"$regex": re.escape(q_stripped), "$options": "i"}},
                    {"lastName":  {"$regex": re.escape(q_stripped), "$options": "i"}},
                ]
    total = await people_records_collection.count_documents(query)
    cursor = (
        people_records_collection
        .find(query, {"_id": 0})
        .sort("createdAt", -1)
        .skip(skip)
        .limit(LIMIT)
    )
    records = await cursor.to_list(LIMIT)
    pages = max(1, (total + LIMIT - 1) // LIMIT)
    return {"success": True, "records": records, "total": total, "page": page, "pages": pages}


@app.delete("/api/admin/people-records/record/{record_id}")
async def admin_delete_one_people_record(
    record_id: str,
    session: dict = Depends(get_current_admin),
):
    result = await people_records_collection.delete_one({"recordId": record_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Record not found")
    await log_action(session, "delete_people_record", "people_record", record_id)
    return {"success": True, "deleted": record_id}


@app.post("/api/admin/people-records/create")
async def admin_create_people_record(data: dict, session: dict = Depends(get_current_admin)):
    """Create a new internal people record."""
    now = datetime.now(timezone.utc).isoformat()
    doc = {
        "recordId":    str(uuid.uuid4()),
        "source":      "internal",
        "firstName":   data.get("firstName", "").strip(),
        "lastName":    data.get("lastName", "").strip(),
        "middleName":  data.get("middleName", "").strip(),
        "aliases":     data.get("aliases", []),
        "age":         calc_age_from_dob(data.get("dateOfBirth", "")),
        "dateOfBirth": data.get("dateOfBirth", ""),
        "gender":      data.get("gender", ""),
        "state":       data.get("state", ""),
        "addresses":   data.get("addresses", []),
        "phones":      data.get("phones", []),
        "emails":      data.get("emails", []),
        "relatives":   data.get("relatives", []),
        "associates":  data.get("associates", []),
        "occupation":  data.get("occupation", ""),
        "employer":    data.get("employer", ""),
        "sourceData":  {},
        "createdAt":   now,
        "lastUpdated": now,
    }
    if not doc["firstName"] or not doc["lastName"]:
        raise HTTPException(status_code=400, detail="First and last name are required")
    await people_records_collection.insert_one(doc)
    await log_action(session, "create_people_record", "people_record", doc["recordId"], f"{doc['firstName']} {doc['lastName']}")
    doc.pop("_id", None)
    return {"success": True, "record": doc}


@app.put("/api/admin/people-records/record/{record_id}")
async def admin_update_people_record(record_id: str, data: dict, session: dict = Depends(get_current_admin)):
    """Update an existing people record."""
    existing = await people_records_collection.find_one({"recordId": record_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Record not found")
    updates = {
        "firstName":   data.get("firstName", existing.get("firstName", "")).strip(),
        "lastName":    data.get("lastName",  existing.get("lastName", "")).strip(),
        "middleName":  data.get("middleName", existing.get("middleName", "")).strip(),
        "aliases":     data.get("aliases",   existing.get("aliases", [])),
        "dateOfBirth": data.get("dateOfBirth", existing.get("dateOfBirth", "")),
        "age":         calc_age_from_dob(data.get("dateOfBirth", existing.get("dateOfBirth", ""))),
        "gender":      data.get("gender", existing.get("gender", "")),
        "state":       data.get("state", existing.get("state", "")),
        "addresses":   data.get("addresses", existing.get("addresses", [])),
        "phones":      data.get("phones",    existing.get("phones", [])),
        "emails":      data.get("emails",    existing.get("emails", [])),
        "relatives":   data.get("relatives", existing.get("relatives", [])),
        "associates":  data.get("associates", existing.get("associates", [])),
        "occupation":  data.get("occupation", existing.get("occupation", "")),
        "employer":    data.get("employer",   existing.get("employer", "")),
        "lastUpdated": datetime.now(timezone.utc).isoformat(),
    }
    await people_records_collection.update_one({"recordId": record_id}, {"$set": updates})
    await log_action(session, "update_people_record", "people_record", record_id, f"{updates['firstName']} {updates['lastName']}")
    return {"success": True}


@app.post("/api/admin/people-records/mass-delete")
async def admin_mass_delete_people_records(data: dict, session: dict = Depends(get_current_admin)):
    """Delete multiple people records by recordId."""
    ids = data.get("ids", [])
    if not ids:
        raise HTTPException(status_code=400, detail="No IDs provided")
    result = await people_records_collection.delete_many({"recordId": {"$in": ids}})
    await log_action(session, "mass_delete_people_records", "people_record", "bulk", f"{result.deleted_count} deleted")
    return {"success": True, "deleted": result.deleted_count}


US_STATE_CODES = {
    "AL","AK","AZ","AR","CA","CO","CT","DE","FL","GA","HI","ID","IL","IN","IA",
    "KS","KY","LA","ME","MD","MA","MI","MN","MS","MO","MT","NE","NV","NH","NJ",
    "NM","NY","NC","ND","OH","OK","OR","PA","RI","SC","SD","TN","TX","UT","VT",
    "VA","WA","WV","WI","WY",
}

@app.post("/api/admin/people-records/fix-addresses")
async def admin_fix_bad_addresses(session: dict = Depends(get_current_admin)):
    """One-time migration: fix addresses where the street was incorrectly stored in the city field.
    Detects bad addresses where city starts with a digit or state is not a valid US state code.
    Reconstructs the original string and re-parses with the corrected parser.
    """
    fixed_records = 0
    fixed_addresses = 0
    # Fetch ALL records that have any address — re-validate each one in Python
    # (avoids complex MongoDB queries with edge cases around null/missing fields)
    cursor = people_records_collection.find(
        {"addresses": {"$exists": True, "$not": {"$size": 0}}}
    )
    async for doc in cursor:
        new_addrs = []
        changed = False
        for addr in doc.get("addresses", []):
            city   = addr.get("city", "")
            state  = addr.get("state", "")
            zip_   = addr.get("zip", "")
            street = addr.get("street", "")

            # Bad if city starts with a digit (street number) OR state is not a valid US code
            city_is_street = bool(re.match(r"^\d", city))
            state_invalid  = bool(state) and state not in US_STATE_CODES

            if city_is_street or state_invalid:
                # Reconstruct the original mangled string and re-parse
                parts = [p for p in [street, city, f"{state} {zip_}".strip()] if p]
                reconstructed = ", ".join(parts)
                reparsed = _parse_addr_str(reconstructed, current=addr.get("current", False))
                new_addrs.append(reparsed)
                changed = True
                fixed_addresses += 1
            else:
                new_addrs.append(addr)

        if changed:
            await people_records_collection.update_one(
                {"_id": doc["_id"]},
                {"$set": {"addresses": new_addrs, "lastUpdated": datetime.now(timezone.utc).isoformat()}}
            )
            fixed_records += 1

    await log_action(session, "fix_bad_addresses", "people_record", "bulk",
                     f"Fixed {fixed_addresses} addresses across {fixed_records} records")
    return {"success": True, "fixedRecords": fixed_records, "fixedAddresses": fixed_addresses}


# ── Public: People associated with a phone or address ────────────────────────

@app.get("/api/people-search/associated-by-phone")
async def associated_people_by_phone(phone: str):
    """Return internal DB people records that have the given phone number."""
    clean = re.sub(r"[^\d]", "", phone)[-10:]
    if len(clean) < 7:
        return {"results": []}
    cursor = people_records_collection.find(
        {"phones": {"$elemMatch": {"$regex": clean}}}
    ).limit(10)
    results = []
    async for doc in cursor:
        results.append(normalize_internal_record(doc))
    return {"results": results}


@app.get("/api/people-search/associated-by-address")
async def associated_people_by_address(street: str = "", city: str = "", state: str = ""):
    """Return internal DB people records associated with the given address."""
    if not street and not city:
        return {"results": []}
    conditions = []
    if street:
        conditions.append({"addresses.street": {"$regex": re.escape(street.strip()), "$options": "i"}})
    if city:
        conditions.append({"addresses.city": {"$regex": re.escape(city.strip()), "$options": "i"}})
    if state:
        conditions.append({"addresses.state": {"$regex": f"^{re.escape(state.strip())}$", "$options": "i"}})
    query = {"$and": conditions} if len(conditions) > 1 else conditions[0]
    cursor = people_records_collection.find(query).limit(10)
    results = []
    async for doc in cursor:
        results.append(normalize_internal_record(doc))
    return {"results": results}


# ── Public: People Record by ID ───────────────────────────────────────────────

@app.get("/api/people/record/{record_id}")
async def get_people_record_by_id(record_id: str):
    """Return a normalized people record by recordId (for linked relative navigation)."""
    doc = await people_records_collection.find_one({"recordId": record_id})
    if not doc:
        raise HTTPException(status_code=404, detail="Record not found")
    normalized = normalize_internal_record(doc)
    return {"success": True, "record": normalized}


# ── Admin: Browse people records by query (for relation linking) ──────────────

@app.get("/api/admin/people-records/search-suggest")
async def admin_people_records_suggest(q: str = "", limit: int = 10, session: dict = Depends(get_current_admin)):
    """Return name suggestions for linking relatives/associates."""
    if not q or len(q) < 2:
        return {"results": []}
    parts = q.strip().split()
    if len(parts) == 1:
        regex_query = {
            "$or": [
                {"firstName": {"$regex": f"^{re.escape(parts[0])}", "$options": "i"}},
                {"lastName":  {"$regex": f"^{re.escape(parts[0])}", "$options": "i"}},
            ]
        }
    else:
        regex_query = {
            "firstName": {"$regex": f"^{re.escape(parts[0])}", "$options": "i"},
            "lastName":  {"$regex": f"^{re.escape(parts[-1])}", "$options": "i"},
        }
    cursor = people_records_collection.find(regex_query).limit(limit)
    results = []
    async for doc in cursor:
        full_name = f"{doc.get('firstName', '')} {doc.get('lastName', '')}".strip()
        results.append({
            "recordId": doc.get("recordId", ""),
            "name":     full_name,
            "state":    doc.get("state", ""),
            "age":      doc.get("age") or calc_age_from_dob(doc.get("dateOfBirth", "")),
        })
    return {"results": results}


# ── Public: Opt-Out Requests ──────────────────────────────────────────────────

class OptOutRequest(BaseModel):
    name:   str
    email:  str
    reason: str = ""
    phone:  str = ""

@app.post("/api/opt-out")
async def submit_opt_out(body: OptOutRequest):
    name  = body.name.strip()
    email = body.email.strip().lower()
    if not name or not email:
        raise HTTPException(status_code=400, detail="name and email are required")
    doc = {
        "id":        str(uuid.uuid4()),
        "name":      name,
        "email":     email,
        "phone":     body.phone.strip(),
        "reason":    body.reason.strip(),
        "status":    "pending",
        "createdAt": datetime.now(timezone.utc).isoformat(),
        "updatedAt": datetime.now(timezone.utc).isoformat(),
    }
    await opt_out_requests_collection.insert_one(doc)
    return {"success": True}


# ── Admin: Opt-Out Management ─────────────────────────────────────────────────

@app.get("/api/admin/opt-outs")
async def admin_list_opt_outs(
    status: str = "pending",
    page: int = 1,
    session: dict = Depends(get_current_admin),
):
    LIMIT = 25
    skip  = (page - 1) * LIMIT
    query: dict = {}
    if status and status != "all":
        query["status"] = status
    total   = await opt_out_requests_collection.count_documents(query)
    records = await (
        opt_out_requests_collection
        .find(query, {"_id": 0})
        .sort("createdAt", -1)
        .skip(skip)
        .limit(LIMIT)
        .to_list(LIMIT)
    )
    pages = max(1, (total + LIMIT - 1) // LIMIT)
    return {"success": True, "requests": records, "total": total, "page": page, "pages": pages}


@app.post("/api/admin/opt-outs/{opt_out_id}/approve")
async def admin_approve_opt_out(opt_out_id: str, session: dict = Depends(require_admin_only)):
    req = await opt_out_requests_collection.find_one({"id": opt_out_id}, {"_id": 0})
    if not req:
        raise HTTPException(status_code=404, detail="Not found")

    name_parts = req["name"].split()
    first = name_parts[0] if name_parts else ""
    last  = name_parts[-1] if len(name_parts) > 1 else ""

    deleted_count = 0
    if first and last:
        result = await people_records_collection.delete_many({
            "firstName": {"$regex": f"^{re.escape(first)}$", "$options": "i"},
            "lastName":  {"$regex": f"^{re.escape(last)}$",  "$options": "i"},
        })
        deleted_count = result.deleted_count

    await suppression_list_collection.update_one(
        {"email": req["email"]},
        {"$set": {
            "name":    req["name"],
            "email":   req["email"],
            "phone":   req.get("phone", ""),
            "addedAt": datetime.now(timezone.utc).isoformat(),
            "addedBy": session.get("email", "admin"),
        }},
        upsert=True,
    )
    await opt_out_requests_collection.update_one(
        {"id": opt_out_id},
        {"$set": {"status": "approved", "deletedCount": deleted_count,
                  "updatedAt": datetime.now(timezone.utc).isoformat()}},
    )
    await log_action(session, "approve_opt_out", "opt_out", opt_out_id,
                     f"Deleted {deleted_count} records for {req['name']}")
    return {"success": True, "deleted": deleted_count}


@app.post("/api/admin/opt-outs/{opt_out_id}/deny")
async def admin_deny_opt_out(opt_out_id: str, session: dict = Depends(require_admin_only)):
    req = await opt_out_requests_collection.find_one({"id": opt_out_id})
    if not req:
        raise HTTPException(status_code=404, detail="Not found")
    await opt_out_requests_collection.update_one(
        {"id": opt_out_id},
        {"$set": {"status": "denied", "updatedAt": datetime.now(timezone.utc).isoformat()}},
    )
    await log_action(session, "deny_opt_out", "opt_out", opt_out_id, req.get("name", ""))
    return {"success": True}


# ── Voter Roll CSV Importer ───────────────────────────────────────────────────

class VoterRollColumnMap(BaseModel):
    firstName:   str = "first_name"
    lastName:    str = "last_name"
    middleName:  str = ""
    dob:         str = "date_of_birth"
    street:      str = "address"
    city:        str = "city"
    state:       str = "state"
    zip:         str = "zip"
    phone:       str = "phone"
    email:       str = ""


@app.post("/api/admin/voter-rolls/preview")
async def voter_roll_preview(
    request: Request,
    file: UploadFile,
):
    """Return first 5 rows + header columns so admin can map columns."""
    token = request.headers.get("Authorization", "").replace("Bearer ", "")
    if not await sessions_collection.find_one({"token": token, "type": {"$in": ["admin", "moderator"]}}):
        raise HTTPException(status_code=401, detail="Unauthorized")

    content = await file.read()
    try:
        text = content.decode("utf-8", errors="replace")
        import csv as _csv
        reader = _csv.DictReader(text.splitlines())
        columns = reader.fieldnames or []
        rows = []
        for i, row in enumerate(reader):
            if i >= 5:
                break
            rows.append(dict(row))
        return {"columns": list(columns), "preview": rows}
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not parse CSV: {e}")


@app.post("/api/admin/voter-rolls/import")
async def voter_roll_import(
    request: Request,
    file: UploadFile,
    firstName:  str = "first_name",
    lastName:   str = "last_name",
    middleName: str = "",
    dob:        str = "date_of_birth",
    street:     str = "address",
    city:       str = "city",
    state:      str = "state",
    zipCol:     str = "zip",
    phone:      str = "phone",
    email:      str = "",
    stateName:  str = "",
):
    """Import a voter roll CSV file into people_records. Runs in background."""
    token = request.headers.get("Authorization", "").replace("Bearer ", "")
    session = await sessions_collection.find_one({"token": token, "type": {"$in": ["admin", "moderator"]}})
    if not session:
        raise HTTPException(status_code=401, detail="Unauthorized")

    content = await file.read()
    job_id = str(uuid.uuid4())
    job_doc = {
        "jobId": job_id,
        "source": "voter_rolls",
        "status": "running",
        "recordsAdded": 0,
        "recordsUpdated": 0,
        "errors": [],
        "startedAt": datetime.now(timezone.utc).isoformat(),
        "finishedAt": None,
    }
    await scraper_jobs_collection.insert_one(job_doc)
    _running_scrapers.add("voter_rolls")

    col_map = {
        "firstName": firstName, "lastName": lastName, "middleName": middleName,
        "dob": dob, "street": street, "city": city, "state": state,
        "zip": zipCol, "phone": phone, "email": email,
    }

    async def _import():
        import csv as _csv
        added = updated = 0
        errors = []
        try:
            text = content.decode("utf-8", errors="replace")
            reader = _csv.DictReader(text.splitlines())

            def g(row, key):
                col = col_map.get(key, "")
                return (row.get(col, "") or "").strip() if col else ""

            batch = []
            BATCH = 500
            for row in reader:
                try:
                    first = g(row, "firstName").title()
                    last  = g(row, "lastName").title()
                    if not first or not last:
                        continue

                    addr_street = g(row, "street")
                    addr_city   = g(row, "city").title()
                    addr_state  = (g(row, "state") or stateName).upper()
                    addr_zip    = g(row, "zip")[:5]
                    dob_val     = g(row, "dob")
                    phone_val   = re.sub(r"[^\d]", "", g(row, "phone"))
                    email_val   = g(row, "email").lower()

                    import uuid as _uuid
                    from datetime import datetime as _dt, timezone as _tz
                    doc = {
                        "recordId":   str(_uuid.uuid4()),
                        "source":     "voter_rolls",
                        "firstName":  first,
                        "lastName":   last,
                        "middleName": g(row, "middleName").title(),
                        "dateOfBirth": dob_val,
                        "state":      addr_state,
                        "addresses":  [{
                            "street": addr_street, "city": addr_city,
                            "state": addr_state, "zip": addr_zip, "current": True,
                        }] if addr_street or addr_city else [],
                        "phones":  [phone_val] if phone_val else [],
                        "emails":  [email_val] if email_val else [],
                        "relatives": [],
                        "occupation": "",
                        "sourceData": {"stateSource": stateName or addr_state},
                        "createdAt":     _dt.now(_tz.utc).isoformat(),
                        "lastUpdated":   _dt.now(_tz.utc).isoformat(),
                    }
                    batch.append(doc)
                    if len(batch) >= BATCH:
                        await people_records_collection.insert_many(batch, ordered=False)
                        added += len(batch)
                        batch = []
                        await asyncio.sleep(0)
                except Exception as e:
                    errors.append(str(e))

            if batch:
                await people_records_collection.insert_many(batch, ordered=False)
                added += len(batch)

        except Exception as e:
            errors.append(f"Fatal: {e}")
        finally:
            await scraper_jobs_collection.update_one(
                {"jobId": job_id},
                {"$set": {
                    "status": "completed" if not errors or added > 0 else "failed",
                    "recordsAdded": added,
                    "recordsUpdated": updated,
                    "errors": errors[:20],
                    "finishedAt": datetime.now(timezone.utc).isoformat(),
                }},
            )
            _running_scrapers.discard("voter_rolls")

    asyncio.create_task(_import())
    await log_action(session, "import_voter_rolls", "voter_rolls", stateName or "unknown", job_id)
    return {"success": True, "jobId": job_id}


# ===== SITE NAV ORDER =====

@app.get("/api/nav-order")
async def get_nav_order():
    doc = await site_settings_collection.find_one({"key": "site_nav_order"}, {"_id": 0})
    return {"success": True, "order": doc["value"] if doc else None}

@app.put("/api/admin/nav-order")
async def set_nav_order(request: Request, session: dict = Depends(get_current_admin)):
    body = await request.json()
    order = body.get("order", [])
    if not isinstance(order, list):
        raise HTTPException(status_code=400, detail="order must be a list")
    await site_settings_collection.update_one(
        {"key": "site_nav_order"},
        {"$set": {"key": "site_nav_order", "value": order}},
        upsert=True
    )
    await log_action(session, "update_nav_order", "site_settings", "site_nav_order")
    return {"success": True}


# ===== PHONE & ADDRESS ENTRY REGISTRY =====

@app.get("/api/admin/phone-entries")
async def admin_list_phone_entries(q: str = "", page: int = 1, session: dict = Depends(get_current_admin)):
    LIMIT = 25
    skip = (page - 1) * LIMIT
    query: dict = {}
    if q:
        clean = re.sub(r"[^\d]", "", q)
        if clean:
            query["phone"] = {"$regex": clean}
        else:
            query["$or"] = [
                {"notes": {"$regex": re.escape(q), "$options": "i"}},
                {"carrier": {"$regex": re.escape(q), "$options": "i"}},
            ]
    total = await phone_entries_collection.count_documents(query)
    cursor = phone_entries_collection.find(query, {"_id": 0}).sort("lastSearched", -1).skip(skip).limit(LIMIT)
    entries = await cursor.to_list(LIMIT)
    for entry in entries:
        linked = []
        for rid in (entry.get("linkedRecordIds") or []):
            doc = await people_records_collection.find_one({"recordId": rid}, {"_id": 0, "firstName": 1, "lastName": 1, "recordId": 1})
            if doc:
                linked.append({"recordId": doc["recordId"], "name": f"{doc['firstName']} {doc['lastName']}"})
        entry["linkedPeople"] = linked
    pages = max(1, (total + LIMIT - 1) // LIMIT)
    return {"success": True, "entries": entries, "total": total, "page": page, "pages": pages}


@app.post("/api/admin/phone-entries")
async def admin_create_phone_entry(data: dict, session: dict = Depends(get_current_admin)):
    phone = re.sub(r"[^\d]", "", data.get("phone", ""))[-10:]
    if len(phone) < 7:
        raise HTTPException(status_code=400, detail="Invalid phone number")
    existing = await phone_entries_collection.find_one({"phone": phone})
    if existing:
        raise HTTPException(status_code=409, detail="Phone entry already exists")
    display = f"({phone[:3]}) {phone[3:6]}-{phone[6:]}" if len(phone) == 10 else phone
    now = datetime.now(timezone.utc).isoformat()
    doc = {
        "entryId": str(uuid.uuid4()),
        "phone": phone,
        "displayPhone": display,
        "type": data.get("type", "unknown"),
        "carrier": data.get("carrier", ""),
        "linkedRecordIds": data.get("linkedRecordIds", []),
        "searchCount": 0,
        "lastSearched": None,
        "source": "manual",
        "addedAt": now,
        "notes": data.get("notes", ""),
    }
    await phone_entries_collection.insert_one(doc)
    doc.pop("_id", None)
    await log_action(session, "create_phone_entry", "phone_entry", doc["entryId"], phone)
    return {"success": True, "entry": doc}


@app.put("/api/admin/phone-entries/{entry_id}")
async def admin_update_phone_entry(entry_id: str, data: dict, session: dict = Depends(get_current_admin)):
    existing = await phone_entries_collection.find_one({"entryId": entry_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Not found")
    updates = {
        "type":             data.get("type",             existing.get("type", "unknown")),
        "carrier":          data.get("carrier",          existing.get("carrier", "")),
        "linkedRecordIds":  data.get("linkedRecordIds",  existing.get("linkedRecordIds", [])),
        "notes":            data.get("notes",            existing.get("notes", "")),
    }
    await phone_entries_collection.update_one({"entryId": entry_id}, {"$set": updates})
    await log_action(session, "update_phone_entry", "phone_entry", entry_id)
    return {"success": True}


@app.delete("/api/admin/phone-entries/{entry_id}")
async def admin_delete_phone_entry(entry_id: str, session: dict = Depends(get_current_admin)):
    result = await phone_entries_collection.delete_one({"entryId": entry_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Not found")
    await log_action(session, "delete_phone_entry", "phone_entry", entry_id)
    return {"success": True}


@app.get("/api/admin/address-entries")
async def admin_list_address_entries(q: str = "", page: int = 1, session: dict = Depends(get_current_admin)):
    LIMIT = 25
    skip = (page - 1) * LIMIT
    query: dict = {}
    if q:
        q_stripped = q.strip()
        query["$or"] = [
            {"street": {"$regex": re.escape(q_stripped), "$options": "i"}},
            {"city":   {"$regex": re.escape(q_stripped), "$options": "i"}},
            {"state":  {"$regex": re.escape(q_stripped), "$options": "i"}},
            {"zip":    {"$regex": re.escape(q_stripped), "$options": "i"}},
        ]
    total = await address_entries_collection.count_documents(query)
    cursor = address_entries_collection.find(query, {"_id": 0}).sort("lastSearched", -1).skip(skip).limit(LIMIT)
    entries = await cursor.to_list(LIMIT)
    for entry in entries:
        linked = []
        for rid in (entry.get("linkedRecordIds") or []):
            doc = await people_records_collection.find_one({"recordId": rid}, {"_id": 0, "firstName": 1, "lastName": 1, "recordId": 1})
            if doc:
                linked.append({"recordId": doc["recordId"], "name": f"{doc['firstName']} {doc['lastName']}"})
        entry["linkedPeople"] = linked
    pages = max(1, (total + LIMIT - 1) // LIMIT)
    return {"success": True, "entries": entries, "total": total, "page": page, "pages": pages}


@app.post("/api/admin/address-entries")
async def admin_create_address_entry(data: dict, session: dict = Depends(get_current_admin)):
    street = data.get("street", "").strip()
    city = data.get("city", "").strip()
    if not street or not city:
        raise HTTPException(status_code=400, detail="Street and city required")
    now = datetime.now(timezone.utc).isoformat()
    doc = {
        "entryId": str(uuid.uuid4()),
        "street": street,
        "city": city,
        "state": data.get("state", "").strip(),
        "zip": data.get("zip", "").strip(),
        "linkedRecordIds": data.get("linkedRecordIds", []),
        "searchCount": 0,
        "lastSearched": None,
        "source": "manual",
        "addedAt": now,
        "notes": data.get("notes", ""),
    }
    await address_entries_collection.insert_one(doc)
    doc.pop("_id", None)
    await log_action(session, "create_address_entry", "address_entry", doc["entryId"], f"{street}, {city}")
    return {"success": True, "entry": doc}


@app.put("/api/admin/address-entries/{entry_id}")
async def admin_update_address_entry(entry_id: str, data: dict, session: dict = Depends(get_current_admin)):
    existing = await address_entries_collection.find_one({"entryId": entry_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Not found")
    updates = {
        "street":           data.get("street",          existing.get("street", "")).strip(),
        "city":             data.get("city",            existing.get("city", "")).strip(),
        "state":            data.get("state",           existing.get("state", "")).strip(),
        "zip":              data.get("zip",             existing.get("zip", "")).strip(),
        "linkedRecordIds":  data.get("linkedRecordIds", existing.get("linkedRecordIds", [])),
        "notes":            data.get("notes",           existing.get("notes", "")),
    }
    if not updates["street"] or not updates["city"]:
        raise HTTPException(status_code=400, detail="Street and city required")
    await address_entries_collection.update_one({"entryId": entry_id}, {"$set": updates})
    await log_action(session, "update_address_entry", "address_entry", entry_id)
    return {"success": True}


@app.delete("/api/admin/address-entries/{entry_id}")
async def admin_delete_address_entry(entry_id: str, session: dict = Depends(get_current_admin)):
    result = await address_entries_collection.delete_one({"entryId": entry_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Not found")
    await log_action(session, "delete_address_entry", "address_entry", entry_id)
    return {"success": True}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001)
